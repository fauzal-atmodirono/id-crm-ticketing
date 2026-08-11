#!/usr/bin/env python3
"""PROTON CRM Feature Guide — operator handbook builder.

Renders the 14 hand-drafted markdown chapters in `feature-guide-src-v3/`
(01-introduction.md ... 14-glossary.md) into the client's Google Docs
template (`Google Docs template - Short version.docx`), producing
`PROTON - CRM Feature Guide v3.docx`.

`SRC_DIR`/`OUT`/`COVER_TITLE`/`COVER_SUBTITLE` are overridable via the
`FG_SRC_DIR`/`FG_OUT`/`FG_COVER_TITLE`/`FG_COVER_SUBTITLE` env vars, so a
further edition can be built without editing this file. The v1 and v2
sources and outputs live under `archive/`; to rebuild one, point
`FG_SRC_DIR`/`FG_OUT` at the archived paths.

The markdown subset implemented here is deliberately narrow — it's exactly
what the chapter files use, no more:
  - `#`.."####" headings -> Heading 1..4
  - paragraphs with inline **bold**, *italic*, `code` (code -> Courier New)
  - `-` bullets, one level of nesting -> List Bullet / List Bullet 2
    (falls back to a manual indent + bullet character if those styles
    aren't defined in the template — which is the case here)
  - `1.` numbered lists -> List Number (same fallback: manual "N. " prefix)
  - pipe tables with a `|---|` separator row -> a bordered table, bold header
  - `> ` blockquotes -> a left-bordered, lightly shaded, italic paragraph
  - `[[SCREENSHOT: id | caption]]` on its own line -> the PNG at
    feature-guide-assets/<id>.png if present, else a bordered placeholder
    box, so the build never fails on a missing screenshot
  - `<!-- ... -->` HTML comments (including inline trailing ones) are
    stripped entirely, whether they sit on their own line or trail real
    content

## Audience filtering — the three training curricula (§2.3.3)

The same source also renders the three role curricula in `training/`.
A `##` section (or a whole chapter) declares its audience with an HTML
comment marker:

    ## Labels
    <!-- TRAINING: audience=agent, exercise -->

The three audiences are cumulative — `agent` < `supervisor` < `admin` —
so a section names the most junior audience that needs it and every more
senior audience inherits it. `exercise` marks a section whose documented
steps are a hands-on lab task.

    python3 build_crm_feature_guide.py                  # the handbook, unchanged
    python3 build_crm_feature_guide.py --curricula      # write training/
    python3 build_crm_feature_guide.py --check          # exit 1 if training/ is stale
    python3 build_crm_feature_guide.py --audience agent # a role-scoped .docx

The markers are HTML comments, which this builder already strips from
every line, so **the default build is unaffected by them** — that is the
whole safety argument for putting them in the chapter files, and
`scripts/test_build_feature_guide_audiences.py` asserts it by rebuilding
the default output with the pre-audience-filter generator and comparing.

Plain stdlib + python-docx script, constants at the top; the only CLI args
are the audience/curricula switches described above, so the default
invocation stays the argument-free build it has always been.
"""
import argparse
import glob
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(BASE, "Google Docs template - Short version.docx")
ASSETS_DIR = os.path.join(BASE, "feature-guide-assets")

# SRC_DIR / OUT / COVER_TITLE / COVER_SUBTITLE are overridable by environment
# variable so a further edition living in its own source dir can be built
# without touching this script. The defaults track the CURRENT edition (v3);
# v1/v2 moved under archive/ on 2026-08-09 and are rebuilt by pointing
# FG_SRC_DIR/FG_OUT at those archived paths.
SRC_DIR = os.environ.get("FG_SRC_DIR", os.path.join(BASE, "feature-guide-src-v3"))
OUT = os.environ.get(
    "FG_OUT", os.path.join(BASE, "PROTON - CRM Feature Guide v3.docx")
)

COVER_TITLE = os.environ.get(
    "FG_COVER_TITLE", "PROTON e.MAS — CRM Feature Guide"
)
COVER_SUBTITLE = os.environ.get(
    "FG_COVER_SUBTITLE", "Operator Handbook — August 2026"
)

BORDER_COLOR = "4472C4"  # a muted blue, used for table borders and the
                          # blockquote/placeholder left accent
NOTE_SHADE = "F2F2F2"    # light grey shading for blockquote paragraphs

TRAINING_DIR = os.path.join(BASE, "training")

# ---------------------------------------------------------------------------
# Audience tagging — one source, three curricula (§2.3.3)
# ---------------------------------------------------------------------------
# CUMULATIVE and in this order, because the design's role table defines each
# senior role as "the above, plus ...": a supervisor is taught everything an
# agent is taught, an administrator everything a supervisor is. So a section
# declares the MOST JUNIOR audience that needs it and the seniors inherit it.
# One token per section beats a three-way list that nobody keeps in step.
AUDIENCES = ("agent", "supervisor", "admin")

AUDIENCE_LABELS = {
    "agent": "Frontline agent",
    "supervisor": "Supervisor / team leader",
    "admin": "Administrator",
}

# Design targets from the spec's §3.1 table, in minutes. Used only to state
# how far the curriculum as scoped is from the length the design assumed —
# never to scale the derived durations into agreement with it.
AUDIENCE_TARGET_MINUTES = {"agent": 120, "supervisor": 180, "admin": 240}

# Exercise-id prefix per audience, so a facilitator can say "AG-04" without
# ambiguity about which cohort's sheet it is on.
AUDIENCE_EXERCISE_PREFIX = {"agent": "AG", "supervisor": "SV", "admin": "AD"}

# A section with no marker of its own, in a chapter with no marker either,
# lands here. `admin` rather than `agent` deliberately: the fallback must
# never quietly teach a frontline cohort a page they cannot open, and admin
# is the widest curriculum, so untagged content still appears SOMEWHERE
# instead of vanishing from all three. Every fallback is counted and named
# in training/tag-coverage.md, so "it defaulted" is visible, not assumed.
FALLBACK_AUDIENCE = "admin"

TRAINING_RE = re.compile(r"^\s*<!--\s*TRAINING:(?P<body>[^>]*?)-->\s*$")

# ---------------------------------------------------------------------------
# Markdown line patterns
# ---------------------------------------------------------------------------
COMMENT_RE = re.compile(r"<!--.*?-->")
HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
SCREENSHOT_RE = re.compile(r"^\[\[SCREENSHOT:\s*(.+?)\s*\|\s*(.+?)\]\]\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{1,}:?\s*(\|\s*:?-{1,}:?\s*)+\|?\s*$")
BULLET_RE = re.compile(r"^(\s*)-\s+(.*)$")
NUMBER_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
INLINE_RE = re.compile(
    r"(?P<bold>\*\*(?P<boldtext>.+?)\*\*)"
    r"|(?P<italic>\*(?P<italictext>.+?)\*)"
    r"|(?P<code>`(?P<codetext>.+?)`)"
)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
def style_exists(document, name):
    try:
        document.styles[name]
        return True
    except KeyError:
        return False


def sanitize_page_margins(document):
    """The template's pgMar twip values are non-integers (a Google Docs
    export quirk, e.g. "1440.0000000000002"), which crashes python-docx's
    integer-only Twips parser the first time anything (like add_table)
    reads section.left_margin. Round them to whole twips in place — this
    doesn't perceptibly change the page setup, just makes it parseable."""
    body = document.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        return
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is None:
        return
    for attr in ("top", "right", "bottom", "left", "header", "footer", "gutter"):
        key = qn("w:%s" % attr)
        val = pgMar.get(key)
        if val is not None:
            try:
                pgMar.set(key, str(int(round(float(val)))))
            except ValueError:
                pass


def clear_body(document):
    """Delete every existing body paragraph/table, keep page setup (sectPr)."""
    body = document.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_borders(table, color=BORDER_COLOR, sz=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:%s" % edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def add_left_border_and_shade(paragraph, color=BORDER_COLOR, fill=NOTE_SHADE):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def scan_headings(chapter_texts):
    """Collect every chapter (#) and section (##) heading, in document order,
    pairing each with the bookmark name the TOC will link to.

    Deliberately a second pass over the same text rather than a hook inside
    process_chapter: the table of contents is written before any chapter is
    rendered, so the full list has to exist first. Takes the already-read
    chapter text (not paths) so that an audience-filtered build lists exactly
    the sections it goes on to render — a TOC scanned from the unfiltered
    file would link to headings that are not in the document.

    Returns [(level, text, bookmark_name)]. Bookmark names use Word's reserved
    `_Toc` prefix and are pure ASCII, because a heading's own text may contain
    em dashes, ampersands and non-breaking punctuation that are not legal in a
    bookmark name."""
    entries = []
    for text in chapter_texts:
        for raw in text.split("\n"):
            line = COMMENT_RE.sub("", raw).rstrip()
            m = HEADING_RE.match(line)
            if not m or len(m.group(1)) > 2:
                continue
            entries.append(
                (len(m.group(1)), m.group(2).strip(), "_Toc%05d" % len(entries))
            )
    return entries


def add_bookmark(paragraph, name, bookmark_id):
    """Wrap a paragraph in a bookmarkStart/bookmarkEnd pair so a TOC entry can
    link to it."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_toc_entry(document, level, text, anchor):
    """One line of the table of contents: an internal hyperlink to `anchor`.

    A static, pre-rendered TOC rather than a Word `TOC` field. A field is only
    text once an application evaluates it — Word does that on F9, but Google
    Docs cannot evaluate Word fields at all, and on import it degrades the
    field into raw heading anchors (`?tab=t.0#heading=...`), which is what the
    reader then sees instead of a contents list. Real paragraphs with real
    internal links render identically in Word, Google Docs, LibreOffice and
    any PDF export, and need no "update field" step from the reader.

    The trade is page numbers: computing them needs a layout engine, which
    python-docx is not. Clickable entries are the better half of that trade
    for a document read on screen."""
    paragraph = document.add_paragraph(style="normal")
    fmt = paragraph.paragraph_format
    fmt.left_indent = Inches(0.0 if level == 1 else 0.35)
    fmt.space_after = Pt(2)

    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if level == 1:
        bold = OxmlElement("w:b")
        rPr.append(bold)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC" if level == 2 else "000000")
    rPr.append(color)
    if level == 2:
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rPr.append(underline)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)

    link.append(run)
    paragraph._p.append(link)
    return paragraph


# ---------------------------------------------------------------------------
# Inline formatting (bold / italic / code)
# ---------------------------------------------------------------------------
def add_inline_runs(paragraph, text, base_italic=False, base_bold=False):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            run.bold = base_bold or None
            run.italic = base_italic or None
        if m.group("bold") is not None:
            run = paragraph.add_run(m.group("boldtext"))
            run.bold = True
            run.italic = base_italic or None
        elif m.group("italic") is not None:
            run = paragraph.add_run(m.group("italictext"))
            run.italic = True
            run.bold = base_bold or None
        elif m.group("code") is not None:
            run = paragraph.add_run(m.group("codetext"))
            run.font.name = "Courier New"
            run.bold = base_bold or None
            run.italic = base_italic or None
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold or None
        run.italic = base_italic or None


# ---------------------------------------------------------------------------
# Screenshots
# ---------------------------------------------------------------------------
def add_screenshot(document, shot_id, caption, found, missing):
    png_path = os.path.join(ASSETS_DIR, "%s.png" % shot_id)
    if os.path.exists(png_path):
        found.append(shot_id)
        document.add_picture(png_path, width=Inches(6))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = document.add_paragraph(style="normal")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
    else:
        missing.append(shot_id)
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, sz=8)
        cell = table.cell(0, 0)
        set_cell_shading(cell, NOTE_SHADE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Screenshot: %s" % caption)
        run.italic = True
        # give the placeholder some visible height/width
        cell.width = Inches(6)
        for _ in range(3):
            cell.add_paragraph("")
    document.add_paragraph(style="normal")


# ---------------------------------------------------------------------------
# Markdown block parsing
# ---------------------------------------------------------------------------
def split_into_blocks(text):
    cleaned = []
    for line in text.split("\n"):
        cleaned.append(COMMENT_RE.sub("", line).rstrip())

    blocks = []
    current = []
    for line in cleaned:
        if line.strip() == "":
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(line)
    if current:
        blocks.append(current)
    return blocks


def classify_block(block):
    first = block[0]

    m = HEADING_RE.match(first)
    if m and len(block) == 1:
        return ("heading", len(m.group(1)), m.group(2).strip())

    m = SCREENSHOT_RE.match(first)
    if m and len(block) == 1:
        return ("screenshot", m.group(1).strip(), m.group(2).strip())

    if first.lstrip().startswith("|") and len(block) >= 2 and TABLE_SEP_RE.match(
        block[1]
    ):
        rows = [block[0]] + block[2:]
        parsed_rows = []
        for row in rows:
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            parsed_rows.append(cells)
        return ("table", parsed_rows)

    if all(line.lstrip().startswith(">") for line in block):
        quote_lines = [re.sub(r"^\s*>\s?", "", line) for line in block]
        return ("blockquote", " ".join(l.strip() for l in quote_lines))

    if BULLET_RE.match(first) or NUMBER_RE.match(first):
        items = []
        for line in block:
            bm = BULLET_RE.match(line)
            nm = NUMBER_RE.match(line)
            if bm is not None:
                indent, rest = bm.group(1), bm.group(2)
                level = 1 if len(indent) > 0 else 0
                items.append({"kind": "bullet", "level": level, "text": rest})
            elif nm is not None:
                indent, num, rest = nm.group(1), nm.group(2), nm.group(3)
                items.append(
                    {"kind": "number", "level": 0, "num": num, "text": rest}
                )
            else:
                # continuation line of the previous item's wrapped text
                if items:
                    items[-1]["text"] += " " + line.strip()
        return ("list", items)

    text = " ".join(line.strip() for line in block)
    return ("paragraph", text)


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------
def render_heading(document, level, text, bookmarks=None):
    heading = document.add_paragraph(style="Heading %d" % level)
    add_inline_runs(heading, text)
    # Anchor for the matching table-of-contents entry. Looked up by
    # (level, text) and popped, rather than by position, so that a heading
    # the block parser and the TOC pre-scan disagree about can only lose its
    # link — never silently hand its bookmark to a different heading.
    if bookmarks:
        pending = bookmarks.get((level, text))
        if pending:
            name = pending.pop(0)
            add_bookmark(heading, name, abs(hash(name)) % 900000 + 1000)
    return heading


def render_paragraph(document, text):
    p = document.add_paragraph(style="normal")
    add_inline_runs(p, text)


def render_blockquote(document, text):
    p = document.add_paragraph(style="normal")
    add_left_border_and_shade(p)
    p.paragraph_format.left_indent = Inches(0.3)
    add_inline_runs(p, text, base_italic=True)


def render_list(document, items, use_bullet_style, use_number_style):
    for item in items:
        level = item.get("level", 0)
        if item["kind"] == "bullet":
            style_name = "List Bullet" if level == 0 else "List Bullet 2"
            if use_bullet_style:
                p = document.add_paragraph(style=style_name)
                add_inline_runs(p, item["text"])
            else:
                p = document.add_paragraph(style="normal")
                p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                bullet_char = "•" if level == 0 else "◦"
                run = p.add_run("%s  " % bullet_char)
                add_inline_runs(p, item["text"])
        else:  # number
            if use_number_style:
                p = document.add_paragraph(style="List Number")
                add_inline_runs(p, item["text"])
            else:
                p = document.add_paragraph(style="normal")
                p.paragraph_format.left_indent = Inches(0.25)
                p.add_run("%s.  " % item["num"])
                add_inline_runs(p, item["text"])


def render_table(document, rows):
    n_cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=n_cols)
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, text, base_bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, NOTE_SHADE)
    document.add_paragraph(style="normal")
    return n_cols


# ---------------------------------------------------------------------------
# Chapter processing
# ---------------------------------------------------------------------------
def process_chapter(document, text, use_bullet_style, use_number_style, stats,
                    bookmarks=None):
    blocks = [classify_block(b) for b in split_into_blocks(text)]

    for block in blocks:
        kind = block[0]
        if kind == "heading":
            _, level, heading_text = block
            render_heading(document, level, heading_text, bookmarks)
            if level == 2:
                stats["sections"] += 1
        elif kind == "paragraph":
            render_paragraph(document, block[1])
        elif kind == "blockquote":
            render_blockquote(document, block[1])
        elif kind == "list":
            render_list(document, block[1], use_bullet_style, use_number_style)
        elif kind == "table":
            n_cols = render_table(document, block[1])
            stats["tables"] += 1
        elif kind == "screenshot":
            _, shot_id, caption = block
            add_screenshot(document, shot_id, caption, stats["found"], stats["missing"])


# ---------------------------------------------------------------------------
# Audience markers: parsing and validation
# ---------------------------------------------------------------------------
class TrainingTagError(Exception):
    """A malformed, duplicated or misspelled TRAINING marker.

    Raised rather than warned, and raised in EVERY mode including the plain
    handbook build. A typo'd audience name would otherwise drop its section
    from all three curricula and produce a quietly thinner deck — which is
    the exact failure this mechanism exists to prevent, and the one a
    facilitator would discover in front of a cohort. Validation emits
    nothing, so making it unconditional cannot change the handbook output.
    """


def parse_training_marker(body, path, lineno):
    """Parse one marker body into (audience, is_exercise).

    Grammar is deliberately tiny: comma-separated tokens, exactly one
    `audience=<name>`, optionally `exercise`. Anything else is an error
    naming the file, the line and the valid alternatives."""
    where = "%s:%d" % (os.path.basename(path), lineno)
    audience = None
    exercise = False
    for token in [t.strip() for t in body.split(",") if t.strip()]:
        if token.startswith("audience="):
            name = token[len("audience="):].strip()
            if name not in AUDIENCES:
                raise TrainingTagError(
                    "%s: unknown audience %r in TRAINING marker; valid names "
                    "are %s" % (where, name, ", ".join(AUDIENCES))
                )
            if audience is not None:
                raise TrainingTagError(
                    "%s: TRAINING marker names more than one audience; a "
                    "section declares only its most junior audience, and the "
                    "seniors inherit it" % where
                )
            audience = name
        elif token == "exercise":
            exercise = True
        else:
            raise TrainingTagError(
                "%s: unknown token %r in TRAINING marker; expected "
                "'audience=<%s>' and optionally 'exercise'"
                % (where, token, "|".join(AUDIENCES))
            )
    if audience is None:
        raise TrainingTagError(
            "%s: TRAINING marker has no 'audience=' token" % where
        )
    return audience, exercise


def parse_chapter_sections(text, path):
    """Split one chapter into its `##` sections and resolve each audience.

    Returns (chapter_title, preamble_lines, sections). Each section is a dict
    with title / audience / audience_source / exercise / lines (its own body,
    `##` heading first) / heading_line.

    A marker before the first `##` sets the chapter default — which is what
    makes a *new* section inherit a sensible audience instead of falling all
    the way through to FALLBACK_AUDIENCE.

    Marker lines are KEPT in the returned body. They are HTML comments, so no
    renderer sees them, and keeping them means audience-filtered source text
    re-parses to the same audiences as the file it came from — a filtered
    chapter that had lost its chapter-level marker would silently re-read as
    FALLBACK_AUDIENCE for every section in it."""
    lines = text.split("\n")
    chapter_title = None
    chapter_default = None
    chapter_default_line = None
    preamble = []
    sections = []

    for lineno, raw in enumerate(lines, start=1):
        marker = TRAINING_RE.match(raw)
        if marker is not None:
            audience, exercise = parse_training_marker(
                marker.group("body"), path, lineno
            )
            if not sections:
                if exercise:
                    raise TrainingTagError(
                        "%s:%d: 'exercise' is a per-section flag; a "
                        "chapter-level TRAINING marker may only set audience="
                        % (os.path.basename(path), lineno)
                    )
                if chapter_default is not None:
                    raise TrainingTagError(
                        "%s:%d: second chapter-level TRAINING marker (the "
                        "first is on line %d)"
                        % (os.path.basename(path), lineno, chapter_default_line)
                    )
                chapter_default = audience
                chapter_default_line = lineno
            else:
                section = sections[-1]
                if section["marker_line"] is not None:
                    raise TrainingTagError(
                        "%s:%d: section %r already has a TRAINING marker on "
                        "line %d"
                        % (
                            os.path.basename(path),
                            lineno,
                            section["title"],
                            section["marker_line"],
                        )
                    )
                section["audience"] = audience
                section["audience_source"] = "section"
                section["exercise"] = exercise
                section["marker_line"] = lineno
            if sections:
                sections[-1]["lines"].append(raw)
            else:
                preamble.append(raw)
            continue

        heading = HEADING_RE.match(COMMENT_RE.sub("", raw).rstrip())
        if heading is not None and len(heading.group(1)) == 1:
            if chapter_title is None:
                chapter_title = heading.group(2).strip()
        if heading is not None and len(heading.group(1)) == 2:
            sections.append(
                {
                    "title": heading.group(2).strip(),
                    "heading_line": lineno,
                    "lines": [],
                    "audience": None,
                    "audience_source": None,
                    "exercise": False,
                    "marker_line": None,
                }
            )
        if sections:
            sections[-1]["lines"].append(raw)
        else:
            preamble.append(raw)

    for section in sections:
        if section["audience"] is None:
            if chapter_default is not None:
                section["audience"] = chapter_default
                section["audience_source"] = "chapter"
            else:
                section["audience"] = FALLBACK_AUDIENCE
                section["audience_source"] = "fallback"

    return chapter_title, preamble, sections


def sees(section_audience, target):
    """Cumulative membership: does `target` get taught this section?"""
    return AUDIENCES.index(section_audience) <= AUDIENCES.index(target)


def read_chapter(path, audience=None):
    """The chapter source, filtered to one audience if asked.

    `audience=None` returns the file verbatim — same string the builder has
    always read — so the default handbook build cannot be perturbed by
    anything in this module. A chapter with no sections for the audience
    returns "" and its caller drops it entirely rather than emitting a
    chapter title with nothing under it."""
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    if audience is None:
        return text
    _, preamble, sections = parse_chapter_sections(text, path)
    kept = [s for s in sections if sees(s["audience"], audience)]
    if not kept:
        return ""
    out = list(preamble)
    for section in kept:
        out.extend(section["lines"])
    return "\n".join(out)


# ---------------------------------------------------------------------------
# Curricula: extracting teachable content from a section
# ---------------------------------------------------------------------------
# Everything a slide, an exercise or a checklist row says comes from the
# handbook section it points at. Nothing here writes operator-facing prose:
# where the source has no summary paragraph or no documented steps, the
# generated file says so, because a curriculum that reads complete while
# teaching from nothing is worse than one that admits the gap.
SUBHEADING_RE = re.compile(r"^###\s+(.+)$")


def split_subsections(body_lines):
    """Group a section body by its `###` sub-headings.

    Returns {name: [lines]}, with the pre-first-`###` remainder under "".
    Chapters 11/12/14 don't use the five-heading template, so that "" bucket
    is what their content lands in."""
    buckets = {"": []}
    current = ""
    for raw in body_lines:
        m = SUBHEADING_RE.match(COMMENT_RE.sub("", raw).strip())
        if m is not None:
            current = m.group(1).strip()
            buckets.setdefault(current, [])
            continue
        buckets[current].append(raw)
    return buckets


def blocks_of(lines):
    return [classify_block(b) for b in split_into_blocks("\n".join(lines))]


def first_paragraph(lines):
    for block in blocks_of(lines):
        if block[0] == "paragraph":
            return block[1]
    return None


def numbered_steps(lines):
    """The first list block's item texts — the documented procedure.

    Bullets count as well as numbers: chapter 12's playbooks use bullets for
    the same "do this, then this" content that chapter 2 numbers."""
    for block in blocks_of(lines):
        if block[0] == "list":
            return [item["text"] for item in block[1]]
    return []


def blockquotes_of(lines):
    return [block[1] for block in blocks_of(lines) if block[0] == "blockquote"]


def table_row_count(lines):
    return sum(
        len(block[1]) - 1 for block in blocks_of(lines) if block[0] == "table"
    )


def slide_minutes(section):
    """DERIVED BY RULE, never measured. Stated as such in every output.

    3 minutes to introduce a topic, one more per two documented steps to
    walk it through, five more if the cohort does it themselves. No delivered
    session has been timed, so a figure that looked measured would be a
    fabricated measurement."""
    minutes = 3 + (len(section["steps"]) + 1) // 2
    if section["exercise"]:
        minutes += 5
    return minutes


def load_curriculum_model():
    """Parse every chapter once into the model all four outputs render from."""
    chapters = []
    for path in chapter_source_paths():
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        title, _, sections = parse_chapter_sections(text, path)
        basename = os.path.basename(path)
        for section in sections:
            body = section["lines"][1:]
            buckets = split_subsections(body)
            fallback = buckets.get("", [])
            what = first_paragraph(buckets.get("What it is", fallback))
            where = first_paragraph(buckets.get("Where to find it", []))
            # Two step sources, and WHICH ONE matters, because it decides what
            # kind of exercise the section can honestly support:
            #   how-to — the five-heading template's imperative procedure, so
            #            a trainee can be asked to carry it out;
            #   body   — chapter 11's scenario narratives, which are told from
            #            the customer's side as much as the agent's, so they
            #            are a role-play with a facilitator, not a checklist.
            # Anything else (chapter 12's per-channel `###` scenarios, chapter
            # 14's glossary table) has no step list at all, and gets none
            # invented: the slide sends the facilitator to the section.
            steps = numbered_steps(buckets.get("How to use it", []))
            steps_source = "how-to" if steps else None
            if not steps:
                steps = numbered_steps(fallback)
                steps_source = "body" if steps else None
            section["steps_source"] = steps_source
            section["what"] = what
            section["where"] = where
            section["steps"] = steps
            section["notes"] = blockquotes_of(body)
            section["table_rows"] = table_row_count(body)
            section["chapter_file"] = basename
        chapters.append(
            {
                "path": path,
                "file": basename,
                "number": basename[:2],
                "title": title or basename,
                "sections": sections,
            }
        )
    return chapters


# ---------------------------------------------------------------------------
# Curricula: what the source cannot teach yet
# ---------------------------------------------------------------------------
# These live in the generator, not in the markdown, so a regeneration cannot
# silently drop them — the same reasoning as the caveat list in
# scripts/generate-config-doc.py. Each row states why the topic is absent
# from the handbook source, so nobody "fixes" the curriculum by writing the
# missing section from a spec instead of from a running tenant.
CURRICULUM_GAPS = (
    {
        "topic": "Agent availability and the workforce dashboard",
        "audiences": ("supervisor", "admin"),
        "why": (
            "Held out of the handbook source on 2026-08-09 because fork "
            "patches `0053`/`0054` have never been built into an image — "
            "\"My status\" and \"Workforce\" do not appear in the deployed "
            "JS bundle. The written section is parked in "
            "`feature-guide-v3-pending.md`."
        ),
        "unblocks": "P6 · a Cloud Build of patches 0053+0054",
    },
    {
        "topic": "Performance targets and attainment",
        "audiences": ("supervisor", "admin"),
        "why": (
            "P5 built a targets store and an attainment view, and the "
            "deployed backend's own OpenAPI document has no "
            "`/metrics/targets`. There is no handbook section, and inventing "
            "one would teach a page no supervisor can open."
        ),
        "unblocks": "P5 · backend rebuilt past `e6dc537`, then a handbook section",
    },
    {
        "topic": "Alert preferences / inbound alerts",
        "audiences": ("admin",),
        "why": (
            "The admin surface is fork patch `0057`, unbuilt. The feature is "
            "also behind two independent switches (blocked-work register "
            "§3h, §3i), so even once the patch ships, what an operator sees "
            "depends on a second gate the UI does not mention."
        ),
        "unblocks": "P9 · a Cloud Build of patch 0057, plus both switches on",
    },
    {
        "topic": "Case taxonomy administration",
        "audiences": ("admin",),
        "why": (
            "The five-field taxonomy is taught from the agent's side "
            "(chapter 5) because that is what exists on the tenant. The "
            "admin page that edits the taxonomy is fork patch `0060`, "
            "unbuilt (blocked-work register §3m), so today an administrator "
            "still edits those lists as Custom Attributes."
        ),
        "unblocks": "P10 · a Cloud Build of patch 0060",
    },
    {
        "topic": "The redesigned Roles & Permissions page",
        "audiences": ("admin",),
        "why": (
            "The handbook documents the page as it exists on the tenant "
            "(patches `0027`/`0028`). Patch `0059`'s redesign is unbuilt, so "
            "the admin curriculum teaches the current page and must be "
            "regenerated when `0059` ships."
        ),
        "unblocks": "P10 · a Cloud Build of patch 0059, then a handbook update",
    },
    {
        "topic": "Data scopes (row-level data access)",
        "audiences": ("admin",),
        "why": (
            "**Deliberately not taught, and not merely missing.** "
            "`DATA_SCOPED_RBAC_ENABLED` restricts nothing: the scope logic "
            "has no caller and no query applies it (risk R16, blocked-work "
            "register §3j). Teaching an administrator to rely on it would "
            "teach a control that does not exist."
        ),
        "unblocks": "Enforcement wiring — not a documentation task",
    },
    {
        "topic": "AI conversational quality (Translate, FAQ composer)",
        "audiences": ("admin",),
        "why": (
            "Held out of the handbook source: patches `0055`/`0056` unbuilt "
            "and the deployed backend has no `/assist/translate`. Parked in "
            "`feature-guide-v3-pending.md`."
        ),
        "unblocks": "P7 · Cloud Build of 0055+0056, backend rebuilt",
    },
    {
        "topic": "AI cost and performance measurement",
        "audiences": ("admin",),
        "why": (
            "Held out of the handbook source: the eleven BigQuery views were "
            "never created, so the reports have nothing to read even once "
            "the code ships. Parked in `feature-guide-v3-pending.md`."
        ),
        "unblocks": "P8 · backend rebuilt plus `ensure_views()` run",
    },
    {
        "topic": "Hands-on voice and phone practice",
        "audiences": ("agent", "supervisor", "admin"),
        "why": (
            "The voice and phone topics are taught from the handbook, but no "
            "real Twilio call has ever been placed (risk R10) and every "
            "`PHONE_*` capability switch is off on the tenant. The channel "
            "topics are therefore presentation-only: there is no exercise "
            "for them, and inventing one would be a lab that cannot run."
        ),
        "unblocks": "R10 · one real call, then a sandbox phone number",
    },
)


# ---------------------------------------------------------------------------
# Curricula: rendering
# ---------------------------------------------------------------------------
GENERATED_BANNER = (
    "<!-- GENERATED FILE — do not edit by hand.\n"
    "     Source: docs/client-materials/feature-guide-src-v3/ "
    "(the operator handbook)\n"
    "     Regenerate: python3 docs/client-materials/build_crm_feature_guide.py"
    " --curricula\n"
    "     Drift check: python3 docs/client-materials/build_crm_feature_guide.py"
    " --check\n"
    "-->\n"
)

GENERATED_NOTE = (
    "> **Generated from the operator handbook — do not edit.** Every line "
    "below is rendered from `feature-guide-src-v3/`; an edit here is "
    "overwritten by the next run. To change what a cohort is taught, change "
    "the handbook section this points at, or its `<!-- TRAINING: ... -->` "
    "marker, and regenerate.\n"
)


def duration_text(minutes):
    hours, mins = divmod(minutes, 60)
    if hours and mins:
        return "%d h %d min" % (hours, mins)
    if hours:
        return "%d h" % hours
    return "%d min" % mins


def audience_sections(chapters, audience):
    """[(chapter, section)] for one audience, in document order."""
    pairs = []
    for chapter in chapters:
        for section in chapter["sections"]:
            if sees(section["audience"], audience):
                pairs.append((chapter, section))
    return pairs


def exercise_ids(chapters, audience):
    """{(file, title): id} for the exercise-flagged sections of one audience."""
    prefix = AUDIENCE_EXERCISE_PREFIX[audience]
    ids = {}
    for chapter, section in audience_sections(chapters, audience):
        if section["exercise"]:
            ids[(chapter["file"], section["title"])] = "%s-%02d" % (
                prefix,
                len(ids) + 1,
            )
    return ids


def gaps_for(audience):
    return [gap for gap in CURRICULUM_GAPS if audience in gap["audiences"]]


def render_gap_section(audience, heading_level="##"):
    gaps = gaps_for(audience)
    lines = ["%s What this curriculum cannot teach yet" % heading_level, ""]
    lines.append(
        "%d topic%s this role would be expected to cover %s absent from the "
        "handbook source, or present only as the tenant's current behaviour. "
        "They are listed here rather than written from a specification, "
        "because a curriculum that teaches an unbuilt page loses its cohort "
        "on day one."
        % (len(gaps), "" if len(gaps) == 1 else "s", "is" if len(gaps) == 1 else "are")
    )
    lines.append("")
    lines.append("| Topic | Why it is not here | Unblocked by |")
    lines.append("|---|---|---|")
    for gap in gaps:
        lines.append(
            "| %s | %s | %s |" % (gap["topic"], gap["why"], gap["unblocks"])
        )
    lines.append("")
    return lines


def render_deck(chapters, audience):
    pairs = audience_sections(chapters, audience)
    ids = exercise_ids(chapters, audience)
    total = sum(slide_minutes(section) for _, section in pairs)
    target = AUDIENCE_TARGET_MINUTES[audience]

    out = [GENERATED_BANNER]
    out.append(
        "# %s curriculum — facilitator deck\n" % AUDIENCE_LABELS[audience]
    )
    out.append(GENERATED_NOTE)
    out.append(
        "**Audience:** %s · **Topics:** %d of %d handbook sections · "
        "**Hands-on exercises:** %d\n"
        % (
            AUDIENCE_LABELS[audience],
            len(pairs),
            sum(len(c["sections"]) for c in chapters),
            len(ids),
        )
    )
    out.append(
        "**Rule-derived length:** %s. **Design target (spec §3.1):** %s. "
        "**Difference: %s%s.**\n"
        % (
            duration_text(total),
            duration_text(target),
            "+" if total >= target else "-",
            duration_text(abs(total - target)),
        )
    )
    out.append(
        "> Durations are **derived by rule** — 3 min per topic, plus 1 min "
        "per two documented steps, plus 5 min where the cohort does it "
        "themselves — and are **not measured**: no session has been "
        "delivered or timed. Where the derived length exceeds the design "
        "target, either the target or the topic list has to move; the "
        "generator will not scale one into agreement with the other.\n"
    )
    out.append("## How to run a slide\n")
    out.append(
        "Each topic below is one slide. **Say** is the handbook section's "
        "opening paragraph — read the whole section before delivering it. "
        "**Show** is where the feature lives in the CRM. **Walk through** is "
        "the documented procedure, verbatim, so a demo cannot drift from the "
        "handbook the cohort takes away. **Say out loud** carries the "
        "section's own caveats; skipping those is how a cohort learns a "
        "limitation from a customer instead.\n"
    )

    current_file = None
    for chapter, section in pairs:
        if chapter["file"] != current_file:
            current_file = chapter["file"]
            chapter_pairs = [p for p in pairs if p[0]["file"] == current_file]
            chapter_minutes = sum(slide_minutes(s) for _, s in chapter_pairs)
            out.append(
                "## Module %s — %s  ·  %d topics  ·  %s\n"
                % (
                    chapter["number"],
                    chapter["title"],
                    len(chapter_pairs),
                    duration_text(chapter_minutes),
                )
            )
        exercise_id = ids.get((chapter["file"], section["title"]))
        out.append(
            "### %s  ·  %s  ·  %s+\n"
            % (
                section["title"],
                duration_text(slide_minutes(section)),
                section["audience"],
            )
        )
        out.append(
            "**Source:** `%s` → `## %s`\n"
            % (chapter["file"], section["title"])
        )
        if section["what"]:
            out.append("**Say:** %s\n" % section["what"])
        elif section["table_rows"]:
            out.append(
                "**Say:** _(reference table of %d rows — hand it out rather "
                "than present it)_\n" % section["table_rows"]
            )
        else:
            out.append(
                "**Say:** _(the handbook section has no summary paragraph — "
                "the steps below are the whole of it)_\n"
            )
        if section["where"]:
            out.append("**Show:** %s\n" % section["where"])
        if section["steps"] and section["steps_source"] == "how-to":
            out.append("**Walk through:**\n")
            for i, step in enumerate(section["steps"], start=1):
                out.append("%d. %s" % (i, step))
            out.append("")
        elif section["steps"]:
            out.append("**Walk the scenario through, in order:**\n")
            for i, step in enumerate(section["steps"], start=1):
                out.append("%d. %s" % (i, step))
            out.append("")
        elif not section["table_rows"]:
            out.append(
                "**No single procedure to demo:** this section is structured "
                "as sub-scenarios rather than one set of steps. Deliver it "
                "from the handbook section itself.\n"
            )
        for note in section["notes"]:
            out.append("**Say out loud:** %s\n" % note)
        if exercise_id:
            out.append(
                "**Hands-on:** exercise `%s` — see `exercises.md`.\n"
                % exercise_id
            )

    out.extend(render_gap_section(audience))
    out.append(
        "See `../delivery-plan.md` for the session schedule, prerequisites, "
        "the sandbox reset between cohorts and the refresher cadence.\n"
    )
    return "\n".join(out).rstrip() + "\n"


def render_exercises(chapters, audience):
    ids = exercise_ids(chapters, audience)
    pairs = [
        (c, s)
        for c, s in audience_sections(chapters, audience)
        if s["exercise"]
    ]
    out = [GENERATED_BANNER]
    out.append(
        "# %s curriculum — hands-on exercises\n" % AUDIENCE_LABELS[audience]
    )
    out.append(GENERATED_NOTE)
    out.append(
        "**%d exercises.** Every step is the handbook's own documented "
        "procedure for that feature, so an exercise cannot drift from the "
        "guide the cohort keeps.\n" % len(pairs)
    )
    out.append(
        "> **Run these on the sandbox tenant only.** Training an agent to "
        "escalate by escalating a real customer's complaint is not a viable "
        "exercise. Reset between cohorts with "
        "`./reset-sandbox-tenant.sh` (see `../delivery-plan.md` §4).\n"
    )
    out.append(
        "> **NOT YET DRY-RUN.** No exercise in this set has been executed "
        "against a sandbox tenant — no sandbox tenant has been provisioned, "
        "and the environment these were generated in has no live Chatwoot, "
        "Gemini or Twilio. \"Completable as written\" is therefore **owed, "
        "not verified**. Dry-run the set once before the first cohort and "
        "record the result in `../delivery-plan.md`.\n"
    )

    procedures = [p for p in pairs if p[1]["steps_source"] == "how-to"]
    roleplays = [p for p in pairs if p[1]["steps_source"] == "body"]
    out.append(
        "**%d are procedures** the trainee carries out themselves, taken from "
        "a handbook section's *How to use it* steps. **%d are role-plays** "
        "drawn from the end-to-end scenarios, where the facilitator plays the "
        "customer — those narratives are told from both sides, so they are "
        "not a checklist one trainee can work through alone.\n"
        % (len(procedures), len(roleplays))
    )

    for chapter, section in pairs:
        if not section["steps"]:
            # A tagging error, not a rendering edge case: `exercise` promises
            # the cohort something to carry out, and this section documents no
            # steps at all. Loud, because the alternative is an exercise sheet
            # with an empty task on it.
            raise TrainingTagError(
                "%s: section %r is marked `exercise` but documents no steps; "
                "drop the flag or add the procedure to the handbook"
                % (chapter["file"], section["title"])
            )
        exercise_id = ids[(chapter["file"], section["title"])]
        procedure = section["steps_source"] == "how-to"
        out.append(
            "## %s — %s%s\n"
            % (
                exercise_id,
                section["title"],
                "" if procedure else "  *(role-play)*",
            )
        )
        out.append(
            "**Module:** %s %s · **Source:** `%s` → `## %s`\n"
            % (
                chapter["number"],
                chapter["title"],
                chapter["file"],
                section["title"],
            )
        )
        if section["where"]:
            out.append("**Where:** %s\n" % section["where"])
        if procedure:
            out.append("**Do this:**\n")
        else:
            out.append(
                "**Reproduce this scenario on the sandbox tenant**, with the "
                "facilitator playing the customer:\n"
            )
        for i, step in enumerate(section["steps"], start=1):
            out.append("%d. %s" % (i, step))
        out.append("")
        if procedure:
            # Deliberately not "step N is complete — <text of step N>": a
            # handbook procedure's last numbered item is sometimes a caveat
            # rather than an action, and quoting it turns the completion
            # criterion into something the trainee cannot do.
            out.append(
                "**Done when:** all %d steps have been carried out on the "
                "sandbox tenant without help.\n" % len(section["steps"])
            )
        else:
            out.append(
                "**Done when:** the cohort has walked all %d steps and can say "
                "which of them the CRM did on its own.\n" % len(section["steps"])
            )
        for note in section["notes"]:
            out.append("**Expect this limitation:** %s\n" % note)

    out.append("## Exercises this set does not contain\n")
    out.append(
        "An exercise exists only where the handbook documents steps a cohort "
        "can actually carry out on a sandbox tenant. The topics below are "
        "presentation-only, and the reason is recorded rather than "
        "papered over with a lab that cannot run.\n"
    )
    out.extend(render_gap_section(audience, heading_level="###"))
    return "\n".join(out).rstrip() + "\n"


def render_checklist(chapters, audience):
    pairs = audience_sections(chapters, audience)
    ids = exercise_ids(chapters, audience)
    out = [GENERATED_BANNER]
    out.append(
        "# %s curriculum — competency checklist\n" % AUDIENCE_LABELS[audience]
    )
    out.append(GENERATED_NOTE)
    out.append(
        "One row per topic in this curriculum, in delivery order. Sign a row "
        "when the trainee has done it unaided — for exercise rows, without "
        "the exercise sheet in front of them.\n"
    )
    out.append(
        "**Trainee:** ________________  **Assessor:** ________________  "
        "**Date:** ____________\n"
    )
    out.append("| # | Competency | Assessed by | Handbook source |")
    out.append("|---|---|---|---|")
    for i, (chapter, section) in enumerate(pairs, start=1):
        exercise_id = ids.get((chapter["file"], section["title"]))
        if section["steps"] and section["steps_source"] == "how-to":
            competency = "Carry out the %d documented steps of **%s** unaided" % (
                len(section["steps"]),
                section["title"],
            )
        elif section["steps"]:
            competency = (
                "Walk **%s** through end to end and say which of its %d steps "
                "the CRM performed on its own"
                % (section["title"], len(section["steps"]))
            )
        elif section["table_rows"]:
            competency = (
                "Use **%s** as a reference (%d rows) and find an entry in it"
                % (section["title"], section["table_rows"])
            )
        else:
            competency = "Explain **%s** in their own words" % section["title"]
        assessed = (
            "Exercise `%s`" % exercise_id if exercise_id else "Q&A / observation"
        )
        out.append(
            "| %d | %s | %s | `%s` § %s |"
            % (i, competency, assessed, chapter["file"], section["title"])
        )
    out.append("")
    out.append(
        "**%d competencies, %d of them assessed by a hands-on exercise.** The "
        "remainder are assessed by question and observation, because the "
        "handbook documents them as behaviour to understand rather than a "
        "procedure to perform.\n" % (len(pairs), len(ids))
    )
    out.extend(render_gap_section(audience))
    return "\n".join(out).rstrip() + "\n"


def render_tag_coverage(chapters):
    """Every section, its audience, and where that audience came from.

    The audit artefact for the filter: it is how a reviewer sees that no
    section fell through to the fallback, and that each curriculum contains
    what it should. A count in a report cannot be checked; this table can."""
    all_sections = [(c, s) for c in chapters for s in c["sections"]]
    fallbacks = [
        (c, s) for c, s in all_sections if s["audience_source"] == "fallback"
    ]
    out = [GENERATED_BANNER]
    out.append("# Training audience tag coverage\n")
    out.append(GENERATED_NOTE)
    out.append(
        "**%d handbook sections across %d chapters.** The audiences are "
        "cumulative (`agent` < `supervisor` < `admin`), so a section tagged "
        "`agent` is taught to all three cohorts.\n"
        % (len(all_sections), len(chapters))
    )
    out.append("| Curriculum | Topics | Exercises | Rule-derived length | Design target |")
    out.append("|---|---|---|---|---|")
    for audience in AUDIENCES:
        pairs = audience_sections(chapters, audience)
        out.append(
            "| %s | %d | %d | %s | %s |"
            % (
                AUDIENCE_LABELS[audience],
                len(pairs),
                len(exercise_ids(chapters, audience)),
                duration_text(sum(slide_minutes(s) for _, s in pairs)),
                duration_text(AUDIENCE_TARGET_MINUTES[audience]),
            )
        )
    out.append("")
    out.append(
        "**Sections with no marker of their own or their chapter's: %d.** "
        "Those fall back to `%s`, the widest curriculum, so untagged content "
        "still reaches a cohort instead of vanishing from all three — and it "
        "is named here so the fallback is visible rather than assumed.\n"
        % (len(fallbacks), FALLBACK_AUDIENCE)
    )
    if fallbacks:
        for chapter, section in fallbacks:
            out.append(
                "- `%s` § %s" % (chapter["file"], section["title"])
            )
        out.append("")
    out.append("| Chapter | Section | Audience | Tagged by | Exercise | Agent | Supervisor | Admin |")
    out.append("|---|---|---|---|---|---|---|---|")
    for chapter, section in all_sections:
        out.append(
            "| `%s` | %s | %s | %s | %s | %s | %s | %s |"
            % (
                chapter["file"],
                section["title"],
                section["audience"],
                section["audience_source"],
                "yes" if section["exercise"] else "—",
                "yes" if sees(section["audience"], "agent") else "—",
                "yes" if sees(section["audience"], "supervisor") else "—",
                "yes" if sees(section["audience"], "admin") else "—",
            )
        )
    out.append("")
    return "\n".join(out).rstrip() + "\n"


def build_curricula():
    """{relative path: content} for every generated training file."""
    chapters = load_curriculum_model()
    files = {os.path.join("tag-coverage.md"): render_tag_coverage(chapters)}
    for audience in AUDIENCES:
        files[os.path.join(audience, "facilitator-deck.md")] = render_deck(
            chapters, audience
        )
        files[os.path.join(audience, "exercises.md")] = render_exercises(
            chapters, audience
        )
        files[os.path.join(audience, "competency-checklist.md")] = (
            render_checklist(chapters, audience)
        )
    return files


def write_curricula(check=False):
    files = build_curricula()
    stale = []
    for relative, content in sorted(files.items()):
        path = os.path.join(TRAINING_DIR, relative)
        if check:
            if not os.path.exists(path):
                stale.append("%s does not exist" % relative)
                continue
            with open(path, "r", encoding="utf-8") as f:
                if f.read() != content:
                    stale.append("%s is stale" % relative)
            continue
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

    if check:
        if stale:
            for line in stale:
                print("training/%s" % line, file=sys.stderr)
            print(
                "Run: python3 docs/client-materials/build_crm_feature_guide.py"
                " --curricula",
                file=sys.stderr,
            )
            return 1
        print("training/ is current (%d generated files)." % len(files))
        return 0

    chapters = load_curriculum_model()
    print("Training curricula build summary")
    print("=" * 40)
    for audience in AUDIENCES:
        pairs = audience_sections(chapters, audience)
        print(
            "%-24s %3d topics  %2d exercises  %s (target %s)"
            % (
                AUDIENCE_LABELS[audience],
                len(pairs),
                len(exercise_ids(chapters, audience)),
                duration_text(sum(slide_minutes(s) for _, s in pairs)),
                duration_text(AUDIENCE_TARGET_MINUTES[audience]),
            )
        )
    fallbacks = [
        s
        for c in chapters
        for s in c["sections"]
        if s["audience_source"] == "fallback"
    ]
    print("Sections with no tag of their own or their chapter's: %d" % len(fallbacks))
    for section in fallbacks:
        print("  - %s § %s" % (section["chapter_file"], section["title"]))
    for relative in sorted(files):
        print("Wrote: training/%s" % relative)
    return 0


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def chapter_source_paths():
    return sorted(
        p
        for p in glob.glob(os.path.join(SRC_DIR, "*.md"))
        if re.match(r"^\d{2}-.*\.md$", os.path.basename(p))
    )


def build_handbook(audience=None):
    os.makedirs(ASSETS_DIR, exist_ok=True)
    gitkeep = os.path.join(ASSETS_DIR, ".gitkeep")
    if not os.path.exists(gitkeep):
        open(gitkeep, "w").close()

    document = Document(TEMPLATE)
    sanitize_page_margins(document)
    clear_body(document)

    use_bullet_style = style_exists(document, "List Bullet") and style_exists(
        document, "List Bullet 2"
    )
    use_number_style = style_exists(document, "List Number")

    # --- Cover page ---
    title_p = document.add_paragraph(COVER_TITLE, style="Title")
    subtitle_p = document.add_paragraph(COVER_SUBTITLE, style="Subtitle")
    document.add_page_break()

    # --- Chapters (resolved first: the TOC is written before them) ---
    # `audience=None` is the default handbook and reads each file verbatim;
    # a role-scoped build drops the sections that role is not taught, and
    # drops a chapter entirely once nothing in it survives the filter.
    chapters = []
    for path in chapter_source_paths():
        text = read_chapter(path, audience)
        if text.strip():
            chapters.append((path, text))
    chapter_paths = [path for path, _ in chapters]
    toc_entries = scan_headings([text for _, text in chapters])
    bookmarks = {}
    for level, text, name in toc_entries:
        bookmarks.setdefault((level, text), []).append(name)

    # --- Table of contents ---
    # Deliberately NOT "Heading 1"/"Heading 2": those are reserved for the
    # 14 chapter titles and their sections, and this label would otherwise
    # appear as an entry in its own list.
    toc_heading = document.add_paragraph(style="normal")
    toc_run = toc_heading.add_run("Table of Contents")
    toc_run.bold = True
    toc_run.font.size = Pt(20)
    for level, text, name in toc_entries:
        add_toc_entry(document, level, text, name)
    document.add_page_break()

    stats = {"sections": 0, "tables": 0, "found": [], "missing": []}

    for i, (path, text) in enumerate(chapters):
        if i > 0:
            document.add_page_break()
        process_chapter(
            document, text, use_bullet_style, use_number_style, stats, bookmarks
        )

    document.save(OUT)

    print("PROTON CRM Feature Guide build summary")
    print("=" * 40)
    if audience is not None:
        print("Audience filter    : %s" % audience)
    print("Chapters processed : %d" % len(chapter_paths))
    for path in chapter_paths:
        print("  - %s" % os.path.basename(path))
    print("Sections (##)      : %d" % stats["sections"])
    print("TOC entries        : %d" % len(toc_entries))
    unlinked = sum(len(v) for v in bookmarks.values())
    print("TOC entries unlinked: %d" % unlinked)
    print("Tables rendered    : %d" % stats["tables"])
    print(
        "Screenshots found  : %d/%d"
        % (len(stats["found"]), len(stats["found"]) + len(stats["missing"]))
    )
    if stats["missing"]:
        print("Screenshots missing (rendered as placeholders):")
        for shot_id in stats["missing"]:
            print("  - %s" % shot_id)
    print("Output: %s" % OUT)
    return 0


def validate_markers():
    """Parse every chapter's markers and let a bad one raise.

    Run in EVERY mode, including the plain handbook build: parsing emits
    nothing, so it cannot change the output, and a misspelled audience name
    should stop the nearest build rather than wait for whoever next runs
    `--curricula`."""
    for path in chapter_source_paths():
        with open(path, "r", encoding="utf-8") as f:
            parse_chapter_sections(f.read(), path)


def main(argv=None):
    parser = argparse.ArgumentParser(description=__doc__.split("\n\n")[0])
    parser.add_argument(
        "--audience",
        choices=AUDIENCES,
        help="render a role-scoped handbook instead of the full one "
        "(cumulative: agent < supervisor < admin). Combine with FG_OUT so it "
        "does not overwrite the shipped guide.",
    )
    parser.add_argument(
        "--curricula",
        action="store_true",
        help="write the three training curricula under training/ and skip the "
        ".docx build",
    )
    parser.add_argument(
        "--check",
        action="store_true",
        help="exit 1 if any committed file under training/ differs from what "
        "would be generated",
    )
    args = parser.parse_args(argv)

    # One try around everything, so a marker problem is reported the same way
    # whichever mode surfaced it — including the plain handbook build.
    try:
        validate_markers()
        if args.check:
            return write_curricula(check=True)
        if args.curricula:
            return write_curricula()
        return build_handbook(audience=args.audience)
    except TrainingTagError as exc:
        print("TRAINING marker error: %s" % exc, file=sys.stderr)
        return 2


if __name__ == "__main__":
    sys.exit(main())
