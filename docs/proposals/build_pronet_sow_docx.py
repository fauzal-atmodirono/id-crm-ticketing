#!/usr/bin/env python3
"""Build the Devoteam G Cloud Technical Proposal / SOW for PRO-NET (RFP 2026_028).

The technical-proposal-generator skill normally copies a master Google Doc and
fills 27 tokens through the Docs API via the `gws` CLI. `gws` is not installed
on this machine, so this builder reconstructs the template locally per
references/template-structure.md section 7 ("Rebuilding this template"):

  - preserved prose and tables come from references/boilerplate.md blocks 1-8,
    with the documented defects fixed rather than reproduced;
  - the eight narrative blocks and the scalar values are filled with
    PRO-NET-specific content;
  - anything not confirmed is written as a literal {{TBD - ...}} so it surfaces
    in the residual scan and in the handover report.

Deliberate deviations from the master template, each for a stated reason:
  - Section 2.3's four generic "Google Data Cloud" intro paragraphs are replaced.
    They describe a BigQuery data-warehouse engagement; this is a CRM/contact-
    centre engagement and reproducing them would misdescribe the solution.
  - Section 3.1's Gantt carries Bahasa Indonesia phase labels. PRO-NET is a
    Malaysian client, so the timeline is rebuilt in English against the RFP's
    own eight phases.
  - Section 6.1's heading/body inconsistency (Enhanced vs Premium Support) is
    resolved rather than carried forward - see known template defect 1.
  - Section 6.2's standard Devoteam SLA table is superseded by the RFP's own
    section 9 severity table, which is stricter. Both are shown, with the
    governing one stated.
  - Block 7g's "eg a CEO of one of the operators" telco leftover is corrected.
  - {{SUPPORT_TIMEZONE}} fills as MYT, not the template's Indonesian default.

Run:  python3 build_pronet_sow_docx.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(
    BASE, "PRO-NET - Devoteam G Cloud Technical Proposal SOW - CCMS.docx"
)

DARK = RGBColor(0x1A, 0x1A, 0x2E)
MID = RGBColor(0x5A, 0x5A, 0x6E)
BLUE = RGBColor(0x1F, 0x4E, 0x8C)
POPPY = RGBColor(0xE8, 0x4A, 0x3D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HDR_FILL = "1F4E8C"
ZEBRA = "F7F9FC"
BAND = "EDF2F9"
TIER_FILL = "E8EFF9"
NOTE_FILL = "FFF6E5"

FONT = "Calibri"

# ---------------------------------------------------------------- scalars
TOK = {
    "PROJECT_TITLE": "Customer Complaint Management System",
    "CLIENT_LEGAL_NAME": "{{TBD - PRO-NET full legal entity name (contract party)}}",
    "CLIENT_SHORT_NAME": "PRO-NET",
    "PROPOSAL_DATE": "{{TBD - proposal date}}",
    "SUPPORT_TIMEZONE": "MYT (Malaysia Time, UTC+8)",
    "SUPPORT_PORTAL_URL": "{{TBD - Devoteam support portal URL}}",
    "SUPPORT_EMAIL": "{{TBD - Devoteam support email}}",
    "SDM_NAME": "{{TBD - Service Delivery Manager name}}",
    "SDM_EMAIL": "{{TBD - SDM email}}",
    "TECH_LEAD_NAME": "{{TBD - Technical Lead name}}",
    "TECH_LEAD_EMAIL": "{{TBD - Technical Lead email}}",
    "TAM_NAME": "{{TBD - Technical Account Manager name}}",
    "TAM_EMAIL": "{{TBD - TAM email}}",
    "ESCALATION_L3_NAME_1": "{{TBD - L3 escalation contact 1}}",
    "ESCALATION_L3_EMAIL_1": "{{TBD - L3 contact 1 email}}",
    "ESCALATION_L3_NAME_2": "{{TBD - L3 escalation contact 2}}",
    "ESCALATION_L3_EMAIL_2": "{{TBD - L3 contact 2 email}}",
}
CS = TOK["CLIENT_SHORT_NAME"]

# The industry credential must be one Devoteam can actually evidence. The
# capability claim below is evidenced by this engagement itself; the sector
# reference claim is left open rather than invented.
INDUSTRY_CREDENTIAL = (
    "We work with customer-operations and contact-centre platforms in the "
    "automotive and mobility sector, and are familiar with Malaysia's Personal "
    "Data Protection Act 2010 and its implications for customer data held in "
    "cloud infrastructure. {{TBD - add one verifiable Devoteam automotive or "
    "contact-centre client reference, or delete this sentence; do not ship an "
    "unevidenced sector claim}}"
)


# ---------------------------------------------------------------- helpers
def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hexfill)
    tcPr.append(el)


def cell_text(cell, text, *, size=9, bold=False, color=DARK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align is not None:
        p.alignment = align
    for i, line in enumerate(str(text).split("\n")):
        tgt = p if i == 0 else cell.add_paragraph()
        if i:
            tgt.paragraph_format.space_before = Pt(0)
            tgt.paragraph_format.space_after = Pt(2)
            if align is not None:
                tgt.alignment = align
        r = tgt.add_run(line)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT
    r.font.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = BLUE


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DARK


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT
    r.font.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = POPPY


def para(doc, text, *, size=10, italic=False, color=DARK, after=7, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = bold
    r.font.color.rgb = color
    return p


def lead(doc, label, text, *, size=10):
    """A bolded lead-in followed by running prose - the house bullet style."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = DARK
    r2 = p.add_run(text)
    r2.font.name = FONT
    r2.font.size = Pt(size)
    r2.font.color.rgb = DARK


def bullet(doc, text, *, size=10, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = DARK


def table(doc, headers, rows, widths, *, size=9, header_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, hh in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, HDR_FILL)
        cell_text(c, hh, size=header_size, bold=True, color=WHITE)
    for n, r in enumerate(rows):
        row = t.add_row()
        for i, v in enumerate(r):
            c = row.cells[i]
            if n % 2:
                shade(c, ZEBRA)
            cell_text(c, v, size=size, bold=(i == 0 and len(r) > 2))
    t.autofit = False
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    return t


def callout(doc, title, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    shade(c, NOTE_FILL)
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run(title)
    r.font.name = FONT
    r.font.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = POPPY
    p2 = c.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.name = FONT
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = DARK
    t.rows[0].cells[0].width = Inches(6.9)
    doc.add_paragraph()


def arch_tier(doc, label, sublabel, items):
    """One horizontal band of the architecture figure."""
    t = doc.add_table(rows=2, cols=max(len(items), 1))
    t.style = "Table Grid"
    head = t.rows[0].cells[0]
    for c in t.rows[0].cells[1:]:
        head = head.merge(c)
    shade(head, HDR_FILL)
    cell_text(head, f"{label}    ·    {sublabel}", size=8.5, bold=True, color=WHITE)
    for i, (name, desc) in enumerate(items):
        c = t.rows[1].cells[i]
        shade(c, TIER_FILL)
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(name)
        r.font.name = FONT
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = BLUE
        p2 = c.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(desc)
        r2.font.name = FONT
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = MID
    t.autofit = False
    w = 6.9 / max(len(items), 1)
    for row in t.rows:
        for c in row.cells:
            c.width = Inches(w)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("▼")
    r.font.size = Pt(9)
    r.font.color.rgb = MID


# =====================================================================
def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10)
    st.font.color.rgb = DARK

    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.75)
    sec.left_margin = sec.right_margin = Inches(0.8)

    # ============================================ front matter
    for _ in range(5):
        doc.add_paragraph()
    for line, size, bold, color in [
        ("TECHNICAL PROPOSAL  ·  STATEMENT OF WORK", 11, True, POPPY),
        ("", 8, False, DARK),
        (f"For {TOK['PROJECT_TITLE']}", 22, True, BLUE),
        ("in Google Cloud Platform", 16, False, MID),
        ("", 10, False, DARK),
        ("Prepared for:", 10, False, MID),
        (TOK["CLIENT_LEGAL_NAME"], 13, True, DARK),
        ("PROTON e.MAS  ·  Customer Operations", 10, False, MID),
        ("", 10, False, DARK),
        ("Prepared by:", 10, False, MID),
        ("Devoteam G Cloud", 13, True, DARK),
        ("", 10, False, DARK),
        (f"Date: {TOK['PROPOSAL_DATE']}", 10, False, MID),
        ("In response to RFP 2026_028  ·  Confidential", 9, False, MID),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color

    doc.add_page_break()

    # ============================================ 1 Company Profile
    h1(doc, "1.  Company Profile")
    para(doc,
         "Devoteam is a leading consulting firm focused on digital strategy, tech "
         "platforms, cybersecurity and business transformation. Devoteam G Cloud is "
         "Devoteam's Google pillar, and has more than 400+ talents globally, with over "
         "600+ certifications. We are a premier managed services partner and authorized "
         "training partner, and we handle all end to end from enablement, solution "
         f"design, implementation, and project operations. {INDUSTRY_CREDENTIAL}")
    para(doc, "Here are some reasons why you should collaborate with Devoteam:")

    for group, items in [
        ("Expertise and Experience:", [
            ("Deep Google Cloud Expertise: ",
             "Devoteam has over 600 Google Cloud experts and 13 years of experience "
             "working with the Google Cloud platform. They have a deep understanding of "
             "the Google Cloud ecosystem and can help you navigate the complexities of "
             "cloud migration, implementation, and management."),
            ("Global Reach: ",
             "Devoteam has a global presence with offices in 18 countries, allowing them "
             "to provide local support and expertise to clients around the world."),
            ("Industry-Specific Expertise: ",
             "Devoteam has experience working with a wide range of industries, including "
             "Telco, Financial Services, and Retail. They understand the unique "
             "challenges and opportunities faced by different industries and can tailor "
             "their solutions accordingly."),
        ]),
        ("Comprehensive Services:", [
            ("End-to-End Cloud Solutions: ",
             "Devoteam offers a comprehensive suite of services, including cloud "
             "strategy, cloud migration, cloud optimization, cloud security, and cloud "
             "managed services. They can help you with every aspect of your cloud "
             "journey, from planning to implementation and ongoing support."),
            ("Managed Services: ",
             "Devoteam offers managed services that can help you free up your internal "
             "resources and focus on your core business. They can manage your Google "
             "Cloud environment, ensuring that it is always secure, reliable, and "
             "performing at its best."),
            ("Innovation and Agility: ",
             "Devoteam is committed to innovation and uses the latest technologies and "
             "methodologies to deliver value to its clients."),
        ]),
        ("Strong Partner Ecosystem:", [
            ("Google Cloud Partner of the Year: ",
             "Devoteam has been recognized as a 5x Google Cloud Partner of the Year, "
             "demonstrating their commitment to customer success and their expertise in "
             "Google Cloud."),
            ("Strong Partnerships: ",
             "Devoteam has a strong network of partners, one of them is Google Cloud. "
             "This allows them to provide a wide range of solutions and services to "
             "their clients."),
        ]),
        ("Value for Money:", [
            ("Competitive Pricing: ",
             "Devoteam offers competitive pricing and flexible payment options to meet "
             "the needs of its clients."),
            ("Value-Driven Approach: ",
             "Devoteam is committed to delivering value to its clients. They focus on "
             "helping clients achieve their business goals and providing a return on "
             "investment."),
        ]),
        ("Cybersecurity Expertise:", [
            ("Cybersecurity Consulting and Audit: ",
             "Devoteam has a strong emphasis on cybersecurity and offers services such "
             "as penetration testing, application development, and security audits."),
        ]),
    ]:
        bullet(doc, group)
        for label, text in items:
            lead(doc, label, text, size=9.5)

    callout(doc, "Before submission - verify",
            "The corporate figures in this section (400+ talents, 600+ certifications, "
            "600 Google Cloud experts, 13 years, 18 countries, 5x Partner of the Year) "
            "are carried from Devoteam collateral and are internally inconsistent on the "
            "talent count. Refresh them against current collateral before this document "
            "goes to PRO-NET, and resolve the {{TBD}} in the industry credential "
            "sentence above.")

    doc.add_page_break()

    # ============================================ 2 Proposed Solution
    h1(doc, "2.  Proposed Solution")

    # ---- PROBLEM_STATEMENT
    para(doc,
         "PROTON e.MAS is scaling an electric-vehicle business in Malaysia, and customer "
         "expectation is scaling with it. Buyers who research, book and service a vehicle "
         "digitally expect the same immediacy when something goes wrong - an answer in "
         "the language they wrote in, on the channel they chose, without repeating "
         "themselves. As the fleet grows, complaint volume grows with it, and the "
         "operation that handles those complaints becomes a visible part of the brand "
         "rather than a back-office function.")
    para(doc,
         f"Today {CS} runs customer complaint handling across five channels - WhatsApp, "
         "email, telephony, Facebook Messenger and Instagram - on a per-seat SaaS CRM. "
         "PRO-NET's own assessment, stated in RFP 2026_028, is that the current system "
         "is limited in integration and dashboards and lacks major functions. Those three "
         "limitations compound each other: because the CRM is not integrated with the "
         "Dealer Management System or the Vehicle Telematics platform, an agent taking a "
         "call cannot see the customer's vehicle, service history or open repair order, "
         "and must ask the customer for information PRO-NET already holds. Because the "
         "dashboards are limited, the pattern behind repeated complaints is not visible "
         "until it is large. And because escalation to dealers is not systematically "
         "timed, the moment a case leaves PRO-NET's own team is the moment it stops being "
         "measurable.")
    para(doc,
         "The operational cost of that is specific and measurable. RFP 2026_028 sets a "
         "2-hour dealer first-response target, an 8-hour escalation to the higher-level "
         "manager and a 48-hour unresolved-case alert; Appendix B sets a 20-second call "
         "answer target, a 2-working-hour agent acknowledgement on transferred social "
         "cases and a 4-working-hour email status update; and the SSI process targets an "
         "SSI score above 90% with a response rate above 45%. Each of those is a number "
         "the current platform cannot reliably enforce or evidence - the escalation "
         "timestamps, business-hours-aware SLA clocks and dealer league tables required "
         "to manage against them do not exist. Where a dealer disputes a turnaround "
         "figure, there is no system-recorded escalation timestamp to settle it. "
         "{{TBD - insert PRO-NET's current monthly case volume, agent headcount, average "
         "first-response time and current SLA attainment; these figures materially "
         "strengthen this section and should be requested at Discovery}}")
    para(doc,
         "The direction PRO-NET has set in the RFP is the right one: consolidate every "
         "channel into one platform, route automatically to the right owner, put a "
         "knowledge base and AI in front of the routine questions, and instrument the "
         "whole lifecycle so the operation can be managed on evidence. Devoteam proposes "
         "to meet that with a self-managed, AI-native complaint management platform "
         "running on Google Cloud inside PRO-NET's own tenancy: an open-source CRM core "
         "extended with a Google Gemini and Vertex AI layer, configured directly against "
         "the Appendix A case taxonomy and the Appendix B channel process flows. It is "
         "not a greenfield build - the platform was demonstrated to PRO-NET's team on "
         "28 July 2026, and this engagement is configuration, gap closure, integration "
         "and migration against a working system.")
    para(doc,
         "The architecture choice is also a commercial one, and it is what makes this "
         "scale. A per-seat SaaS CRM prices growth: every additional agent, every premium "
         "feature and every AI resolution adds recurring cost, which is precisely the "
         "wrong cost curve for a contact centre that must expand as the fleet grows. A "
         "self-managed platform prices infrastructure instead. PRO-NET can add agents at "
         "a seasonal peak with no licensing consequence, retain seven years of operational "
         "data without a per-gigabyte archive charge, add the sixth and seventh channel "
         "without a contract negotiation, and extend the same platform to other business "
         "units - because the conversations, the knowledge base, the AI prompts and the "
         "analytics warehouse all sit in infrastructure PRO-NET owns.")

    # ---- 2.2 Architecture
    h2(doc, "2.2  Proposed Architecture")
    para(doc,
         "The target architecture, tier by tier, from customer contact through to "
         "management reporting:", size=9.5, italic=True, color=MID)

    arch_tier(doc, "CUSTOMER CHANNELS", "inbound and outbound", [
        ("WhatsApp", "text + voice notes, attachments"),
        ("Telephony / IVR", "1300-888-877, RSA routing"),
        ("Email", "e.mascentre@pronet.my"),
        ("Facebook & Instagram", "Meta Business messaging"),
    ])
    arch_tier(doc, "EDGE", "TLS, routing, per-tenant entry", [
        ("Cloud Load Balancing + Caddy", "TLS termination, routing, per-tenant entry point"),
    ])
    arch_tier(doc, "APPLICATION", "GKE Standard - containerised, node-pool control", [
        ("CRM core (GKE)", "agent workspace, case management, live chat"),
        ("agent service (GKE)", "webhook sync, AI orchestration, SLA timers"),
        ("backend service (GKE)", "Gemini agent, knowledge base, routing, metrics"),
    ])
    arch_tier(doc, "AI  ·  DATA  ·  STATE", "managed Google Cloud services", [
        ("Vertex AI - Gemini", "drafts, classify, same-language reply, multimodal"),
        ("Vertex AI Search + pgvector", "KB grounding - traceable, not a black box"),
        ("Cloud SQL (PostgreSQL, HA)", "CRM database, per-tenant knowledge base"),
        ("Memorystore · Cloud Storage", "queues, attachments, call recordings"),
    ])
    arch_tier(doc, "INTEGRATION", "PRO-NET systems of record", [
        ("Dealer Management System", "customer, vehicle, service history, open RO"),
        ("Vehicle Telematics (TSP)", "vehicle telemetry and status"),
    ])
    arch_tier(doc, "CONSUMPTION", "reporting and analytics", [
        ("BigQuery", "analytics warehouse, 7-year history"),
        ("Power BI", "executive and operational reporting"),
        ("In-platform dashboards", "live supervisor and agent views"),
    ])

    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    shade(c, BAND)
    cell_text(c,
              "CROSS-CUTTING  -  Secret Manager  ·  Cloud KMS (customer-managed keys)  ·  "
              "Cloud Monitoring & Logging  ·  automated backups and cross-region DR  ·  "
              "append-only audit trail on write-once storage  ·  RBAC and MFA/SSO  ·  "
              "multi-tenant isolation (per-tenant namespaces, node pools and isolated "
              "databases)",
              size=8.5, bold=True)
    c.width = Inches(6.9)
    doc.add_paragraph()

    para(doc,
         "Region: asia-southeast1 (Singapore), the nearest full-service Google Cloud "
         "region; the Malaysia region is not yet generally available. A PDPA cross-border "
         "transfer assessment is produced at Discovery, and migration to a Malaysian "
         "region is committed when one reaches general availability, should PRO-NET "
         "require it.",
         size=9, italic=True, color=MID)

    # ---- ARCHITECTURE_NARRATIVE
    h3(doc, "How a complaint moves through the architecture")
    lead(doc, "Customer channels. ",
         "A customer reaches PRO-NET on WhatsApp, by calling 1300-888-877, by email to "
         "e.mascentre@pronet.my, or through Facebook Messenger or Instagram. Every "
         "channel terminates on the same platform and resolves to the same customer "
         "record, so a customer who calls after sending a WhatsApp message is not a new "
         "contact. Voice calls enter through the IVR, where a real-time Gemini Live voice "
         "agent handles the conversation, transcribes it in the language spoken, and "
         "routes to a live agent or to the RSA path based on what the caller says rather "
         "than on menu key presses.")
    lead(doc, "Edge. ",
         "Cloud Load Balancing and Caddy terminate TLS and route traffic to the correct "
         "tenant. This is also where per-tenant isolation begins: each tenant has its own "
         "entry point, its own namespace and its own database, so PRO-NET's data is "
         "separated at the infrastructure layer, not by an application-level filter.")
    lead(doc, "Application. ",
         "Three containerised services run on GKE Standard. The CRM core carries the "
         "unified agent workspace, case management and conversation state. The agent "
         "service handles webhook ingestion, AI orchestration, SLA timers and the "
         "escalation ladder - it is what fires the 2-hour, 8-hour and 48-hour triggers "
         "against business hours. The backend service hosts the Gemini agent, the "
         "knowledge base, the routing rules and the metrics pipeline. Running on GKE "
         "Standard rather than a serverless platform is deliberate: it gives node-pool "
         "level control over the resources a real-time voice workload needs, and keeps "
         "per-tenant isolation enforceable.")
    lead(doc, "AI, data and state. ",
         "Vertex AI serves Gemini for reply drafting, classification, sentiment, "
         "summarisation and multimodal understanding of customer photos and video. "
         "Grounding runs against Vertex AI Search and a pgvector store holding PRO-NET's "
         "own FAQ and knowledge base - so an AI answer can be traced to the source entry "
         "that produced it, which is what makes the accuracy calibration cycle possible. "
         "Cloud SQL for PostgreSQL in a high-availability configuration holds the CRM "
         "database; Memorystore carries queues and session state; Cloud Storage holds "
         "attachments, transcripts and call recordings under lifecycle rules aligned to "
         "the 7-year retention obligation.")
    lead(doc, "Integration. ",
         "The DMS and TSP connectors call PRO-NET's systems of record to populate the "
         "Customer 360 card - personal information, vehicle details, service history and "
         "open repair orders - and write case activity back. The card renders immediately "
         "from data the platform already holds and each external section loads "
         "asynchronously behind a short-lived cache and a circuit breaker, so a slow "
         "system of record degrades one panel rather than blocking the agent.")
    lead(doc, "Consumption. ",
         "Case, interaction, SLA, CSAT/NPS and AI-performance data streams into BigQuery, "
         "which is the analytics warehouse and the single source for Power BI. Supervisors "
         "work from live in-platform dashboards; management works from Power BI reports "
         "built to the Appendix C1 and C2 formats; scheduled exports go out as PDF and "
         "Excel on a per-recipient schedule.")
    lead(doc, "Cross-cutting. ",
         "Secrets are held in Secret Manager, data is encrypted at rest with "
         "customer-managed KMS keys and in transit with TLS 1.2+, and every notification, "
         "acknowledgement, escalation and status change is written to an append-only audit "
         "trail on write-once storage with a retention lock. Monitoring and alerting run "
         "continuously, including on the RSA path, which is engineered to remain available "
         "through the 00:00-04:00 maintenance window.")

    # ---- ARCHITECTURE_SUMMARY
    para(doc,
         "Taken together, this architecture gives PRO-NET a complaint operation where "
         "every channel lands in one place, every case carries the customer's full "
         "context, every escalation is timed and evidenced, and every AI decision is "
         "recorded and measurable - all running on infrastructure PRO-NET owns, priced on "
         "consumption rather than on headcount, and extensible to new channels, new "
         "business units and new markets without a licensing negotiation.")

    doc.add_page_break()

    # ---- 2.3 Solution Components
    h2(doc, "2.3  Solution Components")
    para(doc,
         "The components below are the ones in scope for this engagement, described in "
         "the order a complaint flows through them. Google Cloud managed services are "
         "used wherever one exists, so that PRO-NET's operational burden after handover "
         "is configuration rather than infrastructure maintenance, and so that the "
         "platform inherits Google's own availability commitments (section 5.4).")

    for title, paras in [
        ("2.3.1  Omni-Channel Intake", [
            "All five channels required by RFP section 2.1.2 terminate on one platform "
            "against one customer record: WhatsApp text and voice notes through the "
            "WhatsApp Business API, telephony and IVR with call recording, email over "
            "IMAP/SMTP on PRO-NET's own domain, and Facebook Messenger and Instagram "
            "through Meta Business messaging.",
            "Each channel is configured to its own process flow from Appendix B - the AI "
            "disclaimer, the business-hours check, the idle-warning and auto-close "
            "timings, the resolution confirmation and the rating survey - so the SOP "
            "PRO-NET already operates is what the platform enforces, rather than a "
            "vendor's default behaviour that PRO-NET must adapt to.",
        ]),
        ("2.3.2  Unified Agent Workspace", [
            "One interface for every channel, with the conversation, the case, the "
            "Customer 360 card and the AI assist panel on a single screen. Voice is "
            "handled inside the same workspace - answer, transfer, mute and hold - with "
            "the live transcript rendering as the caller speaks, so an agent never "
            "switches applications mid-interaction.",
            "Inbound work raises an in-app and desktop notification, and agent state "
            "(availability, channel priority, current load) drives assignment so work is "
            "never routed to an agent who cannot take it.",
        ]),
        ("2.3.3  Vertex AI and Gemini - the AI layer", [
            "Gemini serves every AI function in the platform: conversational handling on "
            "chat and voice, reply drafting for agents, case classification against the "
            "Appendix A taxonomy, sentiment and urgency detection, conversation "
            "summarisation, and multimodal understanding of customer-shared photos and "
            "video.",
            "Language coverage spans English, Bahasa Malaysia, Chinese and Tamil, with "
            "the model detecting and answering in the customer's own language including "
            "colloquial and abbreviated Malay. Consumption is billed from PRO-NET's own "
            "Google Cloud account at Google's published rates - there is no per-resolution "
            "or per-conversation AI charge from Devoteam, and token and audio usage is "
            "recorded per conversation so cost is attributable and controllable.",
        ]),
        ("2.3.4  Knowledge Base and FAQ - Vertex AI Search with pgvector", [
            "Operator-authored FAQ entries and uploaded source documents are indexed into "
            "a vector store and served through Vertex AI Search, so AI answers are "
            "grounded on PRO-NET's own content rather than on general web knowledge.",
            "The CRM team maintains the knowledge base directly and in real time, with "
            "bulk CSV import for large revisions. Because every AI answer records the "
            "knowledge it retrieved, an incorrect answer is traceable to the entry that "
            "caused it - which is what makes the calibration cycle a measurement exercise "
            "rather than guesswork.",
        ]),
        ("2.3.5  Rule Engine, Escalation and SLA", [
            "The rule engine resolves the Person-In-Charge from case category, division "
            "and dealer, using routing maps that PRO-NET administrators maintain "
            "themselves. Escalation issues the email to the PIC with the configured CC "
            "group and the optional WhatsApp alert.",
            "SLA timers implement the RFP's escalation matrix - 2-hour first response, "
            "8-hour manager escalation, 48-hour unresolved alert - measured against "
            "configured business hours and public holidays. The escalation timestamp is "
            "stamped on the case when a dealer escalation is raised, which is what makes "
            "the dealer first-response and turnaround reporting in requirement 4.59 "
            "defensible when a figure is disputed.",
        ]),
        ("2.3.6  Customer 360 and DMS/TSP Integration", [
            "A configurable integration layer calls PRO-NET's Dealer Management System "
            "and Vehicle Telematics platform to assemble the Customer 360 card: personal "
            "information, vehicle details, service history and open repair orders, "
            "alongside the complete cross-channel interaction history the platform holds "
            "natively.",
            "The card is triggered automatically on inbound contact and matched on caller "
            "number, WhatsApp number, vehicle registration number or chassis number. "
            "Endpoints, credentials and field mappings are administrator-configurable, so "
            "a change to a DMS field does not require a Devoteam release.",
        ]),
        ("2.3.7  Case Management and Taxonomy", [
            "The Appendix A hierarchical taxonomy - Case Category, Case Division and four "
            "further levels across Sales, Product, Network, Charging and Apps - is loaded "
            "at configuration and maintained thereafter by PRO-NET administrators, with "
            "bulk import for large revisions.",
            "The case lifecycle carries New, Assigned, WIP, Pending Customer, Higher "
            "Escalation, Temporarily Closed, Resolved and Closed, with role-enforced "
            "transitions. Historical cases retain the category assigned at the time, so "
            "reporting stays coherent across a taxonomy change.",
        ]),
        ("2.3.8  BigQuery Analytics Warehouse", [
            "Case, interaction, SLA, CSAT/NPS and AI-performance data streams into "
            "BigQuery, which holds the full operational history for the 7-year retention "
            "period and serves as the single governed source for all reporting.",
            "Because it is PRO-NET's own warehouse, the data is available to PRO-NET's "
            "analysts and to any future BI or data-science initiative without an export "
            "request or a vendor API quota.",
        ]),
        ("2.3.9  Power BI Reporting", [
            "A governed Power BI dataset over BigQuery with row-level security aligned to "
            "the platform's roles, carrying the full report set required by RFP "
            "requirements 4.49 to 4.82 and built to the formats already in use in "
            "Appendix C1 (monthly) and C2 (weekly).",
            "Scheduled distribution sends PDF and Excel exports to management on a "
            "per-recipient schedule with per-recipient data scoping. PRO-NET's own BI team "
            "can build further reports on the same dataset without vendor involvement.",
        ]),
        ("2.3.10  Security, Identity and Audit", [
            "Role-based access control with granular function and data permissions, MFA "
            "and federation to PRO-NET's corporate identity provider, encryption at rest "
            "with customer-managed KMS keys and in transit with TLS 1.2+, and PII masking "
            "applied on display, on export and in AI prompt payloads.",
            "An append-only audit trail on write-once storage with a retention lock "
            "records every notification, acknowledgement, escalation and status change "
            "with actor, timestamp, SLA state and remarks, satisfying the tamper-evidence "
            "requirement at RFP 2.2.5 and 3.2.6.",
        ]),
    ]:
        h3(doc, title)
        for p in paras:
            para(doc, p, size=9.5, after=5)

    doc.add_page_break()

    # ---- 2.4 Scope of Works
    h2(doc, "2.4  Scope of Works")
    para(doc,
         "The following Project has been identified to address the following scope:")

    for phase, items in [
        ("Phase 1  -  Mobilisation and Discovery", [
            "Project kick-off with the Steering Committee and Working Group; governance, "
            "communication plan and RAID log established.",
            "Requirements workshops against RFP 2026_028, the Appendix A case taxonomy "
            "and the Appendix B channel process flows, producing the Business "
            "Requirements Document.",
            "Develop-versus-integrate boundary confirmed per component and recorded.",
            "DMS and TSP data contract, field mapping and response-time expectations "
            "agreed with PRO-NET; sandbox access requested.",
            "PII classification and masking rules agreed; PDPA cross-border transfer "
            "assessment for the Singapore region.",
            "Solution architecture, GCP landing-zone design and the Technical Design "
            "Document.",
            "Legacy data profiling and the migration plan; Test Plan drafted.",
        ]),
        ("Phase 2  -  GCP Foundation and Provisioning", [
            "GCP project, VPC, IAM and landing zone provisioned from infrastructure-as-code.",
            "GKE Standard cluster and node pools; Cloud SQL in high-availability "
            "configuration; Memorystore; Cloud Storage; BigQuery; Artifact Registry.",
            "CI/CD pipeline, Secret Manager, Cloud KMS keys and the monitoring and "
            "alerting baseline.",
            "Three environments stood up - Development, UAT and Production - with "
            "environment smoke tests and access handover.",
        ]),
        ("Phase 3  -  Core Platform Configuration and Channel Integration", [
            "Platform deployed across all three environments.",
            "All five channels wired and configured to their Appendix B process flows: "
            "WhatsApp (text and voice), telephony and IVR including RSA routing and call "
            "recording, email with the one-acknowledgement-per-thread rule, Facebook "
            "Messenger and Instagram.",
            "Appendix A case taxonomy loaded; business hours and the Malaysian public "
            "holiday calendar configured.",
            "SLA policies, the PIC and dealer routing maps, recipient groups and the "
            "escalation ladder configured, including the 2-hour, 8-hour and 48-hour "
            "timers.",
            "Roles, permissions and PII masking rules configured; MFA and SSO federation "
            "to PRO-NET's identity provider.",
        ]),
        ("Phase 4  -  Gap Closure and Build", [
            "DMS and TSP connectors, the Customer 360 card population and the automatic "
            "screen-pop on inbound contact.",
            "Agent management build: the eight named agent statuses, availability "
            "monitoring with the 10-minute and 1-hour threshold alerts, follow-up "
            "reminders, timeout warnings and audible alerting.",
            "AI-powered duplicate case detection and merging on mobile number, email, "
            "vehicle registration number and chassis number, with PIC notification on "
            "merge.",
            "Agent-facing translation surface across English, Bahasa Malaysia, Chinese "
            "and Tamil.",
            "One-click FAQ insertion into the composer; attachment carriage on escalation "
            "emails.",
            "Email delivery-status tracking with failure alerting, and the send-delay and "
            "retraction workflow offered in place of email recall.",
            "AI activity monitor console for administrators.",
        ]),
        ("Phase 5  -  AI Enablement, Analytics and Reporting", [
            "AI persona, guardrails and lifecycle messages configured; PRO-NET FAQ and "
            "knowledge-base content loaded and indexed.",
            "Multimodal image and video understanding, sentiment analysis and "
            "conversation summarisation enabled and tuned.",
            "First AI calibration cycles executed with written scorecards against agreed "
            "baselines for transcription accuracy, language match, sentiment "
            "classification, FAQ match precision and summary accuracy.",
            "BigQuery warehouse, data pipelines and semantic layer.",
            "Power BI workspace, dataset and the full report set at requirements 4.49 to "
            "4.82, built to the Appendix C1 and C2 formats; scheduled distribution and "
            "the anomaly warning dashboard.",
        ]),
        ("Phase 6  -  Testing, Training and UAT", [
            "SIT executed against the PRO-NET-agreed test script: functional coverage, "
            "all five channel flows, escalation matrix and timers, integration contracts, "
            "RBAC, performance against the 3-second Customer 360 budget, and failure "
            "paths. SIT report issued and exit signed.",
            "Role-based on-site training for frontline agents, supervisors and system "
            "administrators, with manuals, hands-on exercises, train-the-trainer and "
            "session recordings.",
            "UAT with Devoteam environment support, daily defect triage and same-cycle "
            "fixes; UAT report and acceptance sign-off.",
        ]),
        ("Phase 7  -  Data Migration", [
            "Legacy assessment and profiling; field mapping and category cross-walk to "
            "the Appendix A taxonomy; cleansing and transformation rules agreed.",
            "Full dress-rehearsal migration into UAT with record-count, checksum and "
            "attachment reconciliation, reviewed and signed.",
            "Final delta synchronisation of open legacy complaints executed at cutover, "
            "with a cutover reconciliation report as a signed gate.",
        ]),
        ("Phase 8  -  Cutover, Go-Live and Post-Implementation Support", [
            "Production readiness review; disaster-recovery restore test executed and "
            "evidenced before go-live.",
            "Cutover executed against the Production Deployment Checklist with a defined "
            "rollback decision point and named owner.",
            "30-day hypercare with a dedicated engineering team, on-site for the first "
            "week and for any Critical (P1) incident, daily stand-ups and a live defect "
            "board.",
            "Legacy platform decommission and licence termination.",
            "180-day warranty support, including AI prompt and knowledge-base "
            "optimisation and monthly calibration cycles.",
            "Documentation handover, knowledge transfer and project close-out.",
        ]),
    ]:
        h3(doc, phase)
        for it in items:
            bullet(doc, it, size=9.5)

    doc.add_page_break()

    # ============================================ 3 Timeline & Deliverables
    h1(doc, "3.  Project Timeline and Deliverables")

    h2(doc, "3.1  Timeline")
    para(doc,
         "Ten weeks from kick-off to go-live, followed by 30 days of hypercare and a "
         "180-day warranty period. Phases overlap where they safely can - migration runs "
         "alongside testing, and the AI and analytics workstream alongside the core build "
         "- which is what makes the schedule achievable without compressing UAT.",
         size=9.5)
    table(doc,
          ["Phase", "W1", "W2", "W3", "W4", "W5", "W6", "W7", "W8", "W9", "W10",
           "Hypercare", "Warranty"],
          [
              ("Mobilisation & Discovery", "█", "█", "", "", "", "", "", "", "", "", "", ""),
              ("GCP Foundation", "", "█", "█", "", "", "", "", "", "", "", "", ""),
              ("Core Platform & Channels", "", "", "█", "█", "█", "█", "", "", "", "", "", ""),
              ("Gap Closure & Build", "", "", "", "█", "█", "█", "█", "", "", "", "", ""),
              ("AI, Analytics & Power BI", "", "", "", "█", "█", "█", "█", "", "", "", "", ""),
              ("Testing, Training & UAT", "", "", "", "", "", "", "█", "█", "█", "", "", ""),
              ("Data Migration", "", "", "", "", "", "█", "█", "█", "█", "", "", ""),
              ("Cutover & Go-Live", "", "", "", "", "", "", "", "", "", "█", "", ""),
              ("Hypercare (30 days)", "", "", "", "", "", "", "", "", "", "", "█", ""),
              ("Warranty (180 days)", "", "", "", "", "", "", "", "", "", "", "", "█"),
          ],
          [1.75] + [0.32] * 10 + [0.62, 0.55], size=8)

    h2(doc, "3.2  Deliverables")
    para(doc,
         "The following deliverables are produced by this engagement. Each is mapped to "
         "the phase that produces it and, where applicable, to the RFP section 10 payment "
         "milestone it evidences. Every deliverable is issued in draft for PRO-NET review "
         "before the formal sign-off request.")
    table(doc,
          ["Deliverable", "Phase", "Milestone evidenced"],
          [
              ("Project Kick Off Deck and project plan", "1", "Mobilization (15%)"),
              ("Business Requirements Document (BRD)", "1", "Requirements Sign-off (15%)"),
              ("Technical Design Document (TSD)", "1", "Requirements Sign-off (15%)"),
              ("Solution Architecture document", "1", "Requirements Sign-off (15%)"),
              ("Migration Plan and legacy data profile", "1, 7", "-"),
              ("PDPA cross-border transfer assessment", "1", "-"),
              ("API Integration Document (DMS, TSP, BI)", "4", "-"),
              ("Approved Test Plan and test scripts", "1, 6", "-"),
              ("AI Calibration Scorecards", "5", "-"),
              ("Power BI report pack (requirements 4.49-4.82)", "5", "-"),
              ("SIT / QA Report with signed exit", "6", "Gates entry to UAT"),
              ("Training Manuals (agent, supervisor, administrator)", "6", "-"),
              ("UAT Report and acceptance sign-off", "6", "UAT Sign-off (40%)"),
              ("Data Reconciliation Report (rehearsal and cutover)", "7", "-"),
              ("Production Deployment Checklist", "8", "Go-Live (20%)"),
              ("System Configuration Document", "8", "Go-Live (20%)"),
              ("System Operation Manual and runbooks", "8", "Go-Live (20%)"),
              ("Fault Handling Document", "8", "-"),
              ("DR runbook and evidenced restore test", "8", "-"),
              ("Hypercare Exit Report", "8", "-"),
              ("Project Completion Certificate", "8", "Project Close Out (10%)"),
              ("Snag / Punch List and Project Close Document", "8", "Project Close Out (10%)"),
          ],
          [3.6, 0.8, 2.5], size=9)
    para(doc,
         "All documents are delivered in source-editable form so PRO-NET can maintain "
         "them after handover.", size=9.5, italic=True, color=MID)

    doc.add_page_break()

    # ============================================ 4 Out of Scope
    h1(doc, "4.  Out of Scope")
    para(doc, "Following activities and tasks are considered out of scope:")
    for it in [
        "Development, enhancement or maintenance of the Dealer Management System, the "
        "Vehicle Telematics platform, or the e.MAS mobile application themselves. This "
        "engagement integrates with them; they remain PRO-NET-owned systems.",
        "Provision of DMS, TSP or e.MAS API specifications, credentials or sandbox "
        "environments - a PRO-NET dependency.",
        "Microsoft Power BI licences (Pro or Premium) for report authors and consumers, "
        "under PRO-NET's existing Microsoft agreement.",
        "Telephony carrier contracts, number porting fees, call charges and per-message "
        "channel charges (WhatsApp Business API, Meta messaging). These are pass-through "
        "third-party costs.",
        "Google Cloud infrastructure and Vertex AI consumption charges, which are billed "
        "directly to PRO-NET's own Google Cloud account.",
        "Meta Business verification of PRO-NET's Facebook and Instagram business "
        "identity, which only PRO-NET can complete; Devoteam supports the process but "
        "the verification turnaround is controlled by Meta.",
        "Migration of data not present in the agreed legacy export, and remediation of "
        "legacy data quality issues beyond the agreed cleansing rules.",
        "Agent workstations, headsets, local network, internet connectivity and desktop "
        "software.",
        "Languages beyond the four specified in RFP requirement 4.3 (English, Bahasa "
        "Malaysia, Chinese, Tamil).",
        "Changes to PRO-NET's corporate email infrastructure or identity provider beyond "
        "establishing the connection to the platform.",
        "Contact centre staffing, workforce management and day-to-day operation of the "
        "complaint handling process.",
        "Regulatory certification, external audit or legal opinion. Devoteam implements "
        "the controls and provides evidence; certification remains PRO-NET's.",
        "Ongoing Application Managed Services after the 180-day warranty period, which "
        "are described in section 6 and priced separately.",
        "Any requirement marked Non-compliant in the RFP compliance response, "
        "specifically requirement 4.40 (email recall to external recipients), for which "
        "an alternative is delivered.",
    ]:
        bullet(doc, it, size=9.5)

    doc.add_page_break()

    # ============================================ 5 Project Implementation
    h1(doc, "5.  Project Implementation")

    h2(doc, "5.1  Agile Methodology")
    para(doc,
         "The team executes in an agile way, using weekly sprints with demos at the end "
         "of each sprint, but still having milestones.")
    para(doc,
         "Starting with a thorough assessment and preparation stage, the team then dives "
         "into brainstorming, generating ideas and defining the core functionalities of "
         "the product. The subsequent phases focus on conceptual design and technical "
         "architecture, ensuring the product's viability and adherence to user needs.")
    para(doc,
         "Finally, the product is built and deployed, marking the successful completion "
         "of the Agile sprint. This iterative approach, using sprints of one week, allows "
         "for continuous improvement and adaptability to changing requirements, while "
         "still maintaining clear milestones for progress tracking. This methodology "
         "promotes a collaborative and flexible development environment, ensuring that "
         "the final product meets the ever-evolving demands of the users.")
    para(doc,
         "For this engagement the hybrid is explicit: the contractual spine - Discovery "
         "and requirements sign-off, the architecture blueprint, the SIT and UAT gates "
         "and the section 10 payment milestones - is sequential and formally signed, "
         "while build and configuration run in weekly sprints with a working "
         "demonstration on the UAT tenant at the end of each. This matters for a "
         "complaint system: the escalation matrix and the case taxonomy always change "
         "once operators see them working, and an iterative build absorbs that without a "
         "change request.",
         size=9.5)

    h2(doc, "5.2  Change Management and Project Communication")
    para(doc,
         "All change requests that arise within this project will be assessed by all "
         f"parties, Change Request Documents will be created by Devoteam and approved and "
         f"informed by {TOK['CLIENT_LEGAL_NAME']} before entered in the development and "
         "implementation state by Devoteam. Document Approval will be Project Manager in "
         f"Devoteam side and Authorized Person by {CS}.")
    para(doc,
         "Proper communication channels are important to discuss, decide, notify and "
         "escalate issues. As we have agile methodology, we conduct sprint planning at "
         "the beginning of each sprint. If required, we conduct daily stand-up.")

    h2(doc, "5.3  Channels")
    lead(doc, "Email: ", "Primary and formal channel for notification, request and escalation.")
    lead(doc, "Sprint planning: ",
         "backlogs grooming, assignment and update/review at the beginning of each sprint.")
    lead(doc, "Sprint review: ", "MVP demo and update/review at the end of each sprint.")
    lead(doc, "Final presentation: ", "Final of the project.")

    h2(doc, "5.4  Google Cloud Platform Service Level Agreement")
    para(doc,
         "Google Cloud Platform (GCP) has a strong commitment to providing reliable and "
         "high-performance services to its customers. To ensure that our products meet "
         "the demanding needs of businesses, we have implemented rigorous Service Level "
         "Agreements (SLAs) that guarantee specific levels of availability, performance, "
         "and operational support. These SLAs provide customers with confidence in the "
         "reliability and stability of GCP's infrastructure, allowing them to focus on "
         "their core business objectives without worrying about disruptions or downtime.")
    table(doc,
          ["Product", "Monthly Uptime Percentage", "Reference"],
          [
              ("Google Kubernetes Engine (regional cluster control plane)", ">= 99.95%",
               "https://cloud.google.com/kubernetes-engine/sla"),
              ("Cloud SQL (Enterprise edition with HA)", ">= 99.95%",
               "https://cloud.google.com/sql/sla"),
              ("Cloud SQL (Enterprise Plus edition with HA)", ">= 99.99%",
               "https://cloud.google.com/sql/sla"),
              ("BigQuery", ">= 99.99%", "https://cloud.google.com/bigquery/sla"),
              ("Cloud Storage (Standard)", ">= 99.9%",
               "https://cloud.google.com/storage/sla"),
              ("Vertex AI", "{{TBD - verify current figure}}",
               "https://cloud.google.com/vertex-ai/sla"),
              ("Memorystore for Redis (Standard tier)", "{{TBD - verify current figure}}",
               "https://cloud.google.com/memorystore/docs/redis/sla"),
          ],
          [2.6, 1.5, 2.8], size=8.5)
    para(doc,
         "The GKE and Cloud SQL figures above were verified against Google's published "
         "SLA pages on 5 August 2026. The BigQuery and Cloud Storage figures are carried "
         "from Devoteam's template of record. The two marked for verification must be "
         "read from their SLA pages and entered before this document is issued - SLA "
         "terms are "
         "versioned and do change, and a figure typed from memory is a contractual "
         "exposure.",
         size=8.5, italic=True, color=POPPY)

    doc.add_page_break()

    # ============================================ 6 Post Implementation
    h1(doc, "6.  Post Implementation (Optional)")
    para(doc,
         "The 30-day hypercare period and the 180-day post-production warranty are "
         "included in the engagement priced in section 2.4 and are not optional. This "
         "section describes the ongoing Application Managed Service that applies after "
         "the warranty expires, which is quoted separately.",
         size=9.5, italic=True, color=MID)

    h2(doc, "6.1  Google Cloud Platform Premium Support")
    callout(doc, "Tier confirmation required",
            "Devoteam's template carries an inconsistency here, resolved rather than "
            "reproduced: the heading previously said \"Enhanced Support\" while the body "
            "described Premium Support. These are different Google Customer Care tiers "
            "with different response commitments and different prices. The section below "
            "describes Premium Support, whose 1-hour P1 response is what PRO-NET's own "
            "section 9 severity model depends on. Confirm the tier PRO-NET is purchasing "
            "before issue.")
    para(doc,
         "Google Premium Support offers unlimited technical support for outages and "
         "defects, unexpected product behavior, product usage questions and billing "
         "issues.")
    para(doc, "When managing support cases as a Premium Support customer, you have access "
              "to the following features:")
    lead(doc, "P1 response SLO: ",
         "For Priority 1 (P1) support cases, receive the first meaningful response within "
         "1 hour.")
    lead(doc, "24/7 availability: ",
         "Receive support 24 hours a day, 7 days a week (24/7) for cases of certain "
         "priority and language.")
    lead(doc, "Language support: ",
         "Request support across multiple languages, including English, Japanese, "
         "Mandarin Chinese, Korean, and French.")
    lead(doc, "Case escalation: ",
         "Escalate to request additional attention for ongoing support cases.")
    table(doc, ["Priority", "Target Initial Response Times"],
          [("P1", "1 hour"), ("P2", "4 hours"), ("P3", "8 hours*"), ("P4", "8 hours*")],
          [1.6, 5.3], size=9)
    para(doc, "* during the Hours of Operation", size=8.5, italic=True, color=MID)

    h2(doc, "6.2  Devoteam Support and Managed Services")
    para(doc,
         "Devoteam Support offers unlimited technical support for outages and defects, "
         "unexpected product behavior, product usage questions and billing issues.")

    h3(doc, "Incident Management")
    table(doc, ["Incident Management", "Devoteam", "Customer"],
          [
              ("Perform 24x7 incident management support", "R, A", "C, I"),
              ("Generate incident tickets based on events", "R, A, C, I", "R, C, I"),
              ("Evaluate and categorize incidents for prioritization", "R, A, C", "I"),
              ("Respond to and remediate incidents within agreed SLO / KPI", "R, A, C", "I"),
              ("Manage incident escalation to Google via Partner-led Premium Support",
               "R, A, C", "I"),
              ("Support and troubleshoot incidents for tier 1 end-users", "I", "R, A, C"),
              ("Manage and/or resolve incidents outside the scope of the managed assets "
               "on cloud platform", "I", "R, A, C"),
          ],
          [4.5, 1.2, 1.2], size=9)
    para(doc,
         "R = Responsible, A = Accountable, C = Consulted, I = Informed. Tier-1 end-user "
         f"support and anything outside the managed assets remain {CS}'s responsibility.",
         size=8.5, italic=True, color=MID)

    h3(doc, "Service Level Agreement - governing terms")
    para(doc,
         "PRO-NET's own severity model and response and resolution targets, set out in "
         "RFP 2026_028 section 9.0, are accepted in full and govern this engagement. They "
         "are stricter than Devoteam's standard support table and supersede it. They "
         "apply during hypercare, throughout the 180-day warranty, and under the managed "
         "service thereafter.",
         size=9.5)
    table(doc,
          ["Severity", "Target Response Time", "Target Resolution Time", "Coverage"],
          [
              ("Critical (P1)", "< 30 minutes", "< 2 hours",
               "24/7 on-call; on-site attendance for every P1"),
              ("Major (P2)", "< 1 hour", "< 4 hours", "24/7 on-call"),
              ("Moderate (P3)", "< 4 hours", "< 24 hours", "Staffed business hours"),
              ("Minor (P4)", "< 24 hours", "< 5 days", "Staffed business hours"),
          ],
          [1.4, 1.6, 1.6, 2.3], size=9)
    callout(doc, "Commercial note - stricter than Devoteam standard",
            "Devoteam's standard managed-service table commits to a 3-hour P1 resolution "
            "on an 8x5 basis. RFP section 9.0 requires a 2-hour P1 resolution with 24/7 "
            "availability. This proposal commits to the stricter figure, which carries a "
            "staffing cost that must be reflected in the managed-service price. This is "
            "stated here so the commitment is deliberate rather than inherited from "
            "boilerplate.")
    para(doc,
         "Staffed support runs 08:00 to 17:30 " + TOK["SUPPORT_TIMEZONE"] + ", Monday to "
         "Friday, aligned to PRO-NET business hours per RFP 8.1.3, with a 24/7 on-call "
         "rotation for P1 and P2 and 24/7 automated monitoring of the platform and the "
         "RSA path. Weekend and national holiday coverage for P3 and P4 is excluded "
         "unless full 24/7 staffed coverage is contracted as a priced option.",
         size=9.5)

    h3(doc, "Devoteam Support Contact Channel")
    para(doc, "Devoteam provide customer with support contact channel such as,")
    lead(doc, "Support Portal: ",
         "New cases can be opened via the link: " + TOK["SUPPORT_PORTAL_URL"])
    lead(doc, "Email: ",
         "You can contact our services by sending an email to: " + TOK["SUPPORT_EMAIL"])
    lead(doc, "Phone and WhatsApp: ",
         "{{TBD - Devoteam support hotline and WhatsApp number, required by RFP 8.1.2}}")

    h3(doc, "Service Level Infrastructure Uptime")
    para(doc,
         "Devoteam will pass through any applicable service level agreement pertaining to "
         "any Google product or services purchased by the Customer. Devoteam does not "
         "provide any uptime or availability guarantees besides those defined by Google.")

    h3(doc, "Escalation Flow")
    table(doc, ["Flow", "PIC Name"],
          [
              ("Level 1", f"Support Email ( {TOK['SUPPORT_EMAIL']} )"),
              ("Level 2", f"{TOK['SDM_NAME']} ( {TOK['SDM_EMAIL']} )"),
              ("Level 3",
               f"{TOK['ESCALATION_L3_NAME_1']} ( {TOK['ESCALATION_L3_EMAIL_1']} )\n"
               f"{TOK['ESCALATION_L3_NAME_2']} ( {TOK['ESCALATION_L3_EMAIL_2']} )"),
          ],
          [1.4, 5.5], size=9)
    para(doc,
         "Notes : The name of personnel in the table can be adjusted based on Devoteam "
         "policy, and Devoteam will be informed to the customer at least 10 working days "
         "prior.",
         size=8.5, italic=True, color=MID)

    h3(doc, "Hierarchical Escalation")
    para(doc,
         "The hierarchical escalation of the incident should be performed whenever any of "
         "the following situations occur:")
    for it in [
        "the agreed time for the recovery of the service should be approached without any "
        "prospect of such occurrence;",
        "there are obstacles to the recovery of the service within the agreed time, "
        "created by the end customer, other support teams or any other unforeseen factor;",
        "there is a widespread and high impact on the service or business of the end "
        "customer;",
        "a VIP user is affected (eg a member of the Customer's executive leadership);",
        "there is a hierarchical escalation of the incident in the structure of the final "
        "customer;",
    ]:
        bullet(doc, it, size=9.5)
    para(doc,
         "Hierarchical escalation according to priority and SLA evolution is defined in "
         "the matrix below. Escalation is done manually via phone or by email.")
    table(doc,
          ["Priority", "SLA Consumption 50%", "SLA Consumption 75%", "SLA Consumption 100%"],
          [
              ("P1 Critical", "Technical Lead",
               "Technical Lead\nService Delivery Manager",
               "Technical Lead\nService Delivery Manager\nTechnical Account Manager"),
              ("P2 High", "Technical Lead",
               "Technical Lead\nService Delivery Manager",
               "Technical Lead\nService Delivery Manager\nTechnical Account Manager"),
              ("P3 Medium", "",
               "Technical Lead\nService Delivery Manager",
               "Technical Lead\nService Delivery Manager\nTechnical Account Manager"),
              ("P4 Low", "", "Technical Lead",
               "Technical Lead\nService Delivery Manager"),
          ],
          [1.2, 1.7, 1.9, 2.1], size=8.5)

    h3(doc, "Roles")
    table(doc, ["Roles Name", "PIC Name"],
          [
              ("Technical Lead", f"{TOK['TECH_LEAD_NAME']} ( {TOK['TECH_LEAD_EMAIL']} )"),
              ("Service Delivery Manager", f"{TOK['SDM_NAME']} ( {TOK['SDM_EMAIL']} )"),
              ("Technical Account Manager", f"{TOK['TAM_NAME']} ( {TOK['TAM_EMAIL']} )"),
          ],
          [2.3, 4.6], size=9)
    para(doc,
         "Notes : The name of personnel in the table can be adjusted based on Devoteam "
         "policy, and Devoteam will be informed to the customer at least 10 working days "
         "prior.",
         size=8.5, italic=True, color=MID)

    # ---- 6.3
    h2(doc, "6.3  Scope of Works for Maintenance and Managed Service")
    para(doc,
         "The scope of the managed service includes both proactive preventive actions and "
         f"reactive corrective measures, ensuring {CS} receives ongoing support for the "
         "Customer Complaint Management System, minimises downtime across all five "
         "customer channels, and sustains the accuracy and effectiveness of the AI layer "
         "and knowledge base as complaint patterns, dealer networks and vehicle models "
         "change. It covers the platform, the channel integrations, the DMS and TSP "
         "connectors, the analytics warehouse and the reporting layer.")

    bullet(doc, "Preventive Maintenance:")
    for label, text in [
        ("System Updates and Security Patches: ",
         "Regularly applying software updates and security patches to ensure system "
         "stability, performance, and compliance."),
        ("System Optimization: ",
         "Optimizing system configuration, tuning databases, and allocating resources "
         "effectively for optimal performance."),
        ("Data Integrity Checks: ",
         "Regularly checking the integrity of the data to ensure accuracy and prevent "
         "data corruption."),
        ("Security Audits: ",
         "Performing routine security audits to ensure compliance with security standards "
         "and identify potential vulnerabilities."),
        ("AI Calibration and Knowledge Base Health: ",
         "Monthly calibration cycles scoring transcription accuracy, language match, "
         "sentiment classification, FAQ match precision and summary accuracy against "
         "agreed baselines, with a written scorecard; knowledge-base gap analysis driven "
         "by unresolved enquiries; review of grounding and safety-filter configuration."),
        ("Disaster Recovery Testing: ",
         "Restore tests executed and evidenced every six months, with the result "
         "reported."),
    ]:
        lead(doc, label, text, size=9.5)

    bullet(doc, "Corrective Maintenance:")
    for label, text in [
        ("Prompt Response to Issues: ",
         "Quickly addressing system issues, bugs, or hardware failures to minimize "
         "downtime and impact on operations."),
        ("Data Recovery: ",
         "Implementing recovery procedures from backups to restore data integrity and "
         "system functionality in case of failures."),
        ("Root Cause Analysis: ",
         "Investigating issues to identify the underlying causes and implement corrective "
         "measures to prevent recurrence."),
    ]:
        lead(doc, label, text, size=9.5)

    bullet(doc, "Managed Services:")
    for label, text in [
        ("On-Demand Support: ",
         f"Providing support and assistance to {CS}'s IT team on-demand, including "
         "troubleshooting issues, addressing questions, and resolving problems."),
        ("8x5 Onsite Support: ",
         "Offering on-site support services, including troubleshooting complex issues and "
         "providing expert assistance during system installations or upgrades, and "
         "on-site attendance for every Critical (P1) incident as required by RFP 8.1.1."),
        ("24x7 On-call Support: ",
         f"Providing round-the-clock on-call support for P1 and P2 incidents, enabling "
         f"{CS} to receive help even during off-hours for urgent or critical situations."),
        ("System Monitoring: ",
         "Continuously monitoring the system for performance, stability, and security to "
         "detect and address issues proactively, including the RSA path."),
        ("Operator Configuration Support: ",
         "Assistance with bulk updates to the case taxonomy, dealer and PIC routing maps "
         "and recipient groups, and periodic review of routing accuracy."),
        ("Monthly Account Review: ",
         "A monthly service review covering platform performance and availability, the AI "
         "calibration scorecard, knowledge-base health, support SLA attainment, the "
         "system health and capacity report, and written improvement recommendations with "
         "a tracked action list."),
    ]:
        lead(doc, label, text, size=9.5)

    callout(doc, "On-site commitment - confirm before issue",
            "\"8x5 Onsite Support\" and on-site attendance for every P1 commit Devoteam "
            "personnel to PRO-NET's premises in Malaysia. Confirm this is in the "
            "commercial scope and that travel and accommodation costs are reflected in "
            "the managed-service price before this section is issued.")

    doc.save(OUT)
    print(f"wrote: {OUT}")


if __name__ == "__main__":
    build()
