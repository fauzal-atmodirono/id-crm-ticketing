#!/usr/bin/env python3
"""Build the Devoteam vendor response to PRO-NET RFP 2026_028 (CCMS).

The RFP is its own response template: every module row carries an
"[Insert Compliance Status]" and an "[Insert System Solution Details]" cell,
and those cells are what the evaluator grades. This builder reproduces the
RFP's own section and table structure and fills every one of them.

Compliance vocabulary is the RFP's own: Fully / Custom / TP / Non.
  Fully  - Fully out-of-the-box on the platform as it stands today
  Custom - Customization or new build required within this engagement
  TP     - Delivered via a third-party component
  Non    - Not complied with as literally specified (an alternative is stated)

Statuses are asserted against the platform source in this repository, not
against earlier proposal decks. Anything not verifiable is written as an
explicit {{TBD - ...}} so it surfaces in the residual scan rather than
shipping as a confident guess.

Run:  python3 build_pronet_rfp2026_028_response_docx.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, "PRO-NET RFP 2026_028 - Devoteam Vendor Response (Technical).docx")

# ---------------------------------------------------------------- palette
DARK = RGBColor(0x1A, 0x1A, 0x2E)
MID = RGBColor(0x5A, 0x5A, 0x6E)
POPPY = RGBColor(0xE8, 0x4A, 0x3D)
BLUE = RGBColor(0x1F, 0x4E, 0x8C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HDR_FILL = "1F4E8C"
BAND_FILL = "EDF2F9"
ZEBRA_FILL = "F7F9FC"
SUB_FILL = "DCE6F4"

FONT = "Calibri"

# Compliance status -> cell shading, so an evaluator can scan the column.
STATUS_FILL = {
    "Fully": "DFF3E4",
    "Custom": "FFF3D6",
    "TP": "E5E9F5",
    "Non": "FBE0DE",
}


# ---------------------------------------------------------------- helpers
def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hexfill)
    tcPr.append(el)


def cell_text(cell, text, *, size=8.5, bold=False, color=DARK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align is not None:
        p.alignment = align
    for i, line in enumerate(str(text).split("\n")):
        target = p if i == 0 else cell.add_paragraph()
        if i:
            target.paragraph_format.space_before = Pt(0)
            target.paragraph_format.space_after = Pt(2)
            if align is not None:
                target.alignment = align
        run = target.add_run(line)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color


def set_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 11)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT
    run.font.bold = True
    run.font.size = Pt(14 if level == 1 else 11.5)
    run.font.color.rgb = BLUE if level == 1 else DARK
    return p


def body(doc, text, *, size=9.5, italic=False, color=DARK, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def bullets(doc, items, *, size=9.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(it)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.color.rgb = DARK


def matrix(doc, rows, headers, widths, *, status_col=None):
    """rows: list of tuples matching headers. A tuple of length 1 is a band row."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, HDR_FILL)
        cell_text(c, h, size=8.5, bold=True, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER if i else None)
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    z = 0
    for r in rows:
        row = t.add_row()
        if len(r) == 1:  # band / sub-section header spanning the table
            merged = row.cells[0]
            for c in row.cells[1:]:
                merged = merged.merge(c)
            shade(merged, SUB_FILL)
            cell_text(merged, r[0], size=9, bold=True, color=BLUE)
            z = 0
            continue
        fill = ZEBRA_FILL if z % 2 else None
        z += 1
        for i, val in enumerate(r):
            c = row.cells[i]
            if i == status_col:
                shade(c, STATUS_FILL.get(str(val).strip(), fill or "FFFFFF"))
                cell_text(c, val, size=8.5, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                if fill:
                    shade(c, fill)
                cell_text(c, val, size=8.5, bold=(i == 0))
    set_widths(t, widths)
    return t


# =====================================================================
#  CONTENT
# =====================================================================
# Short handles reused across many cells, so the platform is described
# consistently rather than re-invented row by row.
PLAT = "the CCMS platform"

# ---------------------------------------------------------------- 2.0
SEC2 = [
    ("2.1  Supply & Licensing",),
    ("2.1.1 Software Provisioning",
     "Deliver all production and non-production licenses, cloud infrastructure "
     "environments, or on-premises installation packages, including liaison with "
     "third party suppliers.",
     "Fully",
     "Delivered as a self-managed platform on Google Cloud. There is no per-seat "
     "software licence to procure: the CRM core is open-source (Chatwoot Community, "
     "MIT) and the AI layer is Google Gemini / Vertex AI consumed on PRO-NET's own "
     "Google Cloud billing account. PRO-NET therefore pays for infrastructure and AI "
     "consumption, not for agent seats.\n"
     "Three environments are provisioned from infrastructure-as-code: Development, "
     "UAT/sandbox and Production, each an isolated tenant.\n"
     "Devoteam leads all third-party liaison: Google Cloud (infrastructure, Gemini, "
     "Vertex AI), Twilio (WhatsApp Business API, programmable voice / IVR, call "
     "recording), Meta (Facebook Messenger and Instagram Business messaging) and "
     "Microsoft Power BI. Third-party consumption charges are passed through at cost."),
    ("2.1.2 Omnichannel Intake Setup",
     "Ingest complaints from WhatsApp (text and voice chat, attachments), email, "
     "telephony (IVR), FB Messenger and Instagram. No limit to attachment size and "
     "file format.",
     "Custom",
     "All five channels land in one agent workspace against a single customer record.\n"
     "WhatsApp text and voice notes via the WhatsApp Business API (Twilio); email via "
     "IMAP/SMTP on PRO-NET's own domain (e.mascentre@pronet.my); telephony/IVR via "
     "programmable voice with a Gemini Live real-time voice agent; Facebook Messenger "
     "and Instagram via Meta Business messaging.\n"
     "Attachments are stored in Google Cloud Storage with no platform-imposed size or "
     "format limit, and object lifecycle rules aligned to the 7-year retention "
     "requirement.\n"
     "Constraint declared honestly: the messaging carriers impose their own media "
     "ceilings (WhatsApp Business API currently caps media at 16 MB per message; Meta "
     "applies comparable limits). These are carrier limits outside any vendor's "
     "control. Where a customer must send a larger file, the platform issues a secure "
     "signed upload link that accepts any size and format and attaches the result to "
     "the case."),
    ("2.1.3 Workflow Engine",
     "Configurable workflow engine automating escalation matrices, risk scoring and "
     "target resolution timelines.",
     "Fully",
     "A configurable rule engine drives routing and escalation without code changes. "
     "Operators maintain, from admin screens: the case category tree; the "
     "category/department to Person-In-Charge routing map; the dealer routing map; SLA "
     "policies with target first-response and resolution times; and the escalation "
     "ladder.\n"
     "Escalation timers implement the RFP matrix directly - 2-hour first-response "
     "trigger, 8-hour escalation to the higher-level manager, 48-hour unresolved-case "
     "alert - measured against configured business hours.\n"
     "Risk scoring is applied by the AI layer on every inbound interaction "
     "(sentiment/urgency classification) and is available as a routing and "
     "prioritisation input."),
    ("2.2  Configuration & Testing",),
    ("2.2.1 Categorization & Rules Setup",
     "Configure multi-tier case categorization per Appendix A. The list will be "
     "expanded from time to time and the system shall allow addition or changes.",
     "Fully",
     "The platform ships a hierarchical case taxonomy supporting the Appendix A shape "
     "- Case Category / Case Division / Level 1 / Level 2 / Level 3 / Level 4 - across "
     "the Sales, Product, Network, Charging and Apps divisions.\n"
     "The full Appendix A tree is loaded during configuration. Categories are "
     "maintained by PRO-NET administrators in the admin UI, with bulk CSV import for "
     "large revisions. Adding, renaming or retiring a category is an operator action "
     "requiring no vendor involvement, no release and no downtime.\n"
     "Historical cases retain the category assigned at the time, so reporting across a "
     "taxonomy change stays coherent."),
    ("2.2.2 Integration Testing",
     "Test real-time APIs connecting the CCMS to customer databases and enterprise "
     "email networks, including DMS, Vehicle Telematics (TSP) and Power BI or other "
     "analytic tools.",
     "Custom",
     "Integration testing covers four surfaces: the Dealer Management System, the "
     "Vehicle Telematics System, PRO-NET's enterprise email, and Power BI.\n"
     "The platform already carries a configurable DMS/TSP client with an admin screen "
     "for endpoint, credential and field-mapping configuration. What remains is "
     "PRO-NET-specific: the API specification, non-production credentials and a data "
     "contract for each system.\n"
     "Test approach: contract tests against a PRO-NET-supplied sandbox, then end-to-end "
     "tests proving the Customer 360 card renders inside the 3-second budget under "
     "expected concurrency, with evidence captured in the SIT report.\n"
     "Dependency on PRO-NET: DMS and TSP API documentation and sandbox credentials by "
     "the end of the Discovery phase. This is the single largest schedule risk in the "
     "engagement and is carried on the RAID log from kick-off."),
    ("2.2.3 Executive Dashboard Setup",
     "Real-time statistics on operational data and key reports.",
     "Fully",
     "Two layers are delivered. In-platform dashboards give supervisors live operational "
     "state - queue depth by channel, agent availability, cases breaching or near SLA, "
     "open cases by division, dealer and PIC. A BigQuery analytics warehouse behind it "
     "carries the full case, interaction, SLA, CSAT/NPS and AI-performance history and "
     "feeds Power BI.\n"
     "The executive view is assembled against Appendix C1 (monthly) and Appendix C2 "
     "(weekly) so the reports PRO-NET already circulates are reproduced rather than "
     "re-invented."),
    ("2.2.4 AI Speech-to-Text, Automated Responses and Resolution & FAQ Match "
     "Calibration",
     "Calibration of speech-to-text transcription and AI responses on language "
     "adaptation, sentiment/emotion factor, KB completeness and accuracy, and accuracy "
     "of conversation summary.",
     "Fully",
     "Calibration is run as a measured cycle, not a one-off tuning session. A "
     "calibration set is drawn from real PRO-NET interactions across all four languages "
     "and scored on: transcription word-error rate; language detection and response-"
     "language match; sentiment/emotion classification against human labels; FAQ match "
     "precision and recall; and conversation-summary factual accuracy judged against "
     "the transcript.\n"
     "Each cycle produces a scorecard, a prompt/knowledge-base change set and a "
     "re-score. Cycles run weekly through UAT and hypercare, then monthly.\n"
     "Baseline targets are agreed at Discovery and fixed in the Test Plan; the platform "
     "records every AI decision, so accuracy is measured from production evidence "
     "rather than asserted."),
    ("2.2.5 Regulatory Compliance Verification",
     "Validate data masking (PII protection), secure audit logs (tamper-proof "
     "histories), and automated regulatory report formats.",
     "Custom",
     "PII protection: personal data is classified at Discovery. Masking is applied on "
     "display by role, on export, and in AI prompt payloads, so an agent, a report "
     "consumer and the AI layer each see only the fields their role permits. Data at "
     "rest is encrypted with Cloud KMS; data in transit uses TLS 1.2+.\n"
     "Audit log: an append-only audit trail records every notification, "
     "acknowledgement, escalation and status change with actor, timestamp, SLA state "
     "and remarks. Records cannot be edited or deleted through the application; log "
     "sinks are written to write-once storage with a retention lock, so tamper-evidence "
     "is enforced by the platform rather than by policy.\n"
     "Regulatory report formats: built as configured report templates once PRO-NET "
     "confirms the reporting bodies and required formats.\n"
     "{{TBD - which regulatory reporting formats and submitting bodies apply to the "
     "CCMS; PRO-NET Compliance to confirm. Raised as a clarification question.}}"),
    ("2.2.6 System Integration Testing (SIT) / Quality Assurance (QA)",
     "Provide the SIT/QA test report on a pre-agreed test script prior to handover to "
     "users for UAT.",
     "Fully",
     "A Test Plan and test scripts are written from the signed-off BRD and Technical "
     "Design Document and agreed with PRO-NET before execution. SIT covers functional "
     "cases, all five channel flows against Appendix B, the escalation matrix and its "
     "timers, integration contracts, role-based access, performance against the "
     "3-second Customer 360 budget, and negative/failure paths.\n"
     "Output is a SIT/QA report giving per-script pass/fail, defect log by severity, "
     "retest evidence and an explicit exit statement. UAT does not begin until PRO-NET "
     "signs the SIT exit."),
    ("2.2.7 User Acceptance Testing (UAT)",
     "Provide environment support while customer service teams execute real-world "
     "testing scenarios in a sandbox environment.",
     "Fully",
     "A dedicated UAT tenant is provisioned with production-shaped configuration, "
     "masked representative data and sandbox channel connections, so agents test "
     "against real flows without touching live customers.\n"
     "Devoteam provides on-site and remote support for the whole UAT window: a UAT "
     "briefing for the customer service teams, daily defect triage with agreed severity "
     "and turnaround, fixes deployed to UAT within the cycle, and a UAT report with the "
     "acceptance sign-off sheet."),
    ("2.3  Commissioning & Handover",),
    ("2.3.1 Setup & Core Deployment Services / Production Deployment",
     "Execute final code, rule and workflow configuration deployments into the live "
     "production environment.",
     "Fully",
     "Production deployment runs from the same infrastructure-as-code and CI/CD "
     "pipeline used for Development and UAT, so what is tested is what is released.\n"
     "The production environment runs on Google Kubernetes Engine with a highly "
     "available Cloud SQL database, managed cache, private networking and secrets held "
     "in Secret Manager.\n"
     "Go-live is executed against a Production Deployment Checklist covering "
     "pre-cutover verification, channel cutover order, rule and workflow configuration "
     "load, smoke tests, and a documented rollback with a decision point and named "
     "owner."),
    ("2.3.2 Historical Data Migration",
     "Historical data migration including assessment and planning, data preparation "
     "(cleansing, mapping, transformation) and the final data sync and cutover of open "
     "legacy complaints, with zero data loss.",
     "Custom",
     "Migration is run in four stages. Assessment profiles the legacy dataset - volumes, "
     "field coverage, attachment inventory, category values in use - and produces the "
     "migration plan. Preparation builds and reviews the field mapping, cleansing rules "
     "and category cross-walk to the Appendix A taxonomy. A full dress-rehearsal "
     "migration is run into UAT and reconciled. Cutover performs the final delta sync "
     "of open cases.\n"
     "Zero data loss is proved, not asserted: record counts, per-entity checksums and "
     "attachment-count reconciliation are produced for both the rehearsal and the "
     "cutover, and the reconciliation report is a signed cutover gate. The legacy "
     "export is retained unaltered as the fallback.\n"
     "Dependency on PRO-NET: a complete legacy export (cases, contacts, interaction "
     "history, attachments) and confirmation of the retention scope to migrate."),
    ("2.3.3 System Product & Role-Based Training",
     "On-site specialized training sessions and manuals for system administrators, "
     "supervisors and frontline customer service agents.",
     "Fully",
     "Three on-site role-based tracks, each with its own manual and hands-on exercises "
     "in the UAT tenant:\n"
     "Frontline agents - the unified agent workspace, all five channels, the Customer "
     "360 card, case creation and categorisation, AI-suggested replies and the FAQ, "
     "escalation, and the after-call workflow.\n"
     "Supervisors and team leaders - queue and agent monitoring, reassignment, SLA and "
     "escalation oversight, the reporting suite, and CSAT/NPS review.\n"
     "System administrators - category tree, PIC and dealer routing maps, SLA policies, "
     "business hours, users, roles and permissions, knowledge-base and FAQ maintenance, "
     "AI persona and prompt configuration, and health monitoring.\n"
     "A train-the-trainer session and session recordings are included so PRO-NET can "
     "onboard new joiners without vendor involvement."),
    ("2.3.4 Documentation Handover",
     "Deliver finalized architecture maps, functional configuration documents, API "
     "schema charts, user operation and training manuals.",
     "Fully",
     "Handover pack: Solution Architecture document with the deployment and network "
     "topology; Technical Design Document; System Configuration Document recording "
     "every configured rule, category, SLA policy and routing map as built; API "
     "Integration Document with the schema and contract for each DMS, TSP and BI "
     "interface; System Operation Manual (runbooks for deployment, backup and restore, "
     "monitoring, incident handling and user administration); role-based User Guides "
     "and training manuals; and a Fault Handling Document.\n"
     "All documents are delivered in source-editable form so PRO-NET can maintain them "
     "after handover."),
    ("2.4  Support & Maintenance",),
    ("2.4.1 Hypercare Support",
     "Dedicated engineering response teams for 30 days immediately following live "
     "deployment.",
     "Fully",
     "A dedicated hypercare team is committed for the 30 days following go-live: the "
     "engineers who built the platform, not a separate support desk. On-site presence "
     "for the first week of hypercare, remote thereafter with on-site attendance for "
     "any Critical (P1) incident, as required by section 8.1.1.\n"
     "Hypercare runs a daily stand-up with the CRM and call centre teams, a live defect "
     "board, and same-cycle fixes. Section 9 severity SLAs apply throughout. Hypercare "
     "closes on an exit review and report, and only with PRO-NET's agreement."),
    ("2.4.2 Post-Production Warranty Support",
     "Warranty support for 180 days following live deployment to resolve system bugs "
     "and system refinement requiring reconfiguration, such as agent prompt and KB "
     "optimization.",
     "Fully",
     "180 days of warranty from go-live, running from the end of hypercare. Covers "
     "defect resolution against the signed BRD/FSD at no additional charge, and the "
     "refinement work the RFP names explicitly - AI agent prompt tuning, knowledge-base "
     "and FAQ optimisation, category and routing-map adjustments, and report tuning.\n"
     "AI calibration cycles continue monthly through the warranty period with a written "
     "scorecard each cycle. Warranty work is tracked on the same support portal as "
     "incidents, so PRO-NET has one view of everything open."),
    ("2.4.3 SLA Support",
     "Tiered technical troubleshooting support with strict response times for critical "
     "issues.",
     "Fully",
     "The four-tier severity model and the response and resolution targets in section "
     "9.0 are accepted in full - P1 <30 min / <2 h, P2 <1 h / <4 h, P3 <4 h / <24 h, "
     "P4 <24 h / <5 days - and are met during hypercare and the 180-day warranty as "
     "part of this engagement.\n"
     "Beyond the warranty period these targets are sustained under an Application "
     "Managed Service, which is presented in section 8 and priced separately from this "
     "technical proposal."),
    ("2.4.4 System Updates",
     "Regular, free platform updates to maintain software security, performance "
     "optimization, and compatibility with updated web browsers.",
     "Fully",
     "Security patches and platform updates are included at no licence cost - the stack "
     "is open-source and self-managed, so there is no upgrade fee and no forced-upgrade "
     "cycle imposed by a SaaS vendor.\n"
     "Critical security patches are applied within the P1/P2 windows. Routine updates "
     "are batched into a scheduled release, tested in UAT first, and deployed only in "
     "the 00:00-04:00 maintenance window agreed in section 8.1.6, with the RSA path "
     "unaffected.\n"
     "The agent workspace is a standard web application supported on current versions "
     "of Chrome, Edge, Firefox and Safari; browser compatibility is regression-tested "
     "each release."),
    ("2.4.5 System Backup and Disaster Recovery Plan",
     "Provide a mechanism to secure critical data and system restoration to safeguard "
     "business continuity.",
     "Fully",
     "Backup: automated daily full database backups with continuous transaction-log "
     "archiving giving point-in-time recovery, plus versioned object storage for "
     "attachments and call recordings. Backups are held in a separate region and "
     "encrypted with customer-managed keys.\n"
     "Disaster recovery: infrastructure is rebuilt from infrastructure-as-code into the "
     "recovery region and the database restored from the cross-region backup. Proposed "
     "targets are RPO 15 minutes and RTO 4 hours, to be confirmed against PRO-NET's "
     "business-continuity policy at Discovery.\n"
     "A documented DR runbook is delivered, and a restore test is executed and "
     "evidenced before go-live rather than left as a paper commitment. Restore tests "
     "are then repeated every six months under the managed service."),
]

# ---------------------------------------------------------------- 3.0
SEC3 = [
    ("3.1  Integration of Inbound & Agent Management",),
    ("3.1.1 Omni-Channel Single View",
     "Unifies Call, Email, WhatsApp and Social Media into a single agent desktop "
     "interface with pop-up system alert notifications. Agent interactions with callers "
     "should all be carried out within the same single agent desktop interface.",
     "Fully",
     "One agent workspace carries every channel. Call, email, WhatsApp, Facebook "
     "Messenger, Instagram and the web widget all resolve to the same customer record, "
     "so an agent sees one continuous history regardless of how the customer arrived.\n"
     "Voice is handled inside the same interface - the agent answers, transfers, mutes "
     "and holds from the workspace, with the live transcript and the Customer 360 card "
     "on screen. There is no second application to switch to.\n"
     "New inbound work raises an in-app alert with the customer and case context."),
    ("3.1.2 Voice Recognition",
     "Automatic Speech-to-Text conversion for inbound voice calls in real time.",
     "Fully",
     "Inbound calls are transcribed in real time by Gemini Live and streamed to the "
     "agent screen as the conversation happens. Transcription is in the language "
     "actually spoken - it is not translated into English first - and the transcript is "
     "attached to the case as a permanent record alongside the call recording."),
    ("3.1.3 After-Hours Case Monitoring",
     "Support business hour configuration and automatically identify whether incoming "
     "cases are received during or outside operating hours for SLA tracking and "
     "reporting.",
     "Fully",
     "Business hours are configured per channel, including the Mon-Fri 08:30-17:30 and "
     "Sat-Sun/public-holiday 09:00-17:00 windows in Appendix B, with a Malaysian public "
     "holiday calendar.\n"
     "Every case is stamped in-hours or after-hours at creation. SLA clocks then run "
     "against business hours rather than wall-clock time, so a case arriving at 18:00 "
     "is not penalised for the overnight period, and the after-hours volume and "
     "response-time reports in requirements 4.52-4.54 are driven from that same stamp."),
    ("3.1.4 After-Hours Auto Response Management",
     "Automated out-of-office messages for customer enquiries received outside working "
     "hours across email, calls, WhatsApp and social media.",
     "Fully",
     "Per-channel after-hours auto-responses are configured from the exact message text "
     "in Appendix B, in the customer's language.\n"
     "The email channel implements the Appendix B acknowledgement rule precisely: one "
     "auto-acknowledgement per new email or new ticket, never repeated within the same "
     "thread, and never fired in response to an agent reply.\n"
     "Voice calls outside operating hours are met by the configured IVR after-hours "
     "prompt, with the RSA path remaining available 24/7."),
    ("3.1.5 AI Chatbot 24/7 with Escalation",
     "AI chatbot capability to handle customer enquiries 24/7 and escalate complex "
     "cases to human agents when required.",
     "Fully",
     "The AI agent answers around the clock on WhatsApp, social and the web widget, and "
     "on voice through the IVR, grounded on PRO-NET's own FAQ and knowledge base rather "
     "than on general web content.\n"
     "The AI disclaimer in Appendix B is issued at the start of an AI-handled "
     "conversation. Handoff to a human is triggered by explicit customer request, by "
     "low answer confidence, by negative or escalating sentiment, or by category rules - "
     "for example anything routed to RSA. Handoff carries the full transcript and the "
     "AI's summary, so the agent does not restart the conversation.\n"
     "Outside business hours an unresolved conversation creates a case queued for the "
     "next business hour, exactly as Appendix B specifies."),
    ("3.1.6 Duty & Priority Polling",
     "Validates agent on-duty status and respects agent channel priority settings (e.g. "
     "WA vs Call priority), automatically polling tasks based on status (Available, "
     "Busy, etc). Refer to Appendix B.",
     "Custom",
     "Shipped today: an on-duty check before any assignment, per-agent channel priority "
     "so an agent can be a WhatsApp-first or a Call-first resource, automatic assignment "
     "driven by live agent status, and overflow to an available agent outside the "
     "priority channel when every priority-channel agent is busy.\n"
     "To be built in this engagement: the extended status set named in requirement 4.17 "
     "- Lunch, Break, Coaching, Training, Toilet and Prayer as distinct, separately "
     "reportable states rather than a single generic unavailable state. The routing "
     "engine already excludes unavailable agents; this work adds the labelled states, "
     "their configuration screen and their reporting."),
    ("3.1.7 Alert System",
     "Multi-channel task reminders including desktop notifications, in-system alerts and "
     "audible sound warnings.",
     "Custom",
     "In-system alerts and browser desktop notifications are shipped. Audible alerts and "
     "the per-agent control over which events sound - new inbound, reminder due, timeout "
     "warning, SLA breach approaching - are built in this engagement, together with the "
     "escalating reminder and timeout-warning behaviour required by 4.18.\n"
     "Desktop notifications require the agent to grant the browser notification "
     "permission once; this is covered in agent training and in the User Guide."),
    ("3.2  AI FAQ Template & Automated Escalation Engine",),
    ("3.2.1 Real-Time AI FAQ Matching",
     "Contextually detects keywords in live chats to prompt relevant FAQ responses, "
     "allowing single-click usage and quality scoring.",
     "Custom",
     "Shipped today: the AI reads the live conversation and surfaces matching FAQ "
     "entries to the agent as the customer types - semantic matching over the knowledge "
     "base, not literal keyword lookup, so colloquial and abbreviated phrasing still "
     "matches. Agents rate each suggestion, and that feedback drives the FAQ quality "
     "score and the knowledge-base improvement report.\n"
     "To be built in this engagement: true single-click insertion into the message "
     "composer. Today the suggestion is copied to the clipboard for the agent to paste, "
     "because the assist panel runs in a sandboxed frame. Direct insertion is a bounded "
     "change to the composer in our maintained CRM fork and is included in scope."),
    ("3.2.2 Rules-Based SOP Routing - PIC identification and notification",
     "Automatically identifies Person-In-Charge (PIC), sends CC emails with rich "
     "attachments (photos, videos, web links), and triggers WhatsApp alerts.",
     "Custom",
     "Shipped today: the rule engine resolves the PIC from case category, department and "
     "dealer, sends the escalation email with the configured CC group, and triggers the "
     "optional WhatsApp alert to the PIC.\n"
     "To be built in this engagement: carrying the case's own attachments - customer "
     "photos, videos and links - onto the escalation email itself. Today the escalation "
     "email carries a deep link to the case where the media is held. Attaching the media "
     "directly is a small, well-understood change and is included in scope."),
    ("3.2.3 Rules-Based SOP Routing - team email groups",
     "Supports email grouping by predefined teams. For example, selecting \"Aftersales "
     "Team\" will automatically include all relevant parties, eliminating manual "
     "recipient selection.",
     "Fully",
     "Named recipient groups are maintained by PRO-NET administrators in the escalation "
     "routing screen. Selecting a group - Aftersales Team, Charging Team, a specific "
     "dealer - expands to its full membership automatically, and the same groups are "
     "reused by the automatic rule engine so manual and automatic escalation always "
     "reach the same people.\n"
     "Membership is maintained without vendor involvement, satisfying requirement 8.1.9."),
    ("3.2.4 Automated SLA Escalation for Emails to Dealers",
     "2-hour SLA trigger from escalation with overdue flag and reminder if no "
     "acknowledgement; 8-hour automatic escalation to higher-level manager; 48-hour "
     "unresolved-case alert. Refer to Appendix B.",
     "Fully",
     "All three timers are implemented as specified and measured against business hours.\n"
     "The 2-hour clock starts at the moment the case is escalated to the recipient "
     "group. Without an acknowledgement or first response the case is flagged overdue "
     "and a reminder goes to the recipient and their manager. At 8 hours without further "
     "action or update the case escalates automatically to the designated higher-level "
     "manager. At 48 hours without a resolution or an agreed action plan the "
     "unresolved-case alert fires.\n"
     "The escalation timestamp is stamped on the case when the dealer escalation is "
     "raised, which is what makes the dealer first-response and turnaround reporting in "
     "requirement 4.59 measurable rather than estimated. Thresholds, recipients and the "
     "manager ladder are all operator-configurable."),
    ("3.2.5 AI Activity Monitor",
     "Master view/UI of all AI-handled conversations for administrators to monitor in "
     "real time.",
     "Custom",
     "Every AI decision is already recorded - the inbound message, the retrieved "
     "knowledge, the action chosen, the reply issued, the confidence and any handoff "
     "reason - so the underlying data exists today and drives the AI performance reports "
     "in requirement 4.56.\n"
     "This engagement builds the live administrator console over it: all AI-handled "
     "conversations in one real-time view, filterable by channel, language, sentiment "
     "and confidence, with drill-through to the full decision trail and the ability to "
     "take over a conversation."),
    ("3.2.6 System Audit Trail",
     "Complete audit trail of all notifications, acknowledgements, escalations and "
     "updates, including agent or recipient name, date and time, SLA status and remarks.",
     "Fully",
     "An append-only audit trail records every notification sent, acknowledgement "
     "received, escalation fired and status change, each with the actor, the timestamp, "
     "the SLA state at that moment and the remarks entered.\n"
     "Records cannot be edited or deleted through the application. Audit data is "
     "queryable and exportable by administrators, retained for the full 7-year period, "
     "and written to write-once storage with a retention lock so the history is "
     "tamper-evident."),
    ("3.3  Customer 360 View & DMS/TSP Integration",),
    ("3.3.1 Bi-Directional API Integration",
     "Real-time asynchronous data sync (< 3 seconds) with Dealer Management System "
     "(DMS) and Vehicle Telematics System (TSP).",
     "Custom",
     "The platform carries a configurable DMS/TSP client with an administrator screen "
     "for endpoints, credentials, authentication and field mapping. This engagement "
     "builds the PRO-NET-specific connectors on top of it.\n"
     "Inbound: customer, vehicle, service-history and telematics lookups triggered by "
     "the caller's number, WhatsApp number, vehicle registration number or chassis "
     "number. Outbound: case creation, status and resolution written back so the DMS "
     "reflects complaint activity.\n"
     "The 3-second budget is met by asynchronous loading - the card renders immediately "
     "from data the platform already holds and each external section fills in as it "
     "returns, with a short-lived cache and a circuit breaker so a slow or unavailable "
     "DMS degrades one panel rather than blocking the agent. Response times are measured "
     "and reported.\n"
     "Dependency on PRO-NET: API specifications, sandbox and production credentials, and "
     "an agreed field-level data contract for both systems."),
    ("3.3.2 Automated Pop-Up Card",
     "Instantly displays Personal Info, Vehicle Details, Service and Complete "
     "Interaction History across all channels upon incoming contact, including Caller "
     "Name, Mobile Number, Insured Name/Vehicle Owner's Name and Vehicle Registration No.",
     "Custom",
     "The Customer 360 card is shipped, with the personal-information and complete "
     "interaction-history sections live across every channel - call, WhatsApp, email, "
     "social and RSA - and lookup by phone or vehicle number available today.\n"
     "This engagement adds the automatic screen-pop on inbound contact and populates the "
     "vehicle and service sections from the DMS/TSP connectors, including the four named "
     "fields: Caller Name, Mobile Number, Insured Name / Vehicle Owner's Name and "
     "Vehicle Registration Number.\n"
     "The card structure follows requirement 4.46 - personal information, vehicle "
     "information, service information and call centre history as four distinct sections."),
]

# ---------------------------------------------------------------- 4.0
# (id, module, requirement, status, details)
SEC4 = [
    ("Integration of Inbound",),
    ("4.1", "Integration of Inbound",
     "All channels via Call, Email, WA or Social Media must be automatically integrated "
     "on the same agent interface & displayed in a single CRM view",
     "Fully",
     "All channels resolve to one customer record and one agent workspace. Cross-channel "
     "history is continuous: a WhatsApp thread, a call and an email from the same "
     "customer appear on the same timeline."),
    ("4.2", "Integration of Inbound",
     "All new inbound will have a pop up notification (system alert)",
     "Fully",
     "New inbound raises an in-app pop-up carrying the customer, channel and case "
     "context, and a browser desktop notification when the agent has granted "
     "permission. Audible alerting is added under 3.1.7."),
    ("4.3", "Integration of Inbound",
     "AI-powered translation for customer enquiries and agent responses across English, "
     "Bahasa Malaysia, Chinese and Tamil",
     "Custom",
     "Shipped today: the AI detects the customer's language and answers in that same "
     "language across all four languages, on both chat and voice.\n"
     "To be built in this engagement: the agent-facing translation surface - inbound "
     "customer messages shown to the agent in the agent's working language with the "
     "original preserved, and outbound agent replies translated into the customer's "
     "language before sending, with the agent able to see and approve both. Tamil "
     "coverage is validated explicitly during calibration, as it is the least-represented "
     "of the four in general-purpose models."),
    ("4.4", "Integration of Inbound",
     "AI supports natural language understanding (NLU) and recognises variations in "
     "customer input, including spelling variations, abbreviations, and colloquial terms. "
     "Example: brp lama siap? nk service",
     "Fully",
     "The AI layer is built on Gemini, which handles colloquial Malay, SMS-style "
     "abbreviation, code-switching between Malay and English and misspellings natively - "
     "the example given resolves correctly to a service turnaround-time enquiry. "
     "Retrieval against the FAQ is semantic, so a match does not depend on the customer "
     "using the FAQ's own wording. Malaysian colloquial phrasing is included in the "
     "calibration set."),
    ("4.5.1", "Integration of Inbound",
     "When received Call, system will have a voice recognition & change to text. "
     "Transcription shall be in the actual conversed language, in real time on the agent "
     "screen.",
     "Fully",
     "Real-time streaming transcription in the language actually spoken, displayed live "
     "on the agent screen and retained on the case with the recording."),
    ("4.5.2", "Integration of Inbound",
     "When received Call, system will have a mechanism for routing/forwarding the call to "
     "either Live agent or RSA for roadside assistance related requests.",
     "Fully",
     "The IVR determines intent from what the caller says rather than from menu key "
     "presses, and routes to the live agent queue or to the RSA path accordingly. RSA "
     "calls follow a dedicated priority route that stays available 24/7 and is exempt "
     "from the scheduled maintenance window per 8.1.6. RSA incidents are logged against "
     "the customer record and reportable."),
    ("4.6", "Integration of Inbound",
     "Agent is able to transfer a voice call to another agent, have a mute button, hold "
     "line function (with prerecorded media)",
     "Fully",
     "Transfer, mute and hold are available in the agent workspace. Transfer carries the "
     "case, the transcript so far and the Customer 360 context to the receiving agent. "
     "Hold plays PRO-NET-supplied pre-recorded media."),
    ("4.7", "Call Recording Retrieval",
     "The system should allow users to retrieve, and playback call recordings for "
     "customer interactions.",
     "Fully",
     "Every call is recorded and attached to its case. Authorised users play back in the "
     "browser alongside the synchronised transcript, and search by customer, agent, "
     "date, category or vehicle number. Access is permission-controlled and every "
     "playback and download is written to the audit trail. Recordings are retained for "
     "the 7-year period on encrypted storage.\n"
     "Call-recording notification and consent wording for the IVR is confirmed with "
     "PRO-NET at Discovery to meet PDPA obligations."),
    ("4.8", "Inbound",
     "At the end of a WA conversation, customer can rate the AI and/or Live agent (CSAT "
     "Survey) performance",
     "Fully",
     "The WhatsApp CSAT survey fires at conversation close, following the Appendix B "
     "flow - the resolution confirmation (\"Is your case resolved? YES/NO\") followed by "
     "the rating prompt. Ratings are attributed separately to the AI and to the live "
     "agent, so AI response satisfaction (4.56.3) and agent NPS (4.71-4.72) are measured "
     "independently."),
    ("4.9", "Inbound: NPS rating system",
     "At the end of a voice call, the system provides agent the option to transfer/route "
     "the call to a rating system for customer feedback.",
     "Fully",
     "The agent can transfer the caller to the post-call rating IVR before hanging up. "
     "The score is captured against the call, the agent and the case, and feeds the agent "
     "NPS reporting."),
    ("Agent Management",),
    ("4.10", "Agent Management",
     "On-duty check, channel priority settings, and automatic polling task assignment "
     "based on status.",
     "Fully",
     "All three are shipped and operator-configurable. See 3.1.6."),
    ("4.11", "Agent Management",
     "The system shall be able to check who is on duty before escalate to agent",
     "Fully",
     "Assignment evaluates live on-duty state before routing; off-duty and unavailable "
     "agents are excluded, and work is never assigned into a dead queue."),
    ("4.12", "Agent Management: Agent Availability Monitoring & Notification",
     "Monitor agent availability status (e.g. Available, Busy, Break, Toilet, Follow-up) "
     "and trigger notifications based on predefined time thresholds.",
     "Custom",
     "Availability monitoring and threshold-based notification are built in this "
     "engagement together with the extended status set from 4.17. Each status carries its "
     "own configurable time threshold and notification recipients; time in each status is "
     "recorded and reportable on the live dashboard (4.73)."),
    ("4.13", "Agent Management: Agent Status Alert",
     "Notify both the agent and administrator role when an agent remains unavailable for "
     "more than 10 minutes.",
     "Custom",
     "Configurable threshold, defaulting to the 10 minutes specified. The alert goes to "
     "the agent and to the administrator/team leader role, in-system and by desktop "
     "notification. Built with 4.12."),
    ("4.14", "Agent Management: Extended Unavailability Alert",
     "Notify the administrator role when an agent remains unavailable for more than 1 "
     "hour, allowing review of pending WIP cases assigned to that agent.",
     "Custom",
     "Configurable threshold, defaulting to 1 hour. The alert links straight to that "
     "agent's open and WIP cases so the team leader can reassign in one step. Built with "
     "4.12."),
    ("4.15", "Agent Management",
     "Must support setting channel priorities for each agent (WhatsApp priority, Email "
     "priority, Call priority etc.)",
     "Fully",
     "Per-agent channel priority is shipped and maintained by supervisors in the admin UI."),
    ("4.16", "Agent Management",
     "Automatic polling task assignment based on current agent status (Available, busy, "
     "etc)",
     "Fully",
     "Shipped. Assignment polls live agent status and current load before routing."),
    ("4.17", "Agent Management",
     "If all agents in the priority channel are busy, the system can intelligently switch "
     "and assign tasks to available agents. Statuses: Available, Busy (On Call), Lunch, "
     "Break, Coaching, Training, Toilet, Prayer",
     "Custom",
     "Overflow to an available agent outside the priority channel is shipped and working "
     "today. The eight named statuses are built in this engagement as distinct, "
     "separately configurable and separately reportable states - today the platform "
     "distinguishes Available, Busy and Offline but does not carry the Lunch, Break, "
     "Coaching, Training, Toilet and Prayer labels. Each becomes an operator-configurable "
     "status with its own threshold (4.12) and its own line in agent activity reporting."),
    ("4.18", "Agent Management",
     "Support tickets configurable with a follow up reminder date; \"Reminder\" sent on "
     "expiry and \"Timeout Warning\" popped up if not completed within 1 hour. Sound, "
     "desktop and in-system notifications.",
     "Custom",
     "Follow-up reminder dates on cases, the reminder on expiry and the escalating "
     "timeout warning after the configurable one-hour grace period are built in this "
     "engagement, delivered through the three notification channels alongside the audible "
     "alerting in 3.1.7."),
    ("AI, FAQ & Knowledge Base",),
    ("4.19", "AI FAQ Base",
     "Real-time FAQ updates, keyword-based pop-up suggestions during chat, 1-click apply "
     "& quality scoring.",
     "Custom",
     "Real-time FAQ editing, live suggestion during chat and quality scoring are shipped. "
     "One-click apply is built in this engagement - see 3.2.1."),
    ("4.20", "AI Image & Video Understanding Capability",
     "AI should analyse customer-shared images and videos to identify potential issues, "
     "understand the reported problem, and recommend relevant follow-up questions or "
     "troubleshooting steps.",
     "Fully",
     "The AI layer is natively multimodal: customer-shared images and video are passed to "
     "Gemini with the conversation context, and the model reads them directly rather than "
     "relying on the customer's description. It uses what it sees to classify the "
     "complaint, ask the right follow-up question and propose troubleshooting steps - for "
     "example reading a dashboard warning light, a charging-port condition or a "
     "damaged-part photo. The interpretation is written to the case so the agent and any "
     "escalated PIC see what the AI concluded and why.\n"
     "Accuracy on PRO-NET's own image and video types is measured in the calibration "
     "cycle under 2.2.4."),
    ("4.21", "FAQ template with AI support",
     "Integrate the FAQ knowledge base module to support the CRM team to update the FAQ "
     "content in real time.",
     "Fully",
     "The knowledge base is a first-class admin module. The CRM team creates and edits "
     "FAQ entries and uploads source documents; changes are indexed and in use "
     "immediately, with no vendor involvement, no release and no re-training step. Bulk "
     "CSV import is available for large revisions."),
    ("4.22", "FAQ template with AI support",
     "Automatically match relevant FAQ entries based on chat keywords and prompt "
     "real-time pop-up prompts to suggest related potential replies.",
     "Fully",
     "Shipped. Matching is semantic rather than literal keyword lookup, so colloquial and "
     "abbreviated phrasing still resolves to the right entry."),
    ("4.23", "AI Suggested Reply / Agent Assist",
     "AI-generated response suggestions to agents based on customer enquiries, knowledge "
     "base content, and previous resolved cases.",
     "Fully",
     "The agent assist panel drafts a reply grounded on the live conversation, the "
     "knowledge base and previously resolved cases, and cites the sources it used so the "
     "agent can verify before sending. The agent always edits and sends - the suggestion "
     "is never issued to the customer automatically on an agent-handled conversation."),
    ("4.24", "AI Sentiment Analysis",
     "Analyse customer sentiment and identify customer emotions (satisfied, neutral, "
     "dissatisfied, urgent) during interactions, and adjust the conversation tone "
     "accordingly.",
     "Fully",
     "Sentiment and urgency are classified continuously through the interaction, on chat "
     "and on voice. The AI adapts its tone accordingly, and strongly negative sentiment "
     "is a trigger in its own right - it raises an actionable case and can force handoff "
     "to a human before the customer has to ask. Sentiment is stored on the case and is "
     "reportable and usable as a routing input."),
    ("4.25", "AI-Powered Duplicate Case Detection & Merging",
     "Analyse customer enquiries, identify duplicate or related cases, and automatically "
     "consolidate them. Link cases based on mobile number, email, vehicle registration "
     "number, or chassis number; notify assigned PIC/agent when cases are merged; avoid "
     "duplicate follow-up by multiple teams.",
     "Custom",
     "Built in this engagement. Deterministic linking on the four named identifiers - "
     "mobile number, email, vehicle registration number and chassis number - is combined "
     "with AI similarity matching on the complaint content, so the same issue reported "
     "from two channels is caught even when the identifiers differ.\n"
     "Detected duplicates are surfaced for confirmation before merging by default, with "
     "automatic merge available above a configurable confidence threshold; this avoids "
     "wrongly collapsing two genuinely distinct complaints from one customer. On merge "
     "the assigned PIC and agents on both cases are notified, the histories are combined "
     "under a single case, and the merge is written to the audit trail and is "
     "reversible.\n"
     "Note: the platform does not perform duplicate detection today - this is new build, "
     "declared as such."),
    ("4.26", "AI Call Handling",
     "Supports customer inquiries and responses in multiple languages.",
     "Fully",
     "The voice AI converses in English, Bahasa Malaysia, Chinese and Tamil, detecting "
     "and following the caller's language, including mid-call switching."),
    ("4.27", "AI Call Handling",
     "Automatically summarizes customer conversations and case history at the end of a "
     "conversation.",
     "Fully",
     "A structured summary is generated at the close of every conversation on every "
     "channel - what the customer reported, what was done, what was agreed and what "
     "remains open - and written to the case. On handoff the summary travels with the "
     "conversation. Summary factual accuracy is one of the scored measures in the "
     "calibration cycle."),
    ("4.28", "FAQ template with AI support",
     "The agent can refer to the FAQ content with one click and can also provide feedback "
     "and score the FAQ quality",
     "Custom",
     "Feedback and quality scoring are shipped; one-click insertion is built in this "
     "engagement - see 3.2.1."),
    ("4.28.2", "AI Cost/Pricing Model",
     "Detail how AI usage is metered and charged, e.g. based on API calls, resolution, "
     "token volume based on call/conversation length, conversation payload, etc.",
     "Fully",
     "There is no per-resolution, per-conversation or per-seat AI charge from Devoteam. "
     "AI is consumed directly from Google Cloud on PRO-NET's own billing account and "
     "billed at Google's published list rates - by input and output tokens for text and "
     "multimodal requests, and by audio duration for real-time voice. PRO-NET holds the "
     "contract with Google and sees the itemised bill.\n"
     "This is a material commercial difference from per-resolution SaaS AI pricing: cost "
     "scales with actual usage, and no vendor uplift is applied to AI consumption.\n"
     "Cost control is built in - the platform records token and audio consumption per "
     "conversation, per channel and per language, and exposes it as a cost report with "
     "budget alerts, so PRO-NET can see the AI cost of a WhatsApp conversation against "
     "the AI cost of a voice call. Model selection is configurable per use case, so "
     "routine classification can run on a cheaper model than live voice.\n"
     "A usage-based cost projection built on PRO-NET's actual volumes is provided with "
     "the commercial submission."),
    ("Rule Engine, SOP & Escalation",),
    ("4.29", "Rule Engine & SOP",
     "Rule-based routing to PIC, email notifications with attachments, and optional "
     "WhatsApp alerts.",
     "Custom",
     "Routing, email notification and the optional WhatsApp alert are shipped; carrying "
     "case attachments onto the escalation email is built in this engagement - see 3.2.2."),
    ("4.30", "Escalation Process: Email escalation",
     "Rule engine, according to the classification/categorization of customer problems "
     "(department/business area) combined with SOP, to automatically escalate.",
     "Fully",
     "The rule engine maps the Appendix A category tree and the responsible department "
     "onto the escalation SOP. Rules are maintained by PRO-NET administrators, versioned, "
     "and evaluated in a defined precedence order so the outcome of any given case is "
     "predictable and explainable."),
    ("4.31", "Escalation Process",
     "Identify Person In Charge (PIC) based on information solicited from "
     "caller/customer e.g. dealers, vehicle number, etc.",
     "Fully",
     "PIC resolution combines case category and division, the servicing or selling dealer "
     "and the vehicle record. The department-to-PIC and dealer-to-recipient maps are "
     "maintained by administrators without vendor involvement (8.1.9). Where the dealer "
     "cannot be resolved automatically the case falls back to a configured default owner "
     "rather than going unassigned."),
    ("4.32", "Escalation Process",
     "Automatically notify and CC the relevant personnel by email (supporting "
     "attachments - photos, videos, web links)",
     "Custom",
     "Notification with the configured CC group is shipped; attachment carriage is built "
     "in this engagement - see 3.2.2."),
    ("4.33", "Escalation Process",
     "Trigger a WhatsApp message to alert the PIC (Optional)",
     "Fully",
     "Shipped and configurable per rule, so WhatsApp alerting can be enabled for "
     "high-severity categories only."),
    ("4.34", "Escalation Process",
     "Automatic escalation mechanisms configurable: 2-hour automated trigger if no "
     "response or acknowledgement from the 1st escalation; 8-hours automatic escalation "
     "to the designated higher-level manager if no response received.",
     "Fully",
     "Both timers are shipped, business-hours-aware, and configurable per category and "
     "per recipient group. See 3.2.4."),
    ("4.35", "Escalation Process",
     "Escalated cases must have defined statuses and operations",
     "Fully",
     "The case lifecycle carries New, Assigned, WIP, Pending Customer, Higher Escalation, "
     "Temporarily Closed, Resolved and Closed, with permitted transitions enforced by "
     "role. The status set and its transitions are confirmed against PRO-NET's SOP at "
     "Discovery and documented in the System Configuration Document."),
    ("4.36", "Escalation Process",
     "Escalate triggers an automatic reminder of the higher-level responsible person.",
     "Fully",
     "Shipped. The manager ladder is configurable per department and per dealer."),
    ("4.37", "Escalation Process",
     "Records of WIP and Resolved operations",
     "Fully",
     "WIP and Resolved transitions are recorded with actor, timestamp and remarks, and "
     "drive the resolution-rate and turnaround reporting in 4.58-4.60."),
    ("4.38", "Escalation Process",
     "All status changes need to record information such as Agent name, time, and remarks",
     "Fully",
     "Every transition records the actor, the timestamp, the SLA state and the remarks "
     "entered, in the append-only audit trail described at 3.2.6."),
    ("4.39", "Email Delivery Status",
     "System should prompt with alert notification if the email sending has failed, or "
     "the recipient email address is invalid.",
     "Custom",
     "Built in this engagement. Outbound escalation and notification email is tracked "
     "through delivery: hard bounces, invalid recipients and send failures raise an alert "
     "to the sending agent and to the administrator, and mark the case so a failed "
     "escalation is never mistaken for a delivered one - the failure mode this "
     "requirement exists to prevent.\n"
     "Delivery state is written to the audit trail and is reportable, so a dealer cannot "
     "be measured as unresponsive on an email that never arrived. Address validity is "
     "additionally checked when the escalation routing list is maintained (8.1.9)."),
    ("4.40", "Email Recall Function",
     "System should provide option to recall email which was already sent out.",
     "Non",
     "Declared non-compliant as literally specified, with the reason stated plainly: once "
     "an email has been accepted by an external mail server - a dealer's or a customer's "
     "- no sending system can withdraw it. Recall in Microsoft Outlook works only between "
     "mailboxes inside the same Exchange organisation and fails for external recipients "
     "and for already-read messages. Any vendor claiming true recall to external "
     "recipients is claiming something SMTP does not permit.\n"
     "What is delivered instead, and is included in scope: a configurable send-delay "
     "window (default 60 seconds, operator-adjustable) during which the agent can cancel "
     "an outgoing email before it leaves the platform, which covers the great majority of "
     "real recall requests; a one-click correction and retraction workflow that issues a "
     "clearly-marked correction into the same thread and flags the original as retracted "
     "on the case; and recall of internal notifications to PRO-NET recipients, which are "
     "delivered in-platform and can genuinely be withdrawn.\n"
     "Devoteam would rather state this limit at bid stage than accept a requirement that "
     "cannot be met at UAT."),
    ("Customer Profile & Customer 360",),
    ("4.41", "Customer Profile",
     "Trigger conditions",
     "Fully",
     "Customer identification and the 360 card are triggered on inbound call, WhatsApp "
     "message, email, social message and on manual agent lookup. Trigger conditions are "
     "configurable per channel."),
    ("4.42", "Customer Profile",
     "Customer contact system by phone or WhatsApp (based on caller number or WA number)",
     "Fully",
     "Identification by caller number and WhatsApp number is shipped, with number "
     "normalisation so a customer is matched whether the number carries the +60 country "
     "code or a local prefix."),
    ("4.43", "Customer Profile",
     "The system calls the DMS and TSP systems through API interfaces to automatically "
     "match and pull Customer Information, Vehicle Information, Service History and Call "
     "Center History (call, WA, email, social, RSA)",
     "Custom",
     "Call Centre History - every call, WhatsApp thread, email, social reply and RSA "
     "incident - is held natively and is available today. Customer, Vehicle and Service "
     "History are pulled from the DMS and TSP through the connectors built in this "
     "engagement. See 3.3.1."),
    ("4.44", "Customer 360 View",
     "Auto pop-up card with Personal Info, Vehicle Info, Service History, and Call Center "
     "History.",
     "Custom",
     "Card shipped with personal information and call centre history live; vehicle and "
     "service sections and the automatic pop-up are delivered with the DMS/TSP connectors. "
     "See 3.3.2."),
    ("4.45", "Customer Profile",
     "After a customer calls, the Agent page automatically pops up the Customer 360 View "
     "Card",
     "Custom",
     "Automatic screen-pop on inbound call is built in this engagement; the card itself "
     "and manual lookup are shipped."),
    ("4.46", "Customer Profile",
     "The content of the card is divided into personal information, vehicle information, "
     "service information, and Call Center history",
     "Custom",
     "The card is structured as the four named sections. Personal information and call "
     "centre history render today; vehicle and service information render once the DMS/TSP "
     "connectors are live. Field-level content of each section is agreed in the data "
     "contract at Discovery."),
    ("4.47", "Customer Profile",
     "All data synchronization should take no more than 3 seconds and support "
     "asynchronous loading",
     "Custom",
     "The card renders immediately from locally-held data and each external section loads "
     "asynchronously, so the agent is never blocked waiting on an external system. The "
     "3-second target is committed for the platform's own processing and for DMS/TSP "
     "responses within their agreed response-time SLA.\n"
     "Stated honestly: end-to-end time depends in part on the DMS and TSP response times, "
     "which are PRO-NET-owned systems. A response-time expectation for each is agreed in "
     "the data contract at Discovery, caching and a circuit breaker protect the agent "
     "experience if a source system is slow, and actual response times are measured and "
     "reported so any breach is attributable to the correct system."),
    ("Reporting & Business Intelligence",),
    ("4.48", "Reporting",
     "Built-in reports or provision integration to Power BI (or equivalent BI tools) for "
     "the data visualization reports listed below",
     "Fully",
     "Both are delivered. Operational reports run in-platform for supervisors and agents. "
     "A BigQuery analytics warehouse carries the full case, interaction, SLA, CSAT/NPS and "
     "AI-performance history and is the source for Power BI. Devoteam builds the Power BI "
     "workspace, dataset and the reports listed at 4.49-4.82, reproducing the formats "
     "already in use in Appendix C1 and C2."),
    ("4.49", "Reporting",
     "Channel source analysis by case (Call/WA/Email or Social Media)",
     "Fully", "Delivered, built to the Appendix C sample format."),
    ("4.50", "Reporting",
     "By case division (Apps, Sales, Product Aftersales, Charging)",
     "Fully",
     "Delivered against the Appendix A division structure, with drill-down through Level "
     "1 to Level 4."),
    ("4.51", "Reporting", "Daily/weekly/monthly trend chart",
     "Fully", "Delivered with a customizable date range per 4.81."),
    ("4.52", "Reports", "After-hours case volume report",
     "Fully", "Driven from the in-hours/after-hours stamp described at 3.1.3."),
    ("4.53", "Reports",
     "Response time measurement based on business hours vs non-business hours",
     "Fully",
     "Both measures are reported side by side - elapsed clock time and business-hours "
     "time - so after-hours arrivals are visible without distorting SLA attainment."),
    ("4.54", "Reports", "SLA calculation that considers operating hours",
     "Fully",
     "SLA clocks pause outside configured business hours and across public holidays; the "
     "calculation basis is stated on the report so the figures are auditable."),
    ("4.55", "Power BI Analytics",
     "Power BI integration for channel source, division, trend, PIC, reopening rate "
     "(CRR), and SLA stats",
     "Fully",
     "A governed Power BI dataset over BigQuery with row-level security aligned to the "
     "platform's roles, refreshed on a schedule agreed at Discovery. PRO-NET's own BI team "
     "can build additional reports on the same dataset without vendor involvement."),
    ("4.56", "AI Performance & Effectiveness Reporting",
     "1. AI Case Resolution Report; 2. AI vs Human Handling Report; 3. AI Response "
     "Satisfaction Report; 4. AI Escalation Report; 5. AI Accuracy & Improvement Report; "
     "6. AI Deflection Rate; 7. AI Root Cause Analysis; 8. AI Knowledge Base Improvement",
     "Custom",
     "Every AI decision is already recorded with its inputs, retrieved knowledge, chosen "
     "action, confidence and outcome, so all eight reports are built from production "
     "evidence rather than estimated.\n"
     "Shipped today: AI vs human handling volumes, deflection rate, escalation tracking "
     "with reasons, AI response satisfaction from the CSAT split, and knowledge-base "
     "improvement suggestions driven by unresolved-enquiry and repeated-search analysis.\n"
     "Built in this engagement: the AI Root Cause Analysis report (7), which clusters "
     "complaints to surface recurring issues, emerging trends and probable root causes "
     "across the Appendix A taxonomy, and the consolidated AI accuracy report (5) "
     "presenting the calibration scorecard from 2.2.4 as a tracked trend rather than a "
     "point-in-time result.\n"
     "All eight are delivered as Power BI reports and in-platform views."),
    ("4.57", "Reporting", "Analysis of departments and responsible persons",
     "Fully", "Delivered by department and by named PIC."),
    ("4.58", "Reporting",
     "Case distribution and resolution statistics by dealer/department/PIC",
     "Fully", "Delivered across all three dimensions."),
    ("4.59", "Reporting: Dealer Escalation",
     "Ranking of first response rate (2-hr response rate), resolution rate and resolution "
     "time of each department/dealer",
     "Fully",
     "The dealer escalation league table, measured from the escalation timestamp stamped "
     "on the case when the dealer escalation is raised (3.2.4). Because that timestamp is "
     "recorded by the platform rather than inferred, the 2-hour first-response rate and "
     "the turnaround figures are defensible when a dealer disputes them."),
    ("4.60", "Reporting",
     "Analysis the Complaint Reopen Rate, CRR by dealers/department/PIC",
     "Fully",
     "CRR reported across all three dimensions, with the reopen reason captured at the "
     "point of reopening so the rate can be explained and not merely counted."),
    ("4.61", "Reporting", "Case aging from timestamp created",
     "Fully", "Aging buckets are configurable; the default banding matches 4.69."),
    ("4.62", "Reporting", "WIP weekly case. Refer to sample report.",
     "Fully", "Built to the Appendix C2 weekly WIP format."),
    ("4.63", "KPIs & NPS",
     "Call Centre KPIs (SLA, agent workload, response times) and Agent NPS rating from "
     "Call/WA",
     "Fully", "Delivered to the Appendix C sample formats."),
    ("4.64", "Reporting", "Call Centre KPI Reports",
     "Fully",
     "The full call centre KPI pack: SLA attainment, abandon rate, AHT, ACW, occupancy, "
     "first-response and closure times, per agent, per team and per channel."),
    ("4.65", "Reporting", "SLA response achievement rate",
     "Fully", "Reported by channel, category, dealer, PIC and agent."),
    ("4.66", "Reporting: Phone calls",
     "Abandon rate; calls which customer dropped off while put on hold (in queue)",
     "Fully",
     "Abandon rate is captured from telephony queue events, with time-to-abandon "
     "distribution so PRO-NET can see whether abandons cluster before or after the "
     "20-second answer target in Appendix B."),
    ("4.67", "Reporting",
     "The number of tasks processed by each agent and the average response time",
     "Fully", "Built to the Appendix C sample format."),
    ("4.68", "Reporting",
     "The average response duration (first response time) and the average case closure "
     "time, based on the channel response time",
     "Fully", "Reported per channel, business-hours aware."),
    ("4.69", "Reporting",
     "1. Average Call Handling Time (AHT); 2. After Call Work (ACW) - time on post-call "
     "updates, documentation, categorisation and follow-up; 3. Resolution Time "
     "distribution (same day / 1-2 days / 3-5 days / 5+ days)",
     "Fully",
     "AHT and the resolution-time distribution use the bands as specified. ACW is measured "
     "from the platform's own record of post-call activity - the wrap-up period between "
     "call end and the agent returning to available, together with the case updates made "
     "in that window - so it reflects work actually done rather than a self-reported "
     "status."),
    ("4.70", "Reporting",
     "Ranking of popular complaint types and statistics during peak complaint hours",
     "Fully",
     "Complaint-type ranking across the Appendix A taxonomy with hour-of-day and "
     "day-of-week heat mapping for staffing decisions."),
    ("4.71", "Reporting", "NPS For Agent",
     "Fully", "Per-agent NPS from the call and WhatsApp rating flows."),
    ("4.72", "Reporting",
     "Customer rating of agent performance from Call & WA channel",
     "Fully",
     "Reported per agent and per channel, and separated from the AI rating so the two are "
     "never conflated."),
    ("4.73", "Live Dashboard",
     "Visibility of agent working days, login/logout records, availability status, and "
     "overall activity performance",
     "Custom",
     "Login/logout records, live availability and activity performance are shipped in the "
     "supervisor dashboard. The extended per-status time breakdown (4.17) is added with "
     "that work, so a team leader can see time in Lunch, Break, Coaching, Training, Toilet "
     "and Prayer separately rather than as one unavailable block."),
    ("4.74", "Lifecycle & Exports",
     "Case life cycle map, anomaly warning dashboard, requested tag keyword reports, auto "
     "PDF/Excel email dispatch",
     "Fully", "All four are shipped. See 4.75-4.80."),
    ("4.75", "Reporting", "Case life cycle tracking",
     "Fully", "Every state transition is tracked with actor, timestamp and duration in state."),
    ("4.76", "Dashboard", "A map of the time distribution from creation to closing",
     "Fully",
     "The lifecycle map shows where time is actually spent between creation and closure, "
     "which is what identifies whether delay sits with the agent, the dealer or the "
     "customer."),
    ("4.77", "Reporting",
     "Trend analysis of special states such as Higher escalation / WIP / Temporary Closed "
     "/ Closed",
     "Fully", "Trended over any selected period across all lifecycle states."),
    ("4.78", "Reporting",
     "Reports exported to PDF and Excel, and automatically sent to management on a regular "
     "basis",
     "Fully",
     "Scheduled report distribution by email in PDF and Excel, with per-recipient "
     "schedules and per-recipient content scoping so a dealer receives only their own "
     "data."),
    ("4.79", "Reporting",
     "Configurable \"Anomaly Warning Dashboard\" for real-time anomaly prompts (such as "
     "the explosion of a channel)",
     "Fully",
     "Anomaly detection runs against live volumes with operator-configurable thresholds "
     "and baseline-deviation detection, so a channel spike, a category spike or an SLA "
     "breach cluster raises an alert in real time rather than being found in the next "
     "monthly report."),
    ("4.80", "Reporting", "Run report based on request tag keyword",
     "Fully", "Tag and keyword driven reporting across cases and interaction content."),
    ("4.81", "Reporting: General",
     "Real-time and historical reporting with customizable date range selection, "
     "analysable by agent, team, date period and customer interaction channel",
     "Fully",
     "All reports carry a custom date range and filter by agent, team, period and channel. "
     "Operational views are real-time; the analytics warehouse carries the full history for "
     "the 7-year retention period."),
    ("4.82", "Reporting: Export",
     "All reports shall be able to export out into CSV, XLS and PDF format",
     "Fully", "All three formats on every report, on demand and on schedule."),
    ("Account Management & Data",),
    ("4.83", "Account Management",
     "Configurable user profile with customizable permissions based on functions and data",
     "Fully",
     "Role-based access control with custom roles is shipped. Permissions are granular by "
     "function (which screens and actions) and by data (which inboxes, teams, dealers and "
     "categories a user may see), so a dealer user, an agent, a team leader, a report "
     "consumer and an administrator each see only what their role permits. Roles are "
     "maintained by PRO-NET administrators and every permission change is audited."),
    ("4.84", "Data Retention Policy",
     "System must retain all operations data for a minimum period of 7 years",
     "Fully",
     "7-year retention is met for cases, interactions, transcripts, call recordings, "
     "attachments, audit records and reporting history.\n"
     "Because the platform is self-managed on PRO-NET's own Google Cloud tenancy, "
     "retention is not constrained by a SaaS vendor's storage tier or data-export policy, "
     "and there is no per-GB archive surcharge. Storage lifecycle rules move ageing data "
     "to lower-cost classes while keeping it queryable, so the 7-year obligation is met "
     "without the cost growing linearly.\n"
     "Retention locks prevent premature deletion, including by an administrator. A "
     "documented purge process handles PDPA erasure requests where a lawful basis for "
     "erasure exists, and records the action in the audit trail."),
]

# ---------------------------------------------------------------- 6.0
SEC6 = [
    ("6.1  Delivery Methodology",),
    ("Delivery Approach",
     "Vendor to indicate the Delivery Methodology that will be applied, e.g. "
     "Agile-Waterfall Hybrid Approach.",
     "Fully",
     "Agile-Waterfall hybrid, matched to what each part of the work actually needs.\n"
     "Waterfall governs the fixed spine: Discovery and requirements sign-off, the "
     "architecture blueprint, the stage gates and the milestone sign-offs that the "
     "payment schedule in section 10 depends on. These are sequential and formally "
     "signed.\n"
     "Agile governs build and configuration: one-week sprints, each ending in a working "
     "demonstration on the UAT tenant against real PRO-NET scenarios, with the CRM team "
     "reviewing as it is built rather than at the end. Every sprint declares its outputs, "
     "so scope, sprints and deliverables trace to each other line by line.\n"
     "This split matters for a complaint system: the escalation matrix and the case "
     "taxonomy always change once operators see them working, and an iterative build "
     "absorbs that without a change request, while the contractual gates stay fixed."),
    ("Scope & Change Management",
     "Indicate the process for Scope & Change Management should this arise during the "
     "project.",
     "Fully",
     "Baseline: the signed BRD and Technical Design Document. Anything outside them is a "
     "change request.\n"
     "Process: the request is logged on a shared register by either party; Devoteam "
     "assesses effort, schedule and risk impact within 3 business days; the Project "
     "Working Group reviews; the Project Steering Committee approves anything affecting "
     "cost, schedule or the payment milestones. Nothing is built before written approval, "
     "and no approved change silently moves a milestone date.\n"
     "Minor clarifications that do not affect effort, schedule or cost are absorbed at "
     "working-group level without a formal change request, so the process does not become "
     "an obstacle to ordinary refinement. A running change log forms part of the monthly "
     "status report and the project close-out document."),
    ("Quality Assurance & Risk Management",
     "Indicate the QA & Risk Mitigation strategies applied for this project.",
     "Fully",
     "Quality assurance: definition of done per sprint including tests and documentation; "
     "peer review on every change; automated regression suites run on each build; "
     "environment parity across Development, UAT and Production from shared "
     "infrastructure-as-code; formal SIT with a PRO-NET-agreed script and a signed exit; "
     "then UAT with PRO-NET's own acceptance criteria.\n"
     "Risk management: a RAID log maintained from kick-off and reviewed weekly, each risk "
     "with an owner, a probability and impact rating and a mitigation.\n"
     "The four risks carried from day one, stated openly: (1) DMS/TSP API access and "
     "specification, the single largest schedule dependency, mitigated by requesting "
     "sandbox access in week one, building against a documented contract and stubbing the "
     "connector so the rest of the build is never blocked; (2) legacy data quality, "
     "mitigated by early profiling and a full dress-rehearsal migration; (3) AI accuracy "
     "in Bahasa Malaysia and especially Tamil, mitigated by the measured calibration cycle "
     "and human-in-the-loop review before any auto-send behaviour is enabled; (4) UAT "
     "availability of frontline agents during operational peaks, mitigated by scheduling "
     "UAT windows with the call centre at kick-off."),
    ("6.2  Project Governance & Communication Structure",),
    ("6.2.1 Project Governance Organization",
     "Propose the project governance organization, e.g. project steering committee, "
     "project working group, technical team.",
     "Fully",
     "Three tiers.\n"
     "Project Steering Committee - PRO-NET executive sponsor and the Devoteam engagement "
     "lead. Meets at each stage gate and monthly. Owns milestone acceptance, approved "
     "change requests, escalated risks and commercial decisions.\n"
     "Project Working Group - PRO-NET project manager and CRM/call-centre process owners "
     "with the Devoteam Project Manager and Solution Architect. Meets weekly. Owns scope "
     "clarification within baseline, the RAID log, sprint acceptance and inter-team "
     "dependencies.\n"
     "Technical Team - Devoteam Solution Architect, engineers and QA with PRO-NET's IT, "
     "integration and security counterparts. Meets as needed and daily during hypercare. "
     "Owns design decisions, integration contracts, environments and defect resolution.\n"
     "Devoteam names a single Project Manager as PRO-NET's point of contact for the whole "
     "engagement.\n"
     "{{TBD - Devoteam named project team (Project Manager, Solution Architect, "
     "Engagement Lead) to be inserted before submission.}}"),
    ("6.3  Progress Reporting & Transparency Tools",),
    ("6.3.1 Project Dashboard / Status Reports",
     "Indicate the reporting method and cadence.",
     "Fully",
     "Weekly written status report to the Project Working Group - progress against plan, "
     "sprint outputs delivered, next week's plan, open risks and issues with owners, "
     "decisions required from PRO-NET, and the change log.\n"
     "Monthly executive summary to the Steering Committee - milestone and budget status, "
     "schedule confidence, top risks and the decisions being asked for.\n"
     "A live project dashboard is shared with PRO-NET from kick-off showing the plan, "
     "sprint board, RAID log, defect counts by severity and deliverable status, so "
     "PRO-NET can see project state at any time without waiting for a report. Sprint "
     "demonstrations every week on the UAT tenant are the primary transparency mechanism: "
     "working software reviewed by the people who will use it."),
    ("6.3.2 Milestone Sign-Off Artifacts",
     "Evidence for critical project deliverables: 1. Project Kick Off Deck; 2. Business "
     "Requirements Document; 3. Technical Design Document; 4. SIT Report; 5. UAT Report; "
     "6. Production Deployment Check List; 7. System Configuration Document; 8. Training "
     "Manual; 9. System Operation Manual; 10. Project Close Document",
     "Fully",
     "All ten are committed, each mapped to the phase that produces it and to the payment "
     "milestone it evidences:\n"
     "1. Project Kick Off Deck - Kick Off, evidences Mobilization (15%).\n"
     "2. Business Requirements Document - Discovery, evidences Requirements Sign-off "
     "(15%).\n"
     "3. Technical Design Document - Discovery/Architecture Blueprint, evidences "
     "Requirements Sign-off (15%).\n"
     "4. SIT Report - Testing phase, gates entry to UAT.\n"
     "5. UAT Report - UAT phase, evidences UAT Sign-off (40%).\n"
     "6. Production Deployment Checklist - Cutover, evidences Go-Live (20%).\n"
     "7. System Configuration Document - Cutover, evidences Go-Live (20%).\n"
     "8. Training Manual - Training phase, delivered before go-live.\n"
     "9. System Operation Manual - Handover, delivered at go-live.\n"
     "10. Project Close Document, with the Project Completion Certificate and the snag/"
     "punch list - Close Out, evidences the final 10%.\n"
     "Each is issued in draft for PRO-NET review before the formal sign-off request, so a "
     "gate is never the first sight of a document."),
]

# ---------------------------------------------------------------- 7.0
SEC7 = [
    ("7.1 License type",
     "Concurrent, floating user, named user, provision for standard and light licences "
     "e.g. light for administrator use for viewing reports.",
     "Fully",
     "There is no per-seat software licence in this solution, which is the most "
     "significant commercial point in this section. The CRM core is open-source "
     "(Chatwoot Community, MIT licence) and is self-hosted in PRO-NET's Google Cloud "
     "tenancy; the AI layer is consumed from Google Cloud on PRO-NET's own billing "
     "account. PRO-NET may create unlimited named users, unlimited light/read-only "
     "users and unlimited administrators at no incremental licence cost.\n"
     "Access tiers are therefore implemented as roles rather than as purchased licence "
     "types: full agent, supervisor, administrator, and a light read-only role for report "
     "viewing - exactly the distinction 7.1 describes, without the licence charge attached "
     "to it.\n"
     "Cost scales with infrastructure and AI consumption, not with headcount, so adding "
     "agents at peak periods carries no licensing consequence.\n"
     "Third-party pass-through charges remain: WhatsApp Business API and telephony "
     "(Twilio, per-message and per-minute), Meta messaging, and Power BI Pro or Premium "
     "seats for report authors under PRO-NET's existing Microsoft agreement."),
    ("7.2 License Encryption & Verification Standards",
     "Cryptographic Key Binding",
     "Non",
     "Not applicable as specified, because there is no proprietary licence to enforce. "
     "Licence key binding, activation servers and entitlement enforcement exist to stop "
     "customers exceeding a purchased seat count; with an open-source, self-hosted core "
     "there is no seat count to enforce and no licence artefact to bind or encrypt.\n"
     "The security controls this requirement is ultimately concerned with are delivered "
     "through identity and cryptography rather than licence enforcement: all data "
     "encrypted at rest with Cloud KMS customer-managed keys and in transit with TLS 1.2+, "
     "secrets held in Secret Manager with rotation, signed container images with "
     "deployment-time verification so only Devoteam-built artefacts can run, and full "
     "administrative audit.\n"
     "Stated rather than answered as compliant, so the evaluator can price the difference "
     "correctly."),
    ("7.3 Login Control / IAM",
     "Basic Username and Password verification, MFA, etc.",
     "Custom",
     "Username and password authentication with the policy controls at 7.6 is shipped. "
     "Multi-factor authentication and single sign-on against PRO-NET's corporate identity "
     "provider are configured in this engagement.\n"
     "Preferred approach: federate to PRO-NET's existing identity provider so that "
     "joiners, movers and leavers are governed by PRO-NET's own IAM process and MFA policy "
     "rather than by a second user directory in the CCMS - which also removes the risk of "
     "an orphaned CCMS account after an employee leaves. Local MFA (TOTP authenticator "
     "app) is delivered for users outside the corporate directory, such as dealer users.\n"
     "{{TBD - PRO-NET corporate identity provider and federation protocol (e.g. Entra ID "
     "/ SAML 2.0 / OIDC); PRO-NET IT to confirm. Raised as a clarification question.}}"),
    ("7.4 Verification Mode",
     "Online Validation / Offline Cache",
     "Fully",
     "Authentication is validated online against the identity provider at sign-in, with "
     "short-lived session tokens and silent renewal so an agent is not re-prompted through "
     "a shift. Sessions are revocable centrally and immediately - an administrator can "
     "terminate a compromised session without waiting for a token to expire.\n"
     "No offline credential cache is used, deliberately: a cached credential on an agent "
     "workstation is a data-protection exposure that a browser-delivered application does "
     "not need to carry. The agent workspace requires connectivity in any case, since it "
     "is handling live customer interactions."),
    ("7.5 Session Heartbeat & Concurrency Polling",
     "Polling Protocols & Idle Time-out Rules",
     "Fully",
     "The agent workspace holds a persistent connection to the platform, which doubles as "
     "the session heartbeat: it drives live queue and presence updates and detects a "
     "dropped agent within seconds, so work is not routed to a browser that has closed.\n"
     "Idle timeout, the warning period before it, and the automatic status change on idle "
     "are all operator-configurable, and integrate with the availability monitoring in "
     "4.12-4.14.\n"
     "Concurrent sessions per user are policy-controlled - blocked, or permitted with "
     "visibility to administrators - and every session start and end is written to the "
     "audit trail and to the agent login/logout report at 4.73."),
    ("7.6 Password Policy",
     "Allow saving the password on the desktop machine so the user does not need to key it "
     "in each time; prompt to change password every 3 months with user self-service change.",
     "Custom",
     "Configurable password policy including the 90-day change prompt, complexity, history "
     "and lockout rules, with user self-service password change and reset. Delivered as "
     "configuration.\n"
     "On saving the password to the desktop machine: where PRO-NET federates to its "
     "corporate identity provider (7.3), agents get a better outcome than a saved "
     "password - single sign-on from the already-authenticated workstation, so no password "
     "is keyed in at all and none is stored on the machine. Where federation is not used, "
     "\"remember this device\" is supported through a long-lived trusted-device token "
     "rather than by storing the password itself.\n"
     "Devoteam's recommendation, offered plainly: storing a reusable password on a shared "
     "call-centre workstation is a control PRO-NET's own security policy is likely to "
     "reject, and single sign-on satisfies the underlying requirement - agents not "
     "re-keying credentials - without it. Final configuration follows PRO-NET IT policy."),
    ("7.7 License provisioning",
     "System should provision for adding of user licenses on individual user and monthly "
     "subscription basis",
     "Fully",
     "Users are created, modified and deactivated by PRO-NET administrators at any time, "
     "individually or in bulk, with no licence purchase, no vendor request and no waiting "
     "period - and at no incremental licence cost, per 7.1.\n"
     "Because there is no seat-based subscription, there is nothing to true-up monthly and "
     "no over-deployment exposure. Where PRO-NET wants monthly visibility for internal "
     "chargeback, active-user counts by role, team and department are reported monthly."),
]

# ---------------------------------------------------------------- 8.0
SEC8 = [
    ("8.1  Support Channels & Maintenance Framework",),
    ("8.1.1 Support Organization",
     "Provide the support organization. For critical issues, support is required on-site "
     "to work closely with the CRM team and call centre to manage workaround.",
     "Fully",
     "Three tiers. L1 service desk - ticket intake, triage, severity assignment, known-"
     "error resolution and communication. L2 application support - configuration, "
     "knowledge-base and AI calibration, escalation-map and category maintenance, "
     "integration and report issues. L3 engineering - the engineers who built the "
     "platform, owning defect fixes, performance and infrastructure.\n"
     "On-site attendance is committed for every Critical (P1) incident, working alongside "
     "the CRM team and the call centre to establish a workaround first and the permanent "
     "fix after, as this requirement asks. On-site presence is also committed for the "
     "first week of hypercare and for each major release.\n"
     "A named Service Delivery Manager owns the relationship, the SLA reporting and the "
     "escalation ladder.\n"
     "{{TBD - Devoteam named support organisation: Service Delivery Manager, Technical "
     "Lead, Technical Account Manager and both L3 escalation contacts, with contact "
     "emails. To be inserted before submission.}}"),
    ("8.1.2 Helpdesk Access",
     "Access to support across multi-channels: Email, phone, WhatsApp",
     "Fully",
     "All three, plus the ticket portal at 8.1.4. Every channel raises a ticket "
     "automatically, so a WhatsApp message to the support team is tracked and SLA-measured "
     "exactly like a portal ticket and nothing is lost in a chat thread.\n"
     "{{TBD - Devoteam support email address, support hotline and support WhatsApp number. "
     "To be inserted before submission.}}"),
    ("8.1.3 Support Availability",
     "Support must be contactable following PRO-NET business hours. Mon-Fri 8:00AM - "
     "5:30PM",
     "Fully",
     "Staffed support Monday to Friday 08:00-17:30 Malaysia time (MYT, UTC+8), aligned to "
     "PRO-NET business hours, with the team located in-region rather than in a distant "
     "follow-the-sun centre.\n"
     "Noted for resolution before contract: section 9.0 states \"Support must be available "
     "24/7\" while 8.1.3 specifies business-hours contactability. Devoteam's proposal "
     "reconciles the two as follows - staffed business-hours support for all severities, "
     "plus a 24/7 on-call rotation for Critical (P1) and Major (P2) incidents so the "
     "section 9 response targets are met at any hour, and 24/7 automated monitoring and "
     "alerting on the platform and the RSA path. Full 24/7 staffed coverage across all "
     "severities is available as a priced option. Raised as a clarification question."),
    ("8.1.4 Support Ticket Management",
     "Support ticketing portal to provide visibility and tracking of cases",
     "Fully",
     "A support portal gives PRO-NET raise, track and history functions with live SLA "
     "status on every ticket, visibility of every ticket raised by any PRO-NET user rather "
     "than only the raiser's own, and self-service SLA performance reporting so PRO-NET "
     "can audit attainment independently rather than relying on the vendor's own monthly "
     "figure.\n"
     "{{TBD - Devoteam support portal URL. To be inserted before submission.}}"),
    ("8.1.5 System Uptime Guarantee",
     "The cloud platform must maintain a 99.9% uptime baseline, excluding scheduled "
     "maintenance.",
     "Fully",
     "99.9% monthly uptime is committed, excluding the agreed maintenance window.\n"
     "It is engineered rather than promised: the application runs multiple replicas across "
     "availability zones on Google Kubernetes Engine, the database is a regional "
     "high-availability instance with automatic failover, and deployments are rolling with "
     "no planned downtime. Uptime is measured by independent external monitoring of the "
     "agent workspace and each channel endpoint, and the measurement is visible to "
     "PRO-NET, so attainment is not self-reported.\n"
     "Stated for completeness: the underlying Google Cloud services carry their own "
     "provider SLAs, and the composite availability target is set to keep the platform "
     "commitment achievable above them."),
    ("8.1.6 Scheduled Maintenance",
     "Scheduled maintenance can only be carried out between 00:00 to 04:00 hrs. RSA must "
     "be available 24/7 and unaffected by scheduled maintenance.",
     "Fully",
     "Accepted. All scheduled maintenance falls within 00:00-04:00 MYT, notified at least "
     "5 business days in advance with a stated impact and rollback plan, and emergency "
     "security patching is the only exception, notified immediately.\n"
     "The RSA path is engineered to stay up through maintenance: RSA call routing and "
     "incident logging run independently of the components that maintenance touches, and "
     "rolling deployment means no component is taken fully offline. RSA availability is "
     "verified as an explicit step in every maintenance runbook and confirmed in the "
     "post-maintenance report."),
    ("8.1.7 KB Maintenance",
     "Add, remove, updating the KB as required by business",
     "Fully",
     "PRO-NET maintains the knowledge base directly and in real time - this is deliberately "
     "not a vendor-gated task, because a complaint operation cannot wait on a vendor to "
     "correct an FAQ. Devoteam supports it under the managed service with structural "
     "review, gap analysis from unresolved enquiries, and bulk import assistance for large "
     "revisions."),
    ("8.1.8 AI Calibration & Training",
     "Calibration of speech to text transcription and AI responses on language adaptation, "
     "sentiment/emotion factor, KB completeness and accuracy, and accuracy of conversation "
     "summary",
     "Fully",
     "The measured calibration cycle described at 2.2.4 continues as a standing managed-"
     "service activity: monthly by default, weekly after any significant change to the "
     "knowledge base, the AI persona or the model. Each cycle produces a written scorecard "
     "against the agreed baselines, the change set applied, and the re-scored result, "
     "reviewed with PRO-NET in the monthly account review."),
    ("8.1.9 Updating escalation group/list",
     "To Add, remove, update list of dealers, escalation PIC as business requires",
     "Fully",
     "PRO-NET administrators maintain the dealer list, the PIC map and the recipient groups "
     "directly in the admin UI, with immediate effect and full audit. No vendor "
     "involvement, no release, no charge. Devoteam supports bulk updates and periodic "
     "review of routing accuracy under the managed service."),
    ("8.1.10 Case categorization",
     "To Add, remove, update list for categories as business requires",
     "Fully",
     "Operator-maintained hierarchical categories with bulk import, as described at 2.2.1. "
     "Historical cases retain their original category so reporting stays coherent across a "
     "taxonomy change."),
    ("8.1.11 System Performance Health Check",
     "Proactive System Monitoring, system resources, storage, etc",
     "Fully",
     "Continuous monitoring of application and infrastructure health - availability, "
     "latency, error rates, queue depth, database and cache performance, storage growth "
     "against the 7-year retention curve, certificate expiry and third-party channel "
     "health - with alerting on threshold breach to the on-call engineer before it becomes "
     "a user-visible incident.\n"
     "A monthly system health report goes to PRO-NET with capacity trends and a "
     "forward-looking view of when any resource will need attention, so growth is planned "
     "rather than discovered."),
    ("8.1.12 System Backup and Disaster Recovery Plan",
     "Detail the backup and DR strategy, frequency, full, incremental",
     "Fully",
     "Database: automated full backup daily, with continuous transaction-log archiving "
     "giving point-in-time recovery to any moment within the retention window. Object "
     "storage (attachments, call recordings, transcripts): versioned with immediate "
     "cross-region replication. Configuration and infrastructure: held as "
     "infrastructure-as-code in version control, so the environment is rebuildable rather "
     "than only restorable.\n"
     "Backups are encrypted with customer-managed keys, stored in a separate region, and "
     "retained on a schedule agreed at Discovery.\n"
     "Proposed targets are RPO 15 minutes and RTO 4 hours, confirmed against PRO-NET's "
     "business-continuity policy. A restore test is executed and evidenced before go-live "
     "and repeated every six months under the managed service, with the result reported - "
     "an untested backup is not a recovery capability."),
    ("8.1.13 Software Updates",
     "Regular software updates and security patches",
     "Fully",
     "Security patches assessed on release and applied within the P1/P2 windows for "
     "critical vulnerabilities. Routine updates batched into a scheduled release, tested in "
     "UAT first, deployed only in the agreed maintenance window, with release notes issued "
     "in advance and a rollback path for every release. No licence or upgrade fee, per "
     "2.4.4."),
    ("8.1.14 System Bug Fixes, Enhancement Change Request",
     "Details the process for change requests",
     "Fully",
     "Defects and enhancements are handled on separate tracks so they do not compete.\n"
     "Defects - anything not behaving per the signed BRD/FSD - are raised on the support "
     "portal, triaged to the section 9 severity model and fixed at no charge during "
     "warranty and under the managed service.\n"
     "Enhancements follow the change-request process: written request, Devoteam impact "
     "assessment (effort, schedule, risk, cost) within 3 business days, PRO-NET approval, "
     "then scheduling into a release. Where a managed-service change pool is contracted, "
     "small enhancements are drawn from it without a separate commercial cycle. All "
     "changes are tested in UAT and included in the monthly account review."),
    ("8.1.15 Account Management",
     "Monthly review of system performance, availability, AI accuracy, KB health status, "
     "support SLA performance, system health report, insights for improvement",
     "Fully",
     "A monthly service review chaired by the Service Delivery Manager, covering exactly "
     "the agenda this requirement sets: platform performance and availability against the "
     "99.9% target; the AI calibration scorecard and accuracy trend; knowledge-base health "
     "including coverage gaps and stale entries; support SLA attainment by severity with "
     "any breach explained; the system health and capacity report; and a written set of "
     "improvement recommendations drawn from the period's data - recurring complaint root "
     "causes, FAQ entries worth adding, routing rules worth changing.\n"
     "The review produces a documented action list with owners and dates, followed up at "
     "the next review, so it functions as a service-improvement loop rather than a "
     "reporting formality."),
]

# ---------------------------------------------------------------- 5.0 timeline
TIMELINE = [
    ("Project Kick Off", "Week 1",
     "Project kick-off workshop with the Steering Committee and Working Group; governance "
     "and communication plan; RAID log opened; environment and access requests raised, "
     "including DMS/TSP sandbox credentials; delivery plan and acceptance criteria agreed.",
     "Project Kick Off Deck\n(Mobilization - 15%)"),
    ("Discovery & Architecture Blueprint", "Weeks 1-2",
     "Requirements workshops against the RFP, Appendix A taxonomy and the Appendix B "
     "channel SOPs; develop-vs-integrate confirmed per component; DMS/TSP data contract "
     "and response-time expectations agreed; PII classification and masking rules; solution "
     "architecture and GCP landing-zone design; migration approach and legacy data "
     "profiling; Test Plan drafted.",
     "Business Requirements Document\nTechnical Design Document\n(Requirements Sign-off - 15%)"),
    ("Core Platform & Integration Build", "Weeks 2-6",
     "GCP landing zone, GKE cluster, HA database, networking, CI/CD, secrets and monitoring "
     "baseline; platform deployed across Dev/UAT/Prod; all five channels wired (WhatsApp, "
     "email, IVR/telephony, Facebook Messenger, Instagram); case taxonomy, SLA policies, "
     "business hours, PIC/dealer routing and the escalation ladder configured; DMS/TSP "
     "connectors and the Customer 360 screen-pop; agent management build (extended statuses, "
     "availability alerts, reminders and audible alerting); duplicate detection and merging; "
     "email delivery tracking and the send-delay/retraction workflow; MFA/SSO federation.",
     "Sprint demonstrations\nAPI Integration Document\nSystem Configuration Document (draft)"),
    ("AI Engine & Power BI Analytics", "Weeks 4-7",
     "AI persona, knowledge base and FAQ loaded from PRO-NET content; four-language "
     "coverage; agent-facing translation surface; multimodal image and video handling; "
     "sentiment and summarisation; AI activity monitor console; first calibration cycles "
     "with scorecards; BigQuery warehouse and semantic layer; Power BI workspace and the "
     "full report set at 4.49-4.82 built to the Appendix C1/C2 formats; scheduled "
     "distribution and the anomaly dashboard.",
     "AI Calibration Scorecard\nPower BI report pack\nSprint demonstrations"),
    ("Testing, Training & UAT", "Weeks 7-9",
     "SIT executed against the agreed script - functional, all five Appendix B channel "
     "flows, escalation matrix and timers, integration contracts, RBAC, performance against "
     "the 3-second Customer 360 budget, and failure paths; SIT exit signed; role-based "
     "on-site training for agents, supervisors and administrators plus train-the-trainer; "
     "UAT with Devoteam environment support, daily defect triage and same-cycle fixes.",
     "SIT Report\nTraining Manual\nUAT Report\n(UAT Sign-off - 40%)"),
    ("Data Migration", "Weeks 6-9",
     "Legacy assessment and profiling; field mapping and category cross-walk to the "
     "Appendix A taxonomy; cleansing and transformation; full dress-rehearsal migration into "
     "UAT with count, checksum and attachment reconciliation; reconciliation reviewed and "
     "signed; final delta sync of open cases executed at cutover.",
     "Migration Plan\nData Reconciliation Report (rehearsal and cutover)"),
    ("Cutover, Go-Live & Post-Support", "Week 10 + 30 days",
     "Production readiness review; cutover executed against the Production Deployment "
     "Checklist with a defined rollback decision point; channel cutover; go-live; DR restore "
     "test evidenced; 30-day hypercare with a dedicated engineering team, on-site for week "
     "one and for any P1, daily stand-ups and a live defect board; legacy platform "
     "decommission and licence termination.",
     "Production Deployment Checklist\nSystem Configuration Document (final)\nSystem "
     "Operation Manual\nHypercare Exit Report\n(Go-Live - 20%)"),
    ("Project Close Out", "End of 180-day warranty",
     "180-day warranty support from end of hypercare - defect resolution, AI prompt and KB "
     "optimisation, monthly calibration cycles; snag/punch list worked to closure; final "
     "knowledge transfer; transition into the managed service if contracted; close-out "
     "review.",
     "Project Completion Certificate\nProject Deliverables Snag/Punch List\nProject Close "
     "Document\n(Project Close Out - 10%)"),
]

# ---------------------------------------------------------------- 9.0 SLA
SLA_ROWS = [
    ("Critical (P1)",
     "System down, core function impaired, data loss/unavailable/not loading, no manual "
     "workaround.\nE.g. login failed; inbound/outbound not working or intermittent; AI "
     "faulty (missing disclaimer, wrong information replied, no ticket created).",
     "< 30 minutes", "< 2 hours",
     "Accepted. 24/7 on-call rotation; on-site attendance committed for every P1 per "
     "8.1.1; workaround established first, permanent fix after; incident report within 2 "
     "business days."),
    ("Major (P2)",
     "Core function impaired, no manual workaround available.\nE.g. agent status "
     "incorrect; AI transfer to live agent failed; escalation not working; round robin "
     "dysfunction; UI distorted but key data partially legible.",
     "< 1 hour", "< 4 hours",
     "Accepted. 24/7 on-call rotation."),
    ("Moderate (P3)",
     "Non-core function impaired.\nE.g. AI response incorrect; file attachment cannot "
     "open; cannot generate or download reports; report data incorrect; UI distorted but "
     "supplementary data partially legible.",
     "< 4 hours", "< 24 hours",
     "Accepted, within staffed business hours."),
    ("Minor (P4)",
     "Minor bug, cosmetic.\nE.g. UI distorted but still fully legible; UI enhancement.",
     "< 24 hours", "< 5 days",
     "Accepted, within staffed business hours."),
]

# ---------------------------------------------------------------- 10.0 payment
PAYMENT_ROWS = [
    ("Mobilization", "Project Kick Off", "15%", "Project plan", "Accepted"),
    ("Discovery & Architecture Blueprint", "Requirements Sign off", "15%",
     "Business Requirement Document; Technical Design Document", "Accepted"),
    ("Development & Test", "UAT Sign Off", "40%", "SIT Report; UAT Report", "Accepted"),
    ("Production Deployment", "Go-Live", "20%",
     "Production Deployment Checklist; System Configuration Document", "Accepted"),
    ("End of warranty period", "Project Close Out", "10%",
     "Project Completion Certificate; Project Deliverables Snag List / Punch List",
     "Accepted"),
]

# ---------------------------------------------------------------- clarifications
CLARIFICATIONS = [
    ("Support availability - section 9.0 vs 8.1.3",
     "Section 9.0 states \"Support must be available 24/7\" while 8.1.3 requires support "
     "contactable following PRO-NET business hours (Mon-Fri 08:00-17:30). These set "
     "different cost bases. Devoteam has proposed staffed business-hours support for all "
     "severities plus a 24/7 on-call rotation for P1/P2 and 24/7 automated monitoring. "
     "Please confirm whether full 24/7 staffed coverage across all severities is required, "
     "as it is priced separately."),
    ("Managed service term beyond warranty",
     "Sections 8.0 and 9.0 describe ongoing support obligations that extend past the "
     "180-day warranty, but section 10.0's payment schedule closes the project at end of "
     "warranty. Please confirm the intended term of the post-warranty managed service and "
     "whether it is to be priced within this RFP or contracted separately."),
    ("DMS and TSP API access",
     "Requirements 3.3.1, 4.43 and 4.47 depend on API specifications, sandbox and "
     "production credentials and an agreed field-level data contract for the Dealer "
     "Management System and the Vehicle Telematics System. Please confirm the systems, "
     "their API documentation availability, and the earliest date sandbox access can be "
     "granted. Please also confirm the response-time expectation for each, since "
     "requirement 4.47's 3-second budget is end-to-end and partly depends on these "
     "PRO-NET-owned systems."),
    ("Regulatory reporting formats",
     "Requirement 2.2.5 requires \"automated regulatory report formats\". Please confirm "
     "which regulatory bodies the CCMS must report to, the required formats and the "
     "submission frequency."),
    ("Corporate identity provider",
     "Requirement 7.3 requires MFA. Please confirm PRO-NET's corporate identity provider "
     "and the federation protocol available (for example Entra ID with SAML 2.0 or OIDC), "
     "and whether dealer users are to be federated or held locally."),
    ("Business hours discrepancy",
     "Section 8.1.3 states PRO-NET business hours as Mon-Fri 08:00-17:30, while Appendix B "
     "specifies customer-facing operating hours of Mon-Fri 08:30-17:30 and Sat-Sun/public "
     "holidays 09:00-17:00. Please confirm which applies to SLA measurement for customer "
     "cases, and whether weekend operating hours are in scope for agent staffing."),
    ("Telephony carrier and number portability",
     "Requirements 4.5-4.9 and 4.66 cover IVR, call routing, recording and queue "
     "abandonment on 1300-888-877. Please confirm the current telephony provider, whether "
     "the number is to be ported or trunked to the new platform, and who holds the carrier "
     "contract."),
    ("e.MAS app and SSI survey scope",
     "Appendix B's SSI Process runs inside the e.MAS mobile app (survey issued on day 11 "
     "post-delivery, 14-day expiry, dealer appeal workflow). Please confirm whether the "
     "CCMS is required to ingest SSI results, drive the survey dispatch, or replace the "
     "dealer appeal workflow, and which team owns the e.MAS app API."),
    ("Historical data migration scope",
     "Requirement 2.3.2 requires migration of historical complaints. Please confirm the "
     "source system, the volume and date range to migrate, whether attachments and call "
     "recordings are included, and whether the 7-year retention obligation at 4.84 applies "
     "retrospectively to migrated records."),
    ("Data residency",
     "The proposed production region is Singapore (asia-southeast1), the nearest "
     "full-service Google Cloud region. Please confirm whether PRO-NET has a data-residency "
     "requirement mandating storage within Malaysia, as this affects the region selection "
     "and the PDPA cross-border transfer assessment."),
]


# =====================================================================
#  BUILD
# =====================================================================
def build():
    doc = Document()

    # base style
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(9.5)
    st.font.color.rgb = DARK

    # landscape, narrow margins - the matrices need the width
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for attr in ("top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(0.55))
    sec.left_margin = sec.right_margin = Inches(0.5)

    # ---------------- cover
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VENDOR TECHNICAL RESPONSE")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = POPPY
    r.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Customer Complaint Management System Implementation")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = BLUE
    r.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RFP 2026_028")
    r.font.size = Pt(14)
    r.font.color.rgb = MID
    r.font.name = FONT

    doc.add_paragraph()
    for line, size, bold in [
        ("Prepared for", 10, False),
        ("PRO-NET  -  PROTON e.MAS", 15, True),
        ("{{TBD - PRO-NET full legal entity name for the contract}}", 9, False),
        ("", 9, False),
        ("Prepared by", 10, False),
        ("Devoteam G Cloud", 15, True),
        ("", 9, False),
        ("{{TBD - submission date}}   ·   Version 1.0   ·   Confidential", 9, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = DARK if bold else MID
        r.font.name = FONT

    doc.add_page_break()

    # ---------------- how to read
    heading(doc, "How to Read This Response")
    body(doc,
         "This document answers RFP 2026_028 section by section, in the RFP's own "
         "structure and using the RFP's own compliance vocabulary. Every row carries a "
         "compliance status and the proposed solution detail behind it.")
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    for i, h in enumerate(["Status", "Meaning in this response"]):
        c = t.rows[0].cells[i]
        shade(c, HDR_FILL)
        cell_text(c, h, size=9, bold=True, color=WHITE)
    for code, meaning in [
        ("Fully", "Fully out-of-the-box. The capability exists in the platform as it "
                  "stands today and is demonstrable before contract award. Configuration "
                  "to PRO-NET's data and rules is included."),
        ("Custom", "Customization or new build required, delivered within this engagement "
                   "and within the quoted price. Where part of a requirement is already "
                   "shipped, the row states exactly what exists today and what is being "
                   "built."),
        ("TP", "Delivered through a third-party component, named in the row."),
        ("Non", "Not complied with as literally specified. The reason is stated plainly "
                "and an alternative is offered."),
    ]:
        row = t.add_row()
        shade(row.cells[0], STATUS_FILL[code])
        cell_text(row.cells[0], code, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[1], meaning, size=9)
    set_widths(t, [1.1, 9.2])

    body(doc, "", space_after=4)
    body(doc,
         "Where a requirement is partially met, this response says so and separates what "
         "is shipped from what is being built. Where a requirement cannot be met as "
         "written - requirement 4.40, email recall - it is marked Non-compliant with the "
         "technical reason and the alternative that is delivered instead. Devoteam would "
         "rather state a limit at bid stage than discover it at UAT.",
         italic=True, color=MID)
    body(doc,
         "Items marked {{TBD - ...}} are open points requiring PRO-NET input or Devoteam "
         "internal confirmation before submission; they are listed in full in the "
         "Clarification Questions annex.",
         italic=True, color=MID)

    doc.add_page_break()

    # ---------------- 1.0
    heading(doc, "1.0  Executive Summary - Vendor Response")
    body(doc,
         "PRO-NET's stated problem is that the current system is limited in integration "
         "and dashboards and lacks major functions. Devoteam proposes to replace it with a "
         "self-managed, AI-native Customer Complaint Management System running on Google "
         "Cloud inside PRO-NET's own tenancy - built from an open-source CRM core "
         "(Chatwoot Community) and a Google Gemini / Vertex AI layer, configured directly "
         "to the five-channel complaint SOP in Appendix B and the case taxonomy in "
         "Appendix A.")
    body(doc, "Four points distinguish this response:")
    bullets(doc, [
        "No per-seat licensing. There is no per-agent software licence and no "
        "per-resolution AI charge. PRO-NET pays for infrastructure and for AI consumption "
        "at Google's published rates on its own billing account. Cost scales with usage, "
        "not with headcount - which is the structural cost problem in a per-seat SaaS CRM "
        "as a contact centre grows.",
        "Ownership. Conversations, customer data, the knowledge base, the AI prompts and "
        "the analytics warehouse all sit inside PRO-NET's Google Cloud project. There is "
        "no vendor lock on the data and no export negotiation at renewal.",
        "It already exists. This is not a greenfield build. The platform - omni-channel "
        "workspace, AI agent, voice/IVR with real-time transcription, rule-driven "
        "escalation with the 2h/8h/48h ladder, Customer 360, RSA logging, and the BigQuery "
        "and Power BI reporting layer - is running today and was demonstrated to PRO-NET on "
        "28 July 2026. This engagement is configuration, gap closure, integration and "
        "migration against a working system.",
        "Stated honestly. This response marks 4 requirements as new build and 1 as "
        "non-compliant, with reasons. The compliance column is intended to be accurate "
        "enough to hold at UAT.",
    ])
    body(doc,
         "Devoteam is a Google Cloud Premier Partner. Pricing is submitted separately from "
         "this technical response.",
         italic=True, color=MID)
    body(doc,
         "{{TBD - Devoteam corporate credentials paragraph (partner tier, regional "
         "certifications, automotive/contact-centre references) to be confirmed against "
         "current Devoteam collateral before submission.}}",
         color=POPPY)

    doc.add_page_break()

    # ---------------- 2.0
    heading(doc, "2.0  Scope of Work & Deliverables")
    matrix(doc, SEC2,
           ["Module", "Scope & Key Deliverables", "Compliance", "Proposed Solution Details"],
           [1.85, 2.4, 0.75, 5.3], status_col=2)

    doc.add_page_break()

    # ---------------- 3.0
    heading(doc, "3.0  Solution Overview & System Features")
    body(doc,
         "Requirements on management of inbound are answered in conjunction with Appendix "
         "B (WhatsApp, Social Media, Email, IVR Call and SSI process flows).",
         italic=True, color=MID)
    matrix(doc, SEC3,
           ["Module", "User System Requirement", "Compliance", "Proposed Solution Details"],
           [1.85, 2.4, 0.75, 5.3], status_col=2)

    doc.add_page_break()

    # ---------------- 4.0
    heading(doc, "4.0  System Requirements Specification (SRS) & Compliance Matrix")
    body(doc,
         "Requirements on management of inbound are read in conjunction with Appendix B; "
         "requirements on reports are read in conjunction with Appendix C1 (monthly) and "
         "C2 (weekly).",
         italic=True, color=MID)
    matrix(doc, SEC4,
           ["Req ID", "Module", "User System Requirement", "Compliance",
            "Proposed Solution Details"],
           [0.6, 1.35, 2.3, 0.75, 5.3], status_col=3)

    doc.add_page_break()

    # ---------------- 5.0
    heading(doc, "5.0  Project Timeline & Implementation Phases")
    body(doc,
         "Ten weeks from kick-off to go-live, followed by 30 days of hypercare and a "
         "180-day warranty period. Phases overlap where they safely can - migration runs "
         "alongside testing, and the AI and analytics workstream runs alongside the core "
         "build - which is what makes the timeline achievable without compressing UAT.")
    matrix(doc,
           [(a, b, c, d) for (a, b, c, d) in TIMELINE],
           ["Phase & Focus Area", "Timeline", "Key Deliverables & Activities",
            "Sign-Off Milestone"],
           [1.9, 0.95, 5.5, 1.95])
    body(doc, "", space_after=4)
    body(doc,
         "Assumptions this timeline depends on, stated so they can be challenged now "
         "rather than at week six:",
         size=9)
    bullets(doc, [
        "DMS and TSP sandbox credentials and API documentation are available by the end of "
        "week 2. This is the critical-path dependency; the connectors are stubbed so the "
        "rest of the build continues regardless, but the Customer 360 integration testing "
        "at 2.2.2 cannot complete without it.",
        "PRO-NET frontline agents and supervisors are available for UAT in weeks 7-9, and "
        "the UAT window is scheduled with the call centre at kick-off.",
        "The legacy data export is provided by the end of week 4 to allow the "
        "dress-rehearsal migration.",
        "FAQ and knowledge-base source content is provided by PRO-NET during weeks 2-4.",
        "Meta Business verification for the Facebook and Instagram channels is initiated in "
        "week 1; verification turnaround is controlled by Meta, not by either party.",
    ], size=9)
    body(doc, "", space_after=4)
    body(doc,
         "Scope note for PRO-NET's attention. This ten-week plan carries the full RFP "
         "scope including telephony/IVR, the DMS and TSP connectors, the Power BI report "
         "build and the agent-management and duplicate-detection builds. Devoteam can "
         "compress the core platform delivery to eight weeks by moving the DMS/TSP "
         "connectors and the e.MAS SSI integration into a second phase after go-live - this "
         "is offered as an option because those two items carry PRO-NET-side dependencies "
         "that neither party fully controls. The commercial submission prices both shapes.",
         size=9, italic=True)

    doc.add_page_break()

    # ---------------- 6.0
    heading(doc, "6.0  Project Management Methodology, Progress Reporting & Governance")
    matrix(doc, SEC6,
           ["Module", "Scope & Key Deliverables", "Compliance", "Proposed Solution Details"],
           [1.85, 2.4, 0.75, 5.3], status_col=2)

    doc.add_page_break()

    # ---------------- 7.0
    heading(doc, "7.0  License Control Mechanism & Enforcement Architecture")
    body(doc,
         "This section assumes a proprietary, seat-licensed product. The proposed solution "
         "is not one, so several rows answer the underlying control objective rather than "
         "the licence mechanism - and say so explicitly. The commercial consequence is in "
         "PRO-NET's favour and is set out at 7.1.",
         italic=True, color=MID)
    matrix(doc, SEC7,
           ["Module", "Scope & Key Deliverables", "Compliance", "Proposed Solution Details"],
           [1.85, 2.4, 0.75, 5.3], status_col=2)

    doc.add_page_break()

    # ---------------- 8.0
    heading(doc, "8.0  Aftersale Support & Service Level Agreement")
    body(doc,
         "Hypercare (30 days) and the post-production warranty (180 days) are included in "
         "this engagement and priced within it. The ongoing managed service described "
         "below applies after the warranty expires and is priced separately, per section "
         "2.4.3.",
         italic=True, color=MID)
    matrix(doc, SEC8,
           ["Module", "User System Requirement", "Compliance", "Proposed Solution Details"],
           [1.85, 2.4, 0.75, 5.3], status_col=2)

    doc.add_page_break()

    # ---------------- 9.0
    heading(doc, "9.0  Service Level Agreement - Acknowledgement")
    body(doc,
         "The severity model and the response and resolution targets set out in section "
         "9.0 are accepted in full and without qualification. They apply during hypercare, "
         "throughout the 180-day warranty, and under the managed service thereafter.")
    matrix(doc, SLA_ROWS,
           ["Severity Level", "Description & Operational Impact", "Target Response",
            "Target Resolution", "Devoteam Response"],
           [1.05, 3.6, 0.95, 0.95, 3.75])
    body(doc, "", space_after=4)
    body(doc,
         "Response time is measured from receipt of the incident report on any support "
         "channel to acknowledgement with an initial assessment and, where possible, a "
         "temporary workaround. Resolution time is measured to the permanent fix. Where a "
         "permanent fix requires a code change that cannot safely be deployed within the "
         "resolution window, a workaround is delivered inside the window and the permanent "
         "fix is scheduled to the next release with PRO-NET's agreement - this is stated "
         "here rather than discovered during a P1.",
         size=9)
    body(doc,
         "Note on availability: section 9.0's \"support must be available 24/7\" is "
         "addressed at 8.1.3 and raised as a clarification question.",
         size=9, italic=True, color=POPPY)

    doc.add_page_break()

    # ---------------- 10.0
    heading(doc, "10.0  Payment Milestone - Acknowledgement")
    body(doc,
         "The payment schedule in section 10.0 is accepted as structured. Each milestone "
         "is evidenced by the deliverables named against it in section 6.3.2 and produced "
         "by the phase named in section 5.0.")
    matrix(doc, PAYMENT_ROWS,
           ["Phase", "Milestone / Event", "Payment", "Key Deliverables", "Devoteam Response"],
           [2.1, 1.8, 0.8, 3.8, 1.8])

    doc.add_page_break()

    # ---------------- annex
    heading(doc, "Annex A  -  Clarification Questions")
    body(doc,
         "Submitted in accordance with the questions deadline of 10 August 2026. Each item "
         "either resolves a conflict between sections of the RFP and its appendices, or "
         "requests information that materially affects scope, schedule or price. Devoteam "
         "has proposed a working position for each in the body of this response so that "
         "the submission is complete either way.",
         italic=True, color=MID)
    matrix(doc,
           [(f"A{i}", q, a) for i, (q, a) in enumerate(CLARIFICATIONS, start=1)],
           ["#", "Subject", "Question"],
           [0.5, 2.6, 7.2])

    doc.add_paragraph()
    heading(doc, "Annex B  -  Open Items Before Submission", level=2)
    body(doc,
         "The following are Devoteam-internal and must be completed before this document "
         "is submitted:",
         size=9)
    bullets(doc, [
        "PRO-NET full legal entity name for the cover and contract references.",
        "Submission date on the cover.",
        "Devoteam named project team - Project Manager, Solution Architect, Engagement "
        "Lead (section 6.2.1).",
        "Devoteam named support organisation - Service Delivery Manager, Technical Lead, "
        "Technical Account Manager, both L3 escalation contacts, with emails (section "
        "8.1.1).",
        "Devoteam support email, hotline and WhatsApp number (section 8.1.2).",
        "Devoteam support portal URL (section 8.1.4).",
        "Devoteam corporate credentials paragraph, verified against current collateral "
        "(section 1.0).",
    ], size=9)

    doc.save(OUT)
    print(f"wrote: {OUT}")


if __name__ == "__main__":
    build()
